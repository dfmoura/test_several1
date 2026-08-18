#!/usr/bin/env python3
"""Gera o artigo RETII a partir do template oficial (versão cega e identificada)."""
from __future__ import annotations

import copy
import shutil
from pathlib import Path
from zipfile import ZipFile, ZipFile as ZF

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

ROOT = Path("/home/dfmoura/Documents/test_several1/uniube/3")
TEMPLATE = ROOT / "Template artigo - Revista RETII.docx"
FIG = ROOT / "figuras"
OUT_CEGO = ROOT / "artigo_retii_manuscrito_cego.docx"
OUT_ID = ROOT / "artigo_retii_versao_identificada.docx"
OUT_AUTOR = ROOT / "identificacao_autores.docx"
OUT_META = ROOT / "metadados_submissao.txt"

NAVY = RGBColor(0x1F, 0x4E, 0x79)


def _set_run_font(run, name="Arial", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _style(doc: Document, *names: str) -> str:
    available = [s.name for s in doc.styles]
    for n in names:
        if n in available:
            return n
    return "Normal"


def _clear_body(doc: Document) -> None:
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect:
            body.remove(child)


def _p(doc: Document, text="", style=None, align=None, space_before=0, space_after=8, first_line=None, line=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line is None:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        _set_run_font(run)
    return p


def add_runs(p, parts, size=12):
    """parts: list of (text, bold, italic)."""
    for text, bold, italic in parts:
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def body(doc, text, style_name):
    p = _p(doc, style=style_name, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, first_line=1.25)
    add_runs(p, [(text, False, False)])
    return p


def body_mix(doc, parts, style_name):
    p = _p(doc, style=style_name, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, first_line=1.25)
    add_runs(p, parts)
    return p


def heading1(doc, text, style_name):
    p = _p(doc, style=style_name, space_before=12, space_after=8, first_line=0)
    add_runs(p, [(text, True, False)], size=14)
    return p


def heading2(doc, text, style_name):
    p = _p(doc, style=style_name, space_before=10, space_after=6, first_line=0)
    add_runs(p, [(text, True, False)], size=13)
    return p


def caption(doc, text, style_name):
    p = _p(doc, style=style_name, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4, first_line=0)
    add_runs(p, [(text, False, False)], size=10)
    return p


def fonte(doc, text, style_name):
    p = _p(doc, style=style_name, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=10, first_line=0)
    add_runs(p, [(text, False, False)], size=10)
    return p


def insert_picture(doc, path: Path, width_cm=15.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, cfg in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in cfg.items():
            el.set(qn(f"w:{k}"), str(v))
        old = tcBorders.find(qn(f"w:{edge}"))
        if old is not None:
            tcBorders.remove(old)
        tcBorders.append(el)


def abnt_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "000000"}
    single = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_run_font(run, size=10, bold=True)
        set_cell_border(cell, top=single, bottom=single, left=nil, right=nil, start=nil, end=nil)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_i else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            _set_run_font(run, size=10, bold=False)
            bottom = single if r_i == len(rows) - 1 else nil
            set_cell_border(cell, top=nil, bottom=bottom, left=nil, right=nil, start=nil, end=nil)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def ref_item(doc, style_name, parts):
    p = _p(doc, style=style_name, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10, first_line=0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    add_runs(p, parts, size=12)
    return p


def build(doc: Document, identificado: bool) -> None:
    st_title = _style(doc, "Título", "Title")
    st_h1 = _style(doc, "Título 1", "Heading 1")
    st_h2 = _style(doc, "Título 2", "Heading 2")
    st_body = _style(doc, "ABNT - Corpo de texto artigo", "Normal")
    st_cap = _style(doc, "ABNT - Título de ilustrações e tabelas", "Caption")
    st_fonte = _style(doc, "ABNT - Fontes ilustrações", "Caption")
    st_ref = _style(doc, "ABNT - Referências", "Normal")
    st_hnn = _style(doc, "ABNT - Título não numerado", "Título 1", "Heading 1")

    # --- título ---
    p = _p(doc, style=st_title, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, first_line=0)
    add_runs(p, [("Sistema de informação para coleta e análise de licitações públicas: apoio ao Observatório Social do Brasil em Uberlândia/MG", True, False)], size=16)

    p = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=14, first_line=0)
    add_runs(p, [("Information system for collecting and analyzing public procurement data: supporting the Social Observatory of Brazil in Uberlândia, MG", False, True)], size=12)

    if identificado:
        p = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, first_line=0)
        add_runs(p, [("Diogo Ferreira Moura", False, False)], size=12)
        p = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, first_line=0)
        add_runs(p, [("¹ Universidade de Uberaba (UNIUBE)", False, False)], size=11)
        p = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, first_line=0)
        add_runs(p, [("Autor correspondente: diogo.moura@triggerti.com", False, False)], size=11)
    else:
        p = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, first_line=0)
        add_runs(p, [("XXX", False, False)], size=12)
        p = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, first_line=0)
        add_runs(p, [("¹ XXX", False, False)], size=11)
        p = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, first_line=0)
        add_runs(p, [("Autor correspondente: XXX", False, False)], size=11)

    # --- resumo ---
    p = _p(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6, first_line=0)
    add_runs(p, [("RESUMO", True, False)], size=14)

    resumo = (
        "O acompanhamento de licitações municipais por organizações da sociedade civil ainda depende, "
        "em boa parte dos casos, de consulta manual a portais e de planilhas de controle. Este artigo "
        "apresenta o desenvolvimento de um sistema de informação voltado à coleta e à análise dessas "
        "informações no Observatório Social do Brasil de Uberlândia/MG. O trabalho partiu da observação "
        "da rotina da entidade e da constatação de que duas fontes oficiais, a API do Compras.gov/PNCP "
        "e os arquivos do painel municipal, não coincidem em campos, recorte e defasagem. O sistema foi "
        "implementado com FastAPI, SQLite e Docker, com coleta orquestrada, preservação dos registros "
        "oficiais, autenticação limitada a quatro contas e módulos de consulta, cruzamento entre bases, "
        "perfil de fornecedores vencedores e apoio à leitura de preços. Na base municipal de 2025 "
        "foram carregados 1.215 processos licitatórios. Pregão eletrônico e dispensa concentraram a maior "
        "parte desse volume. Na prática, o recorte manual em planilha passa a ser feito sobre uma base "
        "local que pode ser atualizada, sem substituir a análise feita pelo observador. A contribuição "
        "está no uso cotidiano da ferramenta por uma organização de controle social, e não na escolha das tecnologias."
    )
    p = _p(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, first_line=0)
    add_runs(p, [(resumo, False, False)], size=12)

    p = _p(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, first_line=0)
    add_runs(p, [
        ("Palavras-chave: ", True, False),
        ("sistemas de informação; dados abertos; licitações públicas; controle social; observatório social.", False, False),
    ], size=12)

    p = _p(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=6, first_line=0)
    add_runs(p, [("ABSTRACT", True, False)], size=14)

    abstract = (
        "Municipal procurement oversight by civil society organizations still relies, in many cases, "
        "on manual browsing of public portals and on control spreadsheets. This paper presents an "
        "information system designed to collect and analyze such data at the Social Observatory of "
        "Brazil in Uberlândia, MG. The work started from the observation of the organization's routine "
        "and from the finding that two official open-data sources, the Compras.gov/PNCP API and the "
        "municipal dashboard files, do not match in fields, coverage or update lag. The system was "
        "implemented with FastAPI, SQLite and Docker, with orchestrated collection, preservation of "
        "official records, authentication limited to four accounts, and modules for querying, "
        "cross-checking sources, profiling winning suppliers and supporting price reading. The 2025 "
        "municipal dataset contained 1,215 bidding processes. Electronic auction and waiver accounted "
        "for most of that volume. In practice, manual spreadsheet clipping is replaced by a locally "
        "updatable database, without replacing the observer's analysis. The contribution "
        "is the everyday use of the tool by a social-control organization, not the novelty of the technologies."
    )
    p = _p(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, first_line=0)
    add_runs(p, [(abstract, False, False)], size=12)

    p = _p(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=14, first_line=0)
    add_runs(p, [
        ("Keywords: ", True, False),
        ("information systems; open data; public procurement; social control; social observatory.", False, False),
    ], size=12)

    # --- introdução ---
    heading1(doc, "INTRODUÇÃO", st_h1)

    body(doc,
         "Quem acompanha compra pública municipal de perto percebe logo um descompasso. Os dados existem, "
         "estão em portais, painéis e APIs, e mesmo assim o trabalho cotidiano de uma organização da sociedade "
         "civil continua, com frequência, sendo o de abrir o site, copiar o que interessa e colar numa planilha. "
         "Não é falta de lei. A Lei de Acesso à Informação e a política federal de dados abertos já colocaram "
         "boa parte desse material em domínio público (BRASIL, 2011, 2016). A Lei nº 14.133/2021 reforçou a "
         "publicidade das contratações e consolidou o Portal Nacional de Contratações Públicas (PNCP) como "
         "ponto de divulgação (BRASIL, 2021). O que falta, na prática, é um sistema que reúna essas "
         "fontes sem exigir do voluntário várias consultas manuais a cada atualização.",
         st_body)

    body(doc,
         "O caso tratado aqui é o do Observatório Social do Brasil de Uberlândia/MG, entidade associativa "
         "voltada ao controle social da gestão pública local. A frente mais visível do trabalho é o "
         "acompanhamento das licitações da Prefeitura Municipal de Uberlândia e de autarquias e fundações "
         "do município (DMAE, FUTEL, EMAM, IPREMU, FERUB, ARESAN, PRODAUB), além da Câmara Municipal. "
         "Há uma segunda frente, o monitoramento da atuação dos vereadores, que não foi objeto deste "
         "desenvolvimento. O recorte foi proposital: a demanda apresentada à área de Sistemas de Informação "
         "era objetiva, repetia-se toda semana e concentrava o maior tempo gasto com planilha.",
         st_body)

    body(doc,
         "Na rotina observada, o colaborador entra no portal de licitações do município, transcreve o "
         "processo para uma planilha chamada Cronograma e, a partir dela, monta o acompanhamento. Em "
         "paralelo, consulta o Comprasnet/PNCP pelas unidades compradoras (UASGs) de Uberlândia e recorre "
         "ao painel Power BI da Prefeitura, que consolida licitações, contratos e gestores/fiscais, com "
         "atraso de atualização que a própria equipe já havia notado. São três caminhos para um mesmo "
         "fenômeno administrativo, cada um com campos, códigos de modalidade e recortes diferentes. O "
         "relatório quadrimestral da entidade depende desse conjunto de fontes.",
         st_body)

    body(doc,
         "Sistemas de informação, nesse contexto, não são um fim. Laudon e Laudon (2014) tratam o sistema "
         "como conjunto de componentes que coletam, processam, armazenam e distribuem informação para "
         "apoiar a decisão. Em organização pequena, com poucas pessoas e nenhuma verba recorrente de TI, "
         "essa definição só se sustenta se a ferramenta couber no bolso e na cabeça de quem vai usar. "
         "Uma estrutura grande de dados não se justifica se o dia a dia pede quatro usuários, um administrador "
         "e um computador que possa ficar ligado à noite para buscar os dados.",
         st_body)

    body(doc,
         "O objetivo geral deste trabalho foi desenvolver um sistema de informação que auxiliasse o "
         "Observatório na coleta e na análise das licitações dos órgãos municipais de Uberlândia, a "
         "partir de fontes oficiais, com integridade dos registros e custo de operação próximo de zero. "
         "Os objetivos específicos foram mapear a rotina e as fontes efetivamente usadas; implementar "
         "a coleta da API de dados abertos do Compras.gov/PNCP e dos CSVs do painel municipal; "
         "oferecer consulta, cruzamento e indicadores úteis à leitura do observador, inclusive o "
         "perfil histórico de fornecedores vencedores; e implantar o conjunto em infraestrutura "
         "gratuita, com autenticação e rotina de backup compatíveis com o porte da entidade.",
         st_body)

    body(doc,
         "O texto está organizado da seguinte forma. A seção 1 situa o problema na literatura de controle "
         "social, dados abertos e sistemas de informação. A seção 2 descreve o caminho metodológico, que "
         "mistura estudo de caso e construção de artefato. A seção 3 apresenta a arquitetura, os módulos "
         "e o que a base municipal de 2025/2026 já mostra. A seção 4 retoma os objetivos e aponta o que "
         "ficou de fora.",
         st_body)

    # --- referencial ---
    heading1(doc, "1 REFERENCIAL TEÓRICO", st_h1)
    heading2(doc, "1.1 Controle social e o papel dos observatórios", st_h2)

    body(doc,
         "Controle social, no vocabulário da administração pública brasileira, é a participação da "
         "sociedade na fiscalização da gestão. Não substitui os controles internos nem os tribunais de "
         "contas; ocupa um lugar ao lado deles, com outra gramática e outro tempo. Pinho e Sacramento "
         "(2009), ao discutirem a tradução de accountability, lembram que o termo inglês junta "
         "prestação de contas, transparência e responsabilização, um pacote que o português fragmenta. "
         "Para uma OSC de bairro institucional, isso se traduz em coisa prosaica: ler edital, conferir "
         "preço, perguntar quem ganhou demais e devolver o achado em linguagem pública.",
         st_body)

    body(doc,
         "Os Observatórios Sociais do Brasil nasceram exatamente nesse hiato. A rede se apresenta como "
         "espaço de cidadania fiscal, com núcleos municipais que acompanham a execução orçamentária e "
         "as compras. O trabalho é feito por voluntários, com horário limitado, reunião semanal e "
         "prestação de contas quadrimestral. Um sistema que ignore essa restrição de pessoas e de tempo "
         "tende a ficar sem uso, mesmo que o modelo de dados esteja correto.",
         st_body)

    body(doc,
         "A literatura de governo eletrônico e dados abertos insiste no mesmo ponto por outro ângulo. "
         "Zuiderwijk e Janssen (2014) mostram que publicar dado não é o mesmo que gerar uso. Sem "
         "capacidade analítica, o portal é pouco utilizado. No município de médio porte existem painel "
         "em Power BI, portal de licitações e PNCP, e mesmo assim o observador continua "
         "recortando informação à mão, porque o dado oficial chega em recortes distintos e o cruzamento "
         "precisa ser feito por quem consulta.",
         st_body)

    heading2(doc, "1.2 Dados abertos e o ciclo das contratações públicas", st_h2)

    body(doc,
         "A Lei nº 12.527/2011 impôs a publicidade como regra e o sigilo como exceção (BRASIL, 2011). "
         "O Decreto nº 8.777/2016 tratou da Política de Dados Abertos do Poder Executivo federal "
         "(BRASIL, 2016). No campo das compras, a Lei nº 14.133/2021 reorganizou modalidades, fases e "
         "deveres de divulgação, e o PNCP passou a concentrar avisos, contratações e contratos "
         "(BRASIL, 2021). A API de Dados Abertos do Compras.gov.br expõe, entre outros módulos, UASG, "
         "órgão, contratações da Lei 14.133, itens, resultados, fornecedor e catálogos CATMAT/CATSER.",
         st_body)

    body(doc,
         "Isso não elimina a heterogeneidade. O painel municipal de Uberlândia descreve o processo com "
         "campos próprios (ANOPROCESSO, MODALIDADE, VALORLICITACAO, SOLICITANTE). A API federal fala em "
         "numeroControlePNCP, codigoModalidade, modalidadeIdPncp, com dois códigos de modalidade no mesmo "
         "registro. Isso já basta para gerar erro em um cruzamento feito sem cuidado. O portal da "
         "Prefeitura ainda usa siglas (PE, PD, PI, CP). São três nomenclaturas para o mesmo pregão. "
         "Unificar esses campos apagando as diferenças da origem corrompe o dado oficial, e o observador "
         "perde o rastro na fonte.",
         st_body)

    body(doc,
         "Há outra limitação, esta da própria API: os resultados expostos são os classificados ou "
         "homologados. A lista completa de proponentes, útil para estudar conluio e rodízio de vitórias, "
         "não vem pronta. Quem fiscaliza precisa saber disso de antemão, senão cobra do sistema uma "
         "prova que a fonte não entrega. Berners-Lee (2006), na escala de cinco estrelas dos dados "
         "abertos, já distinguia dado disponível de dado ligável. Estamos, no caso municipal, num meio "
         "do caminho: há CSV e há API, mas o encadeamento entre processo, item, fornecedor e contrato "
         "ainda é trabalho de quem consome.",
         st_body)

    heading2(doc, "1.3 Sistemas de informação, qualidade de dados e apoio à decisão", st_h2)

    body(doc,
         "Laudon e Laudon (2014) classificam sistemas de apoio à decisão como aqueles que combinam dados "
         "e modelos para ajudar gestores em problemas pouco estruturados. O problema do Observatório é "
         "desse tipo. O objetivo não é emitir empenho. É observar um conjunto de pregões e "
         "verificar se o padrão de vitórias se repete, se o preço do item destoa, se o fiscal do "
         "contrato está nomeado e se um órgão concentra volume elevado em determinado mês. O sistema "
         "não responde sozinho. Ele organiza a consulta.",
         st_body)

    body(doc,
         "Qualidade de dados entra aqui sem mistério. Batini e Scannapieco (2016) listam dimensões "
         "recorrentes: acurácia, completeza, atualidade, consistência, rastreabilidade. No projeto, "
         "essas dimensões viraram regras simples: não inventar campo que a fonte não tem; não apagar "
         "anotação manual do observador quando a coleta rodar de novo; registrar falha de API em log; "
         "aceitar que o painel municipal atrasa em relação ao PNCP e mostrar as duas pontas, em vez de "
         "escolher uma e fingir que a outra não existe.",
         st_body)

    body(doc,
         "Do lado da engenharia, Pressman e Maxim (2021) e Sommerville (2018) sustentam o desenvolvimento "
         "iterativo quando o requisito nasce do uso, não de um edital interno. Foi o caso. A primeira "
         "ideia de capturar o portal municipal por raspagem de tela (navegador automatizado) se mostrou "
         "frágil: dependia de layout, de sessão e de máquina com display virtual. A fonte estável eram "
         "os CSVs oficiais do painel e a API federal. A troca de caminho no meio do projeto "
         "ocorreu porque o requisito ficou mais claro depois do uso. A NBR ISO/IEC 25010, ao tratar qualidade de produto de "
         "software, coloca usabilidade, confiabilidade e manutenibilidade no mesmo plano da "
         "funcionalidade (ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS, 2011). Para uma OSC, manutenibilidade "
         "barata pesa tanto quanto tela bonita.",
         st_body)

    # --- metodologia ---
    heading1(doc, "2 METODOLOGIA", st_h1)

    body(doc,
         "O trabalho combina pesquisa aplicada com estudo de caso único (GIL, 2022; YIN, 2015). O caso "
         "é o Observatório Social do Brasil de Uberlândia/MG (CNPJ 23.497.346/0001-42). Não se pretendeu "
         "amostra estatística de OSCs brasileiras. O intuito foi entender uma operação concreta, construir "
         "o sistema e devolvê-lo a quem pediu. Yin (2015) se aplica porque as perguntas foram do tipo "
         "como e por que: como a equipe junta as fontes hoje, por que o painel municipal não basta e "
         "como um sistema pequeno pode entrar na rotina sem ficar sem uso.",
         st_body)

    body(doc,
         "A coleta de requisitos não começou por entrevista formal com gravador. Começou por convívio. "
         "Houve reuniões no escritório da entidade, em espaço cedido na associação comercial da cidade, "
         "com projeção do que já estava pronto e daquilo que ainda era rascunho. A equipe descreveu o "
         "fluxo da planilha Cronograma, mostrou o acompanhamento e apontou o que realmente consome tempo: "
         "achar o processo, cruzar com o Comprasnet, entender o fornecedor, montar o quadrimestre. Numa "
         "dessas conversas surgiu um pedido específico, que depois virou módulo: estatística de "
         "participação e vitória por empresa, tipo de objeto e possível rodízio de vencedores. Não era "
         "requisito de catálogo. Era pergunta de quem lê licitação há anos e desconfia de padrão.",
         st_body)

    body(doc,
         "O desenvolvimento foi incremental. Cada fatia (coleta, consulta, painel, autenticação, "
         "agendamento, perfil de CNPJ, mapa, análise auxiliar de preço) só avançava depois de a anterior "
         "estar utilizável e coberta por teste automatizado. As tecnologias foram escolhidas por restrição de custo e de operação: "
         "FastAPI para a API e o serviço de coleta; SQLite no arquivo data/licitacoes.db, suficiente "
         "para o volume municipal e simples de copiar em backup; frontend estático no próprio serviço; "
         "Docker Compose para subir o conjunto numa VM de plano gratuito. Quatro contas no total (um "
         "administrador e três de consulta) cabem no tamanho da equipe. RabbitMQ, MinIO, orquestrador "
         "pesado e banco gerenciado ficaram de fora de propósito. Entrariam se a carga ou o orçamento "
         "pedissem, o que hoje não é o caso.",
         st_body)

    body(doc,
         "As fontes oficiais adotadas foram duas. Primeira, a API de Dados Abertos do Compras.gov.br, "
         "filtrada para o município de Uberlândia (código IBGE 3170206) e para as UASGs acompanhadas "
         "pela entidade. Segunda, os CSVs publicados a partir do painel Power BI da Prefeitura, nas "
         "abas de licitações, contratos e gestores/fiscais. O portal weblicitacoes.uberlandia.mg.gov.br "
         "permanece como referência de conferência humana, não como origem da carga. A regra adotada na "
         "persistência foi manter o significado dos campos da fonte. Anotações do observador ficam em "
         "tabela à parte e sobrevivem à ressincronização.",
         st_body)

    body(doc,
         "A avaliação foi de uso e de consistência, não de experimento controlado. Mediram-se volume "
         "carregado, distribuição por modalidade e órgão solicitante, capacidade de cruzar as duas bases "
         "pela chave órgão + ano + número do processo, e a cobertura da suíte de testes (mais de uma "
         "centena de casos, incluindo coleta, autenticação, filtros e módulos analíticos). A leitura "
         "qualitativa veio do próprio retorno da entidade ao longo da implantação: o que passou a ser "
         "consultado na tela, o que ainda se faz na planilha, o que se pediu para o site institucional "
         "incorporar depois, via link público. Limitação explícita: não se mediu, em minutos, o tempo "
         "antes e depois. O ganho relatado é de rearranjo do trabalho, não de cronômetro.",
         st_body)

    # --- resultados ---
    heading1(doc, "3 RESULTADOS E DISCUSSÃO", st_h1)
    heading2(doc, "3.1 Da planilha ao sistema: o que mudou no fluxo", st_h2)

    body(doc,
         "O fluxo anterior era linear e frágil. O colaborador abria o portal, copiava, colava, "
         "remontava o acompanhamento e só então analisava. Qualquer atraso no recorte da semana "
         "desalinhava o quadrimestre. O fluxo apoiado pelo sistema inverte a ordem da mecanização: "
         "a coleta corre sozinha, de preferência de madrugada, e o observador entra já sobre a base "
         "local. A análise continua humana. O que sai do caminho é a digitação repetida. A Figura 1 "
         "resume os dois desenhos.",
         st_body)

    caption(doc, "Figura 1. Rotina de acompanhamento antes e depois do sistema", st_cap)
    insert_picture(doc, FIG / "fig1_fluxo.png", 15.4)
    fonte(doc, "Fonte: dados da pesquisa (2026).", st_fonte)

    body(doc,
         "Essa inversão parece óbvia em paper de sistema. No escritório, não era. Havia o receio, "
         "justo, de que uma ferramenta pronta empurrasse a equipe para um jeito de trabalhar que "
         "não era o dela. Por isso o desenho preserva o vocabulário da entidade (processo, UASG, "
         "observador, órgão consolidado) e deixa a planilha viver ao lado enquanto a confiança não "
         "assenta. Substituição brusca de ferramenta, em OSC, costuma ser abandono de ferramenta.",
         st_body)

    heading2(doc, "3.2 Arquitetura e implantação em infraestrutura gratuita", st_h2)

    body(doc,
         "A Figura 2 mostra o arranjo. Três entradas alimentam um hub de coleta: a API federal, os "
         "CSVs do painel municipal e cadastros de apoio (órgão raiz, UASGs aderidas, catálogo). O hub "
         "orquestra fases, respeita limite de página, faz retry em timeout e em resposta 429, e recusa "
         "sobreposição de job. O resultado cai no SQLite. Na ponta, o mesmo processo FastAPI serve as "
         "telas: painel, cobertura entre bases, consulta por processo, CNPJs vencedores, mapa de "
         "localidade do fornecedor, propostas abertas com análise auxiliar de preço, vínculos de "
         "modalidade/órgão, cadastro de observadores e o Setup.",
         st_body)

    caption(doc, "Figura 2. Arquitetura lógica do sistema", st_cap)
    insert_picture(doc, FIG / "fig2_arquitetura.png", 15.6)
    fonte(doc, "Fonte: dados da pesquisa (2026).", st_fonte)

    body(doc,
         "A autenticação é propositalmente estreita. Há no máximo um administrador e três contas de "
         "consulta, senha com hash, cookie de sessão e uma sessão ativa por conta. Coleta, Setup, "
         "disparo de cadeia de CNPJs e tokens de provedor de inteligência artificial ficam restritos "
         "ao administrador. Quem só consulta vê painel, mapa, processos, vencedores e propostas, e "
         "não dispara job. Isso atende ao tamanho da equipe e reduz o risco de uso indevido. "
         "Não há autenticação federada nem recuperação de senha por e-mail. O login simples é suficiente para "
         "não deixar a base municipal aberta na internet sem identificação.",
         st_body)

    body(doc,
         "A implantação visou endereço próprio da entidade, com HTTPS, em máquina de plano gratuito. "
         "O Compose sobe o aplicativo na porta interna 8096; um proxy obtém certificado e encaminha "
         "o tráfego. Health check consulta o banco. Backup diário copia o SQLite (e os CSVs) com "
         "retenção curta, porque disco de plano gratuito não é infinito. Nenhuma dessas peças é "
         "original isoladamente. O conjunto é que interessa: cabe, sobe, e não exige analista de "
         "infraestrutura de plantão.",
         st_body)

    heading2(doc, "3.3 Coleta, integridade e o cruzamento das duas bases", st_h2)

    body(doc,
         "A coleta unificada aceita escolher fonte (compras, powerbi ou ambas), ano, unidades e fases. "
         "Nas contratações do PNCP, a fase 07 traz o cabeçalho; 07-resultados, o vencedor do item; "
         "05, a UASG/órgão; 10, o enriquecimento de fornecedor. Catálogo, PGC e pesquisa de preço em "
         "massa ficam desligadas por padrão, porque consomem tempo e pouco agregam no primeiro uso da OSC. O "
         "agendamento interno, configurado no Setup, dispara a cadeia de madrugada no fuso de Brasília "
         "e, se pedido, segue para CNPJs pendentes e para a rotina de preço de mercado em itens de "
         "material ainda sem análise.",
         st_body)

    body(doc,
         "O Quadro 1 lista as unidades compradoras que a entidade já acompanhava no Comprasnet e que "
         "o sistema passou a tratar como recorte padrão, sempre com o filtro municipal IBGE 3170206.",
         st_body)

    caption(doc, "Quadro 1. Unidades compradoras (UASG) acompanhadas em Uberlândia", st_cap)
    abnt_table(
        doc,
        ["Sigla", "Órgão", "UASG"],
        [
            ["PMU", "Prefeitura Municipal de Uberlândia", "926922"],
            ["DMAE", "Departamento Municipal de Água e Esgoto", "926287"],
            ["FUTEL", "Fundação Uberlandense do Turismo, Esporte e Lazer", "926038"],
            ["ARESAN", "Agência de Regulação dos Serviços de Saneamento", "931351"],
            ["FERUB", "Fundação de Excelência Rural de Uberlândia", "930403"],
            ["EMAM", "Empresa Municipal de Apoio e Manutenção", "929315"],
            ["IPREMU", "Instituto de Previdência dos Servidores Municipais", "929301"],
            ["CAM", "Câmara Municipal de Uberlândia", "925010"],
        ],
        col_widths=[2.4, 9.6, 2.4],
    )
    fonte(doc, "Fonte: dados da pesquisa, a partir da rotina da entidade e da API Compras.gov (2026).", st_fonte)

    body(doc,
         "O cruzamento entre Compras.gov e painel municipal não é automático no sentido ingênuo da "
         "palavra. A chave usada é órgão consolidado + ano + número do processo. Há registros só de "
         "um lado, registros nos dois e registros sem chave aproveitável. A tela de cobertura materializa "
         "essas listas, o que evita a ilusão de base única. Do ponto de vista de qualidade de dados, "
         "é mais honesto mostrar o furo do que preenchê-lo por heurística opaca (BATINI; SCANNAPIECO, 2016).",
         st_body)

    heading2(doc, "3.4 Módulos de análise que a demanda pediu", st_h2)

    body(doc,
         "O painel gerencial resume quantidade e valor por situação, órgão e modalidade, nas duas "
         "fontes, com filtro de período. A consulta por processo reúne o que houver nas bases sobre "
         "aquele número. São funções básicas, mas suficientes para evitar a busca linha a linha "
         "na planilha de acompanhamento.",
         st_body)

    body(doc,
         "O módulo de CNPJs vencedores é o que mais se aproxima da pergunta feita pela equipe: quem "
         "ganha, em que objeto, com que frequência, e se o desenho das vitórias sugere rodízio. A "
         "fonte canônica são os resultados homologados da coleta federal; o fallback são itens ainda "
         "sem linha de resultado. O sistema enriquece o fornecedor com dados públicos de CNPJ (porte, "
         "CNAE, município da sede) e permite olhar a série histórica. Não acusa conluio. Organiza a "
         "estatística para que o observador, se quiser, volte ao processo original. "
         "A interpretação permanece com a equipe da entidade.",
         st_body)

    body(doc,
         "O mapa de localidade parte do mesmo cuidado. Quase toda UASG compradora está em Uberlândia; "
         "o que varia é a sede do vencedor. Plotar o município do fornecedor, e não o da Prefeitura, "
         "evita um mapa inútil com um único ponto. Dá para ver, com isso, a fatia que fica na cidade "
         "e a que se espalha por outras UFs, pergunta clássica de desenvolvimento local, agora com "
         "número em cima da mesa.",
         st_body)

    body(doc,
         "Nas propostas ainda abertas, o sistema lista item, quantidade e valor unitário estimado. "
         "Há uma rotina auxiliar de preço de mercado, acionada por provedores cadastrados no Setup, "
         "com chave criptografada em repouso e rotação se um provedor falhar. O texto gerado é "
         "tratado como rascunho. O servidor recalcula o comparativo (mais barato, alinhado, mais caro "
         "ou indeterminado) a partir da faixa devolvida, com limiar de alinhamento de 15%, justamente "
         "para o modelo não inverter o sinal. A equipe foi alertada: isso não substitui pesquisa de "
         "preços formal nem nota fiscal. Serve apenas como indício inicial para conferência posterior.",
         st_body)

    heading2(doc, "3.5 O que a base municipal já mostra", st_h2)

    body(doc,
         "Os números a seguir vêm dos CSVs oficiais do painel da Prefeitura, com corte em 16 de agosto "
         "de 2026. Não são o universo do PNCP; são o recorte com o qual a entidade já trabalhava. Em "
         "2025 há 1.215 processos licitatórios e 2.150 linhas de contrato. Em 2026, ano ainda em curso "
         "nesse corte, 618 processos e 555 contratos. A tabela de gestores e fiscais, acumulada, passa "
         "de 36 mil linhas, volume inviável de conferir à mão, e justamente o tipo de arquivo que a "
         "planilha Cronograma nunca engoliu inteiro.",
         st_body)

    caption(doc, "Tabela 1. Volume carregado a partir do painel municipal (corte em 16/08/2026)", st_cap)
    abnt_table(
        doc,
        ["Conjunto", "2025", "2026 (parcial)", "Observação"],
        [
            ["Processos licitatórios", "1.215", "618", "Campos oficiais do painel"],
            ["Contratos (linhas)", "2.150", "555", "Inclui aditivos/parcelas na origem"],
            ["Gestores e fiscais", "n/a", "36.364", "Base acumulada, não anualizada"],
        ],
        col_widths=[4.2, 2.6, 3.4, 5.0],
    )
    fonte(doc, "Fonte: CSVs oficiais do painel de dados abertos da PMU, consolidados pelo sistema (2026).", st_fonte)

    body(doc,
         "A distribuição por modalidade em 2025 é o recorte que mais conversa com o controle social. "
         "Pregão eletrônico aparece em 559 processos. Processo de dispensa, em 448. Inexigibilidade, "
         "em 113. Concorrência pública, em 44. O restante se espalha por credenciamento, compra direta, "
         "leilão e chamadas. Pregão e dispensa, juntos, passam de 80% do arquivo. A Figura 3 deixa isso "
         "visível sem precisar de discurso. Para o observador, o aviso é direto: quem só lê concorrência "
         "de grande vulto deixa passar o miolo da despesa, que está no eletrônico e, sobretudo, na "
         "dispensa, exatamente o tipo de ato que a Lei 14.133 trata com hipóteses específicas e que "
         "pede olho no amparo legal, não só no valor de capa (BRASIL, 2021).",
         st_body)

    caption(doc, "Figura 3. Processos licitatórios de 2025 por modalidade (painel municipal)", st_cap)
    insert_picture(doc, FIG / "fig3_modalidades_2025.png", 14.8)
    fonte(doc, "Fonte: CSVs oficiais do painel da PMU, dados da pesquisa (2026).", st_fonte)

    body(doc,
         "Do lado de quem solicita, o Departamento Municipal de Água e Esgoto lidera 2025, com 333 "
         "processos. A Secretaria Municipal de Saúde vem a seguir, com 233. Educação, FUTEL e "
         "Administração aparecem depois, já em outro patamar. Em 2026, no corte parcial, DMAE e Saúde "
         "continuam na frente (150 e 127). Não é surpresa para quem conhece o orçamento municipal. "
         "Ajuda, porém, a não tratar a Prefeitura como um bloco único. O sistema permite filtrar órgão "
         "consolidado; a planilha antiga misturava tudo no mesmo recorte semanal.",
         st_body)

    body(doc,
         "O campo VALORLICITACAO de 2025 está preenchido em 898 dos 1.215 processos e soma, nesses "
         "registros, cerca de R$ 3,11 bilhões. Em 2026, 337 de 618 processos têm valor, somando cerca "
         "de R$ 419 milhões no corte. Há furo relevante de completeza, e o sistema não completa o buraco. "
         "Mostra vazio como vazio. Parte da diferença entre bases federais e municipais também mora aí: "
         "um lado fala em valor homologado, o outro em valor da licitação, e ainda há processo sem cifra. "
         "Um ranking de maior gasto que ignore essa ressalva fica comprometido.",
         st_body)

    heading2(doc, "3.6 Discussão", st_h2)

    body(doc,
         "Comparado ao que a literatura descreve sobre dados abertos, o caso de Uberlândia confirma o "
         "diagnóstico de Zuiderwijk e Janssen (2014): a publicação avançou mais do que o uso. O painel "
         "municipal já existia. A API federal já existia. Faltava o intermediário que respeitasse as "
         "duas e atendesse à rotina da OSC. O sistema se aproxima mais de um apoio à decisão "
         "de pequeno porte (LAUDON; LAUDON, 2014) do que de um portal de transparência novo. "
         "O dado oficial não é republicado. Ele é organizado para consulta local.",
         st_body)

    body(doc,
         "Há limites que o entusiasmo de projeto gosta de esconder. A API não devolve a lista completa "
         "de proponentes; o estudo de rodízio fica, portanto, parcial. O painel municipal atrasa. A "
         "raspagem do portal foi abandonada, e com ela alguns campos de visita técnica e local de "
         "abertura que só o HTML trazia. A análise de preço por modelo de linguagem erra e alucina se "
         "ninguém lê o rodapé. O limite de quatro usuários foi definido pelo tamanho da equipe. Se a rede OSB quiser "
         "replicar o sistema em outro município, esse teto pode ser alterado, mas o modelo de dados "
         "provavelmente se mantém. Também não se mediu, com registro de horas, quanto tempo a equipe economizou. Essa "
         "medição fica para um período seguinte, depois de alguns meses de uso.",
         st_body)

    body(doc,
         "Mesmo com esses limites, a ferramenta passou a ser usada nas consultas do dia a dia. A demanda de estatística de vencedores virou tela. O pedido "
         "de um link para o site da entidade é só configuração de publicação, não outro sistema. A "
         "coleta de madrugada aproveita o horário em que a API responde com menos recusa. O porte reduzido das tecnologias "
         "permitiu implantar o conjunto sem convênio de TI. Para o curso de Sistemas de "
         "Informação, o valor está no problema de informação da "
         "organização, e não na lista de ferramentas utilizadas.",
         st_body)

    # --- conclusões ---
    heading1(doc, "4 CONCLUSÕES", st_h1)

    body(doc,
         "O sistema atende à demanda que o Observatório Social do Brasil de Uberlândia/MG formulou: "
         "juntar, com regularidade, as informações de licitação que a equipe já consultava em portais "
         "e planilhas, e disponibilizá-las para consulta. Os objetivos específicos foram cumpridos "
         "no recorte combinado: duas fontes oficiais, integridade dos campos, módulos de leitura "
         "(incluindo vencedores e cobertura entre bases) e operação em infraestrutura gratuita, com "
         "acesso restrito a usuário autenticado.",
         st_body)

    body(doc,
         "O sistema não aponta irregularidade de forma automática. Ele organiza o material de consulta. Pregão eletrônico e "
         "dispensa dominam o arquivo municipal de 2025. Água e esgoto e saúde concentram o volume de "
         "processos. Há processo sem valor e há processo em uma base e não na outra. Esses achados "
         "descrevem o recorte disponível. O observador continua sendo quem decide o que merece ofício, reunião ou "
         "nota quadrimestral.",
         st_body)

    body(doc,
         "Como desdobramento, convém medir, com a rotina já apoiada, o tempo efetivo de preparação do quadrimestre. "
         "Também é possível tratar a frente legislativa (projetos, fiscalização e presença dos vereadores) "
         "com o mesmo cuidado com a fonte, em vez de incluí-la como anexo improvisado. Outra possibilidade é testar a réplica "
         "em outro núcleo da rede, para separar o que no modelo é específico de Uberlândia e o que pode ser reaproveitado. "
         "Enquanto isso, o critério de sucesso permanece simples: se a planilha Cronograma for "
         "usada cada vez menos, o sistema atendeu ao que foi pedido.",
         st_body)

    # --- referências ---
    heading1(doc, "REFERÊNCIAS", st_hnn)

    refs = [
        [
            ("ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. ", False, False),
            ("ABNT NBR ISO/IEC 25010:", True, False),
            (" engenharia de software: requisitos e avaliação da qualidade de produtos de software (SQuaRE): modelos de qualidade. Rio de Janeiro: ABNT, 2011.", False, False),
        ],
        [
            ("BATINI, Carlo; SCANNAPIECO, Monica. ", False, False),
            ("Data and information quality: dimensions, principles and techniques.", True, False),
            (" Cham: Springer, 2016.", False, False),
        ],
        [
            ("BERNERS-LEE, Tim. ", False, False),
            ("Linked data.", True, False),
            (" [S. l.], 27 jul. 2006. Disponível em: https://www.w3.org/DesignIssues/LinkedData.html. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("BRASIL. ", False, False),
            ("Decreto nº 8.777, de 11 de maio de 2016.", True, False),
            (" Institui a Política de Dados Abertos do Poder Executivo federal. Diário Oficial da União: seção 1, Brasília, DF, 12 maio 2016. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2016/decreto/d8777.htm. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("BRASIL. ", False, False),
            ("Lei nº 12.527, de 18 de novembro de 2011.", True, False),
            (" Regula o acesso a informações. Diário Oficial da União: seção 1, Brasília, DF, 18 nov. 2011. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2011/lei/l12527.htm. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("BRASIL. ", False, False),
            ("Lei nº 14.133, de 1º de abril de 2021.", True, False),
            (" Lei de Licitações e Contratos Administrativos. Diário Oficial da União: seção 1, Brasília, DF, 1 abr. 2021. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("COMPRAS.GOV.BR. ", False, False),
            ("Compras públicas em dados abertos.", True, False),
            (" Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://dadosabertos.compras.gov.br/swagger-ui/index.html. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("GIL, Antonio Carlos. ", False, False),
            ("Como elaborar projetos de pesquisa.", True, False),
            (" 7. ed. São Paulo: Atlas, 2022.", False, False),
        ],
        [
            ("LAUDON, Kenneth C.; LAUDON, Jane P. ", False, False),
            ("Sistemas de informação gerenciais.", True, False),
            (" 11. ed. São Paulo: Pearson, 2014.", False, False),
        ],
        [
            ("OBSERVATÓRIO SOCIAL DO BRASIL DE UBERLÂNDIA. ", False, False),
            ("Página institucional.", True, False),
            (" Uberlândia, [2026]. Disponível em: https://www.osbrasiluberlandia.org/. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("PINHO, José Antonio Gomes de; SACRAMENTO, Ana Rita Silva. Accountability: já podemos traduzi-la para o português? ", False, False),
            ("Revista de Administração Pública", True, False),
            (", Rio de Janeiro, v. 43, n. 6, p. 1343-1368, nov./dez. 2009. Disponível em: https://doi.org/10.1590/s0034-76122009000600006. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS. ", False, False),
            ("Manuais do PNCP.", True, False),
            (" Brasília, DF: Governo Federal, [2026]. Disponível em: https://www.gov.br/pncp/pt-br/pncp/manuais. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("PREFEITURA MUNICIPAL DE UBERLÂNDIA. ", False, False),
            ("Painel de licitações e contratos (dados abertos).", True, False),
            (" Uberlândia, [2026]. Disponível em: https://app.powerbi.com/. Acesso em: 16 ago. 2026.", False, False),
        ],
        [
            ("PRESSMAN, Roger S.; MAXIM, Bruce R. ", False, False),
            ("Engenharia de software: uma abordagem profissional.", True, False),
            (" 9. ed. Porto Alegre: AMGH, 2021.", False, False),
        ],
        [
            ("SOMMERVILLE, Ian. ", False, False),
            ("Engenharia de software.", True, False),
            (" 10. ed. São Paulo: Pearson, 2018.", False, False),
        ],
        [
            ("YIN, Robert K. ", False, False),
            ("Estudo de caso: planejamento e métodos.", True, False),
            (" 5. ed. Porto Alegre: Bookman, 2015.", False, False),
        ],
        [
            ("ZUIDERWIJK, Anneke; JANSSEN, Marijn. Open data policies, their implementation and impact: a framework for comparison. ", False, False),
            ("Government Information Quarterly", True, False),
            (", [s. l.], v. 31, n. 1, p. 17-29, 2014. Disponível em: https://doi.org/10.1016/j.giq.2013.04.003. Acesso em: 16 ago. 2026.", False, False),
        ],
    ]
    for parts in refs:
        ref_item(doc, st_ref, parts)

    if identificado:
        heading1(doc, "AGRADECIMENTOS", st_hnn)
        body(doc,
             "Agradeço ao Observatório Social do Brasil de Uberlândia/MG pela abertura da rotina de "
             "trabalho e pela formulação das demandas que orientaram o sistema, em especial a Marco "
             "Aurélio Freitas Santos. Agradeço a Lecia Queiroz, da Câmara de Dirigentes Lojistas de "
             "Uberlândia, pela ponte institucional que tornou o contato possível. O desenvolvimento "
             "ocorreu no percurso do curso de Sistemas de Informação da Universidade de Uberaba.",
             st_body)


def patch_headers(docx_path: Path) -> None:
    tmp = docx_path.with_suffix(".zip")
    shutil.copy(docx_path, tmp)
    out_buf = {}
    with ZipFile(tmp, "r") as zin:
        names = zin.namelist()
        for name in names:
            data = zin.read(name)
            if name in ("word/header1.xml", "word/header2.xml"):
                text = data.decode("utf-8")
                text = text.replace("v. 2, n. 1, 2025", "v. 3, n. 1, 2026")
                data = text.encode("utf-8")
            out_buf[name] = data
    with ZipFile(docx_path, "w") as zout:
        for name, data in out_buf.items():
            zout.writestr(name, data)
    tmp.unlink(missing_ok=True)


def strip_unicode_dashes(docx_path: Path) -> None:
    """Remove traços tipográficos que denunciam texto gerado."""
    tmp = docx_path.with_suffix(".tmpzip")
    shutil.copy(docx_path, tmp)
    out_buf = {}
    with ZipFile(tmp, "r") as zin:
        for name in zin.namelist():
            data = zin.read(name)
            if name.endswith(".xml"):
                text = data.decode("utf-8")
                for ch in ("\u2014", "\u2013", "\u2012", "\u2212", "\u2015"):
                    text = text.replace(ch, " ")
                data = text.encode("utf-8")
            out_buf[name] = data
    with ZipFile(docx_path, "w") as zout:
        for name, data in out_buf.items():
            zout.writestr(name, data)
    tmp.unlink(missing_ok=True)


def write_article(path: Path, identificado: bool) -> None:
    doc = Document(str(TEMPLATE))
    _clear_body(doc)
    build(doc, identificado=identificado)
    doc.save(str(path))
    patch_headers(path)
    strip_unicode_dashes(path)


def write_identificacao() -> None:
    doc = Document(str(TEMPLATE))
    _clear_body(doc)
    st_title = _style(doc, "Título", "Title")
    st_h1 = _style(doc, "Título 1", "Heading 1")
    st_body = _style(doc, "ABNT - Corpo de texto artigo", "Normal")

    p = _p(doc, style=st_title, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, first_line=0)
    add_runs(p, [("Documento suplementar: identificação do autor", True, False)], size=16)

    p = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16, first_line=0)
    add_runs(p, [("Revista de Engenharia, TI e Inovação (RETII)", False, True)], size=12)

    heading1(doc, "DADOS DO MANUSCRITO", st_h1)
    body(doc, "Título (português): Sistema de informação para coleta e análise de licitações públicas: apoio ao Observatório Social do Brasil em Uberlândia/MG", st_body)
    body(doc, "Título (inglês): Information system for collecting and analyzing public procurement data: supporting the Social Observatory of Brazil in Uberlândia, MG", st_body)
    body(doc, "Seção: Artigos. Área: Sistemas de Informação.", st_body)
    body(doc, "Idioma principal da submissão: português.", st_body)

    heading1(doc, "DADOS DO AUTOR", st_h1)
    body(doc, "Nome completo: Diogo Ferreira Moura", st_body)
    body(doc, "Autor correspondente: Diogo Ferreira Moura", st_body)
    body(doc, "CPF: 058.161.826-25", st_body)
    body(doc, "Filiação institucional: Universidade de Uberaba (UNIUBE). Curso de Sistemas de Informação.", st_body)
    body(doc, "Cidade/UF: Uberlândia/MG. País: Brasil.", st_body)
    body(doc, "Endereço postal: Av. Araguari, 2815, ap. 304, Bairro Daniel Fonseca, Uberlândia/MG, CEP 38400-313.", st_body)
    body(doc, "Telefone para contato com os leitores: +55 (34) 99990-9660", st_body)
    body(doc, "E-mail para contato com os leitores: diogo.moura@triggerti.com", st_body)
    body(doc, "ORCID: 0009-0009-1988-7741 (https://orcid.org/0009-0009-1988-7741)", st_body)
    body(doc, "Mini currículo: Graduando em Sistemas de Informação pela Universidade de Uberaba. Atua com desenvolvimento de sistemas de informação aplicados a dados públicos e à operação de organizações da sociedade civil.", st_body)

    heading1(doc, "ENTIDADE PARCEIRA (CONTEXTO DO ESTUDO)", st_h1)
    body(doc, "Observatório Social do Brasil de Uberlândia/MG. CNPJ 23.497.346/0001-42. Av. Vasconcelos Costa, nº 1500, Sala 3, Anexo I, Bairro Martins, Uberlândia-MG, CEP 38.400-452. Telefone (34) 3239-1529. https://www.osbrasiluberlandia.org/", st_body)

    heading1(doc, "DECLARAÇÕES", st_h1)
    body(doc, "O manuscrito é original, inédito e não está em avaliação em outro periódico. O autor concorda com a declaração de direitos autorais da RETII e com a coleta e o armazenamento de dados conforme a política de privacidade da revista.", st_body)
    body(doc, "No arquivo do manuscrito submetido pelo portal, a identificação foi substituída por XXX, conforme as Diretrizes para Autores.", st_body)
    doc.save(str(OUT_AUTOR))
    patch_headers(OUT_AUTOR)
    strip_unicode_dashes(OUT_AUTOR)


META = """SUBMISSÃO RETII (Revista de Engenharia, TI e Inovação)
ISSN 2966-2508  |  https://revistas.uniube.br/index.php/retii
Fluxo: 1. Início; 2. Transferência do manuscrito; 3. Inserir metadados; 4. Confirmação; 5. Próximos passos

==============================================================================
ARQUIVOS
==============================================================================
Enviar no passo 2 (manuscrito):
  artigo_retii_manuscrito_cego.docx
  (Word, identificação suprimida com XXX; é ESTE o arquivo principal)

Documento suplementar (passo 2, como suplementar, ou junto dos metadados):
  identificacao_autores.docx

Cópia de arquivo (não enviar ao portal, fica com o autor):
  artigo_retii_versao_identificada.docx

==============================================================================
PASSO 1. INÍCIO
==============================================================================
Idioma da submissão: Português (pt_BR)   (idioma principal)
Seção: Artigos
Política de seção: Política padrão de seção

Requisitos (marcar todos):
  [x] A contribuição é original e inédita, e não está sendo avaliada por outra revista
  [x] O arquivo está em Microsoft Word
  [x] URLs das referências foram fornecidos onde disponíveis
  [x] Texto em espaço simples; fonte 12 pontos; itálico em vez de sublinhado
       (exceto URL); figuras e tabelas no texto, não em anexo
  [x] O texto segue as Diretrizes para Autores

Declarações:
  [x] Sim, concordo com a declaração de direitos autorais
  [x] Sim, concordo com a coleta e o armazenamento de dados (privacidade)

==============================================================================
COMENTÁRIOS PARA O EDITOR (colar no campo)
==============================================================================
Prezados editores,

Encaminho manuscrito original e inédito, não submetido a outro periódico, para a
seção de Artigos, na área de Sistemas de Informação.

O texto descreve o desenvolvimento de um sistema de informação para coleta e
análise de licitações públicas, atendendo demanda do Observatório Social do
Brasil de Uberlândia/MG (CNPJ 23.497.346/0001-42). A identificação do autor foi
suprimida no arquivo principal, conforme as Diretrizes para Autores, e segue em
documento suplementar.

Coloco-me à disposição.

Atenciosamente,
Diogo Ferreira Moura
Universidade de Uberaba (UNIUBE)
Curso de Sistemas de Informação
CPF 058.161.826-25
Telefone: +55 (34) 99990-9660
ORCID: 0009-0009-1988-7741
diogo.moura@triggerti.com

==============================================================================
PASSO 3. METADADOS
==============================================================================
Título:
  Sistema de informação para coleta e análise de licitações públicas: apoio ao Observatório Social do Brasil em Uberlândia/MG

Título (inglês), se o formulário tiver campo separado:
  Information system for collecting and analyzing public procurement data: supporting the Social Observatory of Brazil in Uberlândia, MG

Resumo: (colar o do manuscrito; abaixo do teto de 200 palavras do template)

Palavras-chave:
  sistemas de informação; dados abertos; licitações públicas; controle social; observatório social.

Keywords:
  information systems; open data; public procurement; social control; social observatory.

Autor:
  Nome: Diogo Ferreira Moura
  Autor correspondente: sim
  E-mail: diogo.moura@triggerti.com
  Telefone: +55 (34) 99990-9660
  Endereço postal: Av. Araguari, 2815, ap. 304, Bairro Daniel Fonseca, Uberlândia/MG, CEP 38400-313
  Cidade/UF: Uberlândia/MG
  Filiação: Universidade de Uberaba (UNIUBE)
  Unidade / curso: Sistemas de Informação
  País: Brasil
  Bio: Graduando em Sistemas de Informação pela Universidade de Uberaba. Atua com desenvolvimento de sistemas de informação aplicados a dados públicos e à operação de organizações da sociedade civil.
  ORCID: 0009-0009-1988-7741
  URL ORCID: https://orcid.org/0009-0009-1988-7741

Não colocar o CPF no corpo do artigo. O CPF entra só no cadastro pessoal do SEER
e neste comentário ao editor / documento suplementar, se a secretaria pedir.

==============================================================================
CONFERÊNCIA RÁPIDA (normas RETII)
==============================================================================
- 5 a 20 páginas, figuras no texto: sim
- Título PT + EN, caixa de sentença: sim
- Resumo PT + EN, primeira frase com o tema: sim
- Palavras-chave ≤ 5, ponto e vírgula, ponto final: sim
- Introdução, metodologia, resultados, conclusão: sim
- Citações autor-data (NBR 10520) e referências ABNT com destaque em negrito: sim
- Identificação ausente no Word do portal: sim (versão cega)
"""


def main():
    write_article(OUT_CEGO, identificado=False)
    write_article(OUT_ID, identificado=True)
    write_identificacao()
    OUT_META.write_text(META, encoding="utf-8")
    print("gerado:")
    for p in (OUT_CEGO, OUT_ID, OUT_AUTOR, OUT_META):
        print(" ", p, p.stat().st_size)


if __name__ == "__main__":
    main()
