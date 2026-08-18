#!/usr/bin/env python3
"""Relatorio Semana 1: sistematizacao de evidencias (ACQA / TCC)."""
from __future__ import annotations

import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path("/home/dfmoura/Documents/test_several1/uniube/4")
EVID = ROOT / "evidencias"
FIG_ARTIGO = Path("/home/dfmoura/Documents/test_several1/uniube/3/figuras")
LOGO_OSB = Path(
    "/home/dfmoura/Documents/test_several1/trigger/23/app/static/logo_oficial_udi.png"
)
LOGIN_PNG = EVID / "01_tela_login.png"
MAPA_PNG = EVID / "02_tela_mapa_homologados.png"
CNPJS_PNG = EVID / "07_tela_cnpjs_vencedores.png"
OUT_PDF = ROOT / "Relatorio_Semana1_Sistematizacao_Evidencias.pdf"
OUT_DOCX = ROOT / "Relatorio_Semana1_Sistematizacao_Evidencias.docx"

NAVY = colors.HexColor("#1F4E79")
GRAY = colors.HexColor("#4A4A4A")
FORBIDDEN = "\u2014\u2013\u2012\u2212\u2015\u2026\u00b7\u201c\u201d"

# Texto unico para PDF e DOCX. Sem traco tipografico (em dash).
T = {
    "cover_comp": "Trabalho de Conclusão de Curso:<br/>Relato de Experiência do Projeto Integrado",
    "cover_comp_docx": "Trabalho de Conclusão de Curso: Relato de Experiência do Projeto Integrado",
    "cover_title": (
        "Relatório sobre o desenvolvimento e a implantação de um sistema de informação de "
        "apoio ao Observatório Social do Brasil em Uberlândia/MG: coleta, consulta e "
        "análise de licitações públicas"
    ),
    "cover_sub": "Sistematização de evidências da atividade extensionista",
    "cover_city": "Uberlândia / MG",
    "h1": "1 INTRODUÇÃO: ORGANIZAÇÃO PARCEIRA, PROBLEMA E CONTEXTO",
    "toc1": "1 Introdução: organização parceira, problema e contexto",
    "toc2": "2 Interação com a organização parceira",
    "toc3": "3 Proposta de solução (intervenção) adotada: sistema de apoio",
    "toc4": "4 Registros de evidências da ação realizada",
    "toc5": "5 Referências",
    "intro": (
        "Este relatório sistematiza as evidências da atividade de extensão que desenvolvi no "
        "Projeto Integrado e que agora integra o Trabalho de Conclusão de Curso. A organização "
        "parceira foi o Observatório Social do Brasil de Uberlândia, o OSB. A intervenção que "
        "realizei foi criar, implantar e colocar em uso um sistema de informação de apoio ao "
        "trabalho da entidade: um software web que coleta dados oficiais de licitações dos órgãos "
        "de Uberlândia, organiza a consulta em uma base local e reduz a digitação repetida na "
        "planilha Cronograma. O sistema está no ar no domínio da própria entidade, em "
        "https://licitacoes.osbrasiluberlandia.org/. Nas seções seguintes registro com quem "
        "trabalhei, o problema encontrado na rotina da entidade, como foi o contato, o que o "
        "sistema faz e as provas disso: telas, fluxos, volume carregado e depoimentos. Além do "
        "objetivo da atividade, veio um ganho indireto: a equipe passou a saber o que é uma API "
        "e que o governo já publica, por API, dado aberto para transparência das operações "
        "públicas, em especial a API de Dados Abertos do Compras.gov.br."
    ),
    "org1": (
        "A entidade parceira foi o Observatório Social do Brasil de Uberlândia, o OSB. "
        "Organização da sociedade civil, apartidária, sem fins lucrativos. CNPJ 23.497.346/0001-42. "
        "Endereço que usei no acompanhamento: Av. Vasconcelos Costa, 1500, Sala 3, Anexo I, "
        "Bairro Martins, Uberlândia-MG, CEP 38400-452. Telefone (34) 3239-1529. Site: "
        "https://www.osbrasiluberlandia.org/. A unidade local existe desde 2015, ligada a "
        "associação G7 (a CDL entra nisso) e faz parte da rede nacional do Observatório Social "
        "do Brasil."
    ),
    "org2": (
        "O que eles se propõem é despertar cidadania fiscal, no sentido de a sociedade acompanhar "
        "o uso do dinheiro público. Em Uberlândia isso aparece em duas frentes. Uma é o "
        "acompanhamento das licitações da Prefeitura e das autarquias e fundações (DMAE, FUTEL, "
        "EMAM, IPREMU, FERUB, ARESAN, PRODAUB e a Câmara). A outra é olhar a atuação dos vereadores. "
        "Eu atuei na primeira frente. Foi a que mais consumia tempo de planilha e foi o pedido mais "
        "claro da entidade: um sistema de apoio para reunir, consultar e analisar essas licitações, "
        "em vez de copiar processo por processo do portal para o Cronograma. O pedido era a "
        "ferramenta. O que a API era, e o fato de já existir uma API oficial de transparência para "
        "essas compras, não estava no vocabulário da equipe."
    ),
    "prob1": (
        "Dado público, no papel, não falta. Tem a Lei de Acesso à Informação, o decreto de dados "
        "abertos e a Lei 14.133, com o PNCP (BRASIL, 2011, 2016, 2021). O problema que eu vi no "
        "escritório era outro: o dado existia, mas o uso ainda era manual. Toda semana alguém entra "
        "no portal, copia processo, cola na planilha Cronograma e só depois analisa. Zuiderwijk e "
        "Janssen (2014) falam disso: publicar não é o mesmo que usar. Foi esse desencontro que "
        "justificou a intervenção: criar um sistema de apoio à coleta e à consulta, no tamanho da "
        "entidade."
    ),
    "prob2": (
        "A rotina que me mostraram é mais ou menos assim. Abre o site da prefeitura, transcreve "
        "para o Cronograma, monta o acompanhamento. Ao mesmo tempo consulta o Comprasnet pelas "
        "UASGs de Uberlândia e o Power BI da PMU. Esse painel atrasava, eles mesmos comentaram. "
        "Cada fonte descreve o processo de um jeito, com campo e código diferentes. O relatório "
        "de quadrimestre depende disso. Se a semana atrasava, o quadrimestre atrasava junto. "
        "O Comprasnet, para eles, era o site em que a pessoa pesquisa. Não havia, na rotina, a "
        "ideia de API: um canal em que o dado oficial já sai pronto para o sistema consumir, "
        "sem copiar a tela. Essa API de transparência das compras públicas já existia; o que "
        "faltava era o conhecimento e o uso."
    ),
    "cen1": (
        "É voluntariado. Reunião quase toda semana, em geral na associação comercial. Prestação "
        "de contas a cada quatro meses. Não tem equipe de TI nem verba fixa para servidor. Quatro "
        "usuários no sistema já dão conta. Não fazia sentido inventar arquitetura grande. O sistema "
        "de apoio precisava caber nesse cenário: consulta que funcionasse, cruzamento das bases e "
        "alguma máquina que pudesse buscar os dados de madrugada (LAUDON; LAUDON, 2014). Sem "
        "equipe de TI, conceito de API e de dado aberto consumido por máquina não fazia parte "
        "do dia a dia. Isso só entrou quando o sistema de apoio passou a usar a API de "
        "transparência do Compras.gov."
    ),
    "cen2": (
        "Os nomes também não batem. No painel da prefeitura o processo vem com ANOPROCESSO, "
        "MODALIDADE, VALORLICITACAO, SOLICITANTE. Na API federal é numeroControlePNCP, "
        "codigoModalidade. No portal da PMU aparece PE, PD, PI, CP. Se eu misturasse tudo num "
        "campo só, perdia o rastro da origem. Batini e Scannapieco (2016) tratam isso como "
        "qualidade de dado. Aqui era só não apagar o que a fonte trouxe. Esses nomes da API a "
        "equipe não via, porque trabalhava só na tela do portal. Só passaram a fazer sentido "
        "depois que expliquei o que era API e mostrei a de Dados Abertos do Compras.gov."
    ),
    "int1": (
        "Chegar no Observatório deu volta. Em março eu tinha pensado num treinamento de "
        "inteligência artificial para associados da CDL. Pensei no Antonio Carlos Oliveira "
        "porque ele tinha sido meu professor no curso de Administração, onde me formei, na "
        "ESAMC Uberlândia. Conversei com ele no sentido de pedir orientação: como eu poderia "
        "dar andamento no trabalho, a quem falar, por onde começar. Ele sugeriu procurar a "
        "Lécia Queiroz, da CDL. O resto foi por minha conta. No dia 20 de maio tive reunião "
        "presencial com a Lécia na CDL, Av. Belo Horizonte 1290. Ficou claro que colocar o "
        "projeto dentro da CDL naquele momento emperrava. Ela me passou o Marco Aurélio "
        "Freitas Santos, do OSB. Marcamos no mesmo dia. Fui em 21 de maio, 14h30."
    ),
    "int2": (
        "O que emperrou foi essa articulação inicial, não o convívio depois. O Antonio Carlos "
        "só apontou o caminho. Marcar a Lécia, ir na CDL, falar com o Marco e tocar o sistema "
        "foi comigo. Com o OSB o contato foi direto. Quase toda semana."
    ),
    "fac1": (
        "O Marco recebeu o projeto como sistema de apoio ao trabalho de acompanhamento, não como "
        "visita de faculdade. A gente se via na rotina deles, na ACIUB, com datashow. Mostrei o "
        "que o sistema já fazia e o que ainda estava no meio. Eles me explicaram a planilha "
        "Cronograma, o acompanhamento, e o que mais tomava tempo: achar o processo, cruzar com o "
        "Comprasnet, olhar o fornecedor e fechar o quadrimestre. Foi essa rotina que o software "
        "passou a apoiar. Nas mesmas reuniões eu expliquei, sem jargão de curso, o que é uma "
        "API: o jeito de o sistema buscar o dado direto da fonte, sem a pessoa copiar a tela. "
        "Mostrei que o governo já oferece isso para transparência das operações públicas, "
        "principalmente a API de Dados Abertos do Compras.gov.br. Esse conhecimento não era o "
        "objetivo da atividade; veio junto, de forma indireta."
    ),
    "fac2": (
        "No dia 11 de junho, depois de uma dessas reuniões, o Marco mandou a pergunta que virou "
        "tela no sistema:"
    ),
    "q1": (
        '"Seria possível extrairmos estatísticas de participação de empresas, % de vitória, '
        "quais tipos de objetos ou licitação participa e % de empresas que participam quando "
        "esta participa? o que gostaria de avaliar, existem histórias que empresas participam "
        "para favorecer uma e vao fazendo rodízio nas vitórias, então gostaria de analisar "
        'pela estatística."'
    ),
    "q1src": "Marco Aurélio Freitas Santos, WhatsApp, 11 jun. 2026.",
    "fac3": (
        "Na mesma conversa ele me chamou para ser voluntário. Combinamos de deixar isso mais "
        "formal depois que a faculdade acabar, para não virar obrigação de curso interno da rede. "
        "Também pediram um link no site deles. Ficou https://licitacoes.osbrasiluberlandia.org/. "
        "Em 15 de julho a Lécia me escreveu dizendo que o Marco estava contente com o que estava "
        "saindo:"
    ),
    "q2": '"Que legal! O Marco tem me informado, ele está muito contente com os resultados!"',
    "q2src": "Lécia Queiroz, CDL Uberlândia, WhatsApp, 15 jul. 2026.",
    "fac4": (
        "Isso não ficou só no WhatsApp. No dia 23 de julho, quinta-feira, das 15h30 às 16h, "
        "organizei uma apresentação na sala de reunião da ACIUB. O convite foi para a diretoria "
        "do Observatório, gente da CDL e da superintendência da ACIUB. O tema era o sistema de "
        "apoio que eu havia criado: a busca das licitações e a automação da coleta. Gostaram "
        "bastante da ferramenta. Além da tela, mostrei que a coleta automática não é um truque "
        "no site: usa uma API já existente de transparência das compras públicas, a de Dados "
        "Abertos do Compras.gov.br, que a maior parte da plateia não conhecia pelo nome."
    ),
    "fac5": (
        "Em 3 de agosto, segunda-feira, das 16h às 17h30, apresentei de novo, desta vez por "
        "Microsoft Teams, com o pessoal da sede do Observatório. Também gostaram. Chegaram a "
        "pedir um filtro por porte de empresa, para olhar se quem ganha é ME, EPP ou empresa "
        "maior. Eu coloquei isso depois, na tela de CNPJs vencedores. Na mesma reunião expliquei "
        "de novo o que é API e que essa API federal de transparência já estava disponível, "
        "mesmo antes do nosso sistema. Para quem observa licitação no portal, isso costuma "
        "passar despercebido."
    ),
    "dif1": (
        "Tive dificuldade no começo com a CDL. O Antonio Carlos orientou, mas não intermediou a "
        "reunião. Sem o encaminhamento da Lécia, depois que eu mesmo a procurei, acho que o prazo "
        "da disciplina teria travado. No OSB eu não quis chegar com sistema pronto. Quis entender "
        "a planilha primeiro. Isso atrasou eu mostrar tudo na nuvem, mas evitou empurrar um jeito "
        "de trabalhar que não era o deles."
    ),
    "dif1b": (
        "Hospedagem de graça enrolou de verdade. Testei plano da Oracle de madrugada. Tive uma "
        "reunião com o pessoal da Nuvolli para ver se dava para ter nuvem sem custo. No fim subi "
        "numa VM da AWS, região de São Paulo, Ubuntu, 1 GB de RAM, 2 vCPUs e 40 GB de disco, no "
        "plano gratuito. Tem crédito e prazo. O free tier desta conta acaba em 13 de janeiro de "
        "2027, ou antes se o crédito acabar. Quando isso acontecer, se ninguém pagar o plano ou "
        "mudar de máquina, o sistema sai do ar. Falei isso para eles. Não é solução eterna. É o "
        "que cabia no bolso da entidade agora."
    ),
    "dif2": (
        "Troquei de caminho técnico no meio. Primeiro pensei em raspar o portal da prefeitura com "
        "navegador. Quebrou fácil: layout, sessão, máquina com tela virtual. O sistema nasceu "
        "assim, com os CSVs oficiais do painel municipal. A API de Dados Abertos do Compras.gov.br "
        "entrou depois, quando a ferramenta já existia, para puxar o recorte federal das UASGs "
        "que eles já acompanhavam no portal Comprasnet. Pressman e Maxim (2021) e Sommerville "
        "(2018) tratam isso como desenvolvimento que vai se ajustando pelo uso. Foi o caso. "
        "Quando a API entrou, precisei explicar o que ela é e por que o portal que eles já "
        "usavam e a API não são a mesma coisa. Para a equipe, isso era novidade: existia um "
        "canal oficial de transparência feito para máquina, e não só a tela de pesquisa."
    ),
    "dif3": (
        "Depois que o sistema já estava no ar, apareceu uma dificuldade que não era da rotina "
        "da entidade, e sim do tipo de solução que eu escolhi. O módulo Compras.gov depende de "
        "uma API federal de dados abertos. Se essa API deixa de devolver registro, aquela ponta "
        "da coleta para de atualizar, mesmo com a compra visível no portal. Isso é um risco "
        "passível de qualquer sistema que consome fonte externa. Em agosto ocorreu um caso "
        "desses; o detalhe fica na seção 4.6, só como exemplo, não como o problema que o "
        "Observatório me pediu para resolver. Avisei a entidade do limite. Só deu para avisar "
        "com clareza porque, no meio do projeto, a ideia de API já tinha sido apresentada: não "
        "é o site que a pessoa vê; é o canal que o sistema consome."
    ),
    "sol0": (
        "A proposta de solução que adotei foi desenvolver, implantar e disponibilizar um sistema "
        "de informação de apoio ao Observatório Social do Brasil de Uberlândia. Não foi oficina, "
        "não foi cartilha e não foi só um diagnóstico: foi um software em produção, feito para a "
        "rotina de acompanhamento das licitações dos órgãos de Uberlândia. O sistema junta as "
        "fontes oficiais, grava numa base local e oferece telas de consulta (painel, processo, "
        "CNPJs vencedores, mapa de localidade e cobertura entre bases). Quem observa deixa de "
        "começar o trabalho copiando portal para planilha e passa a começar já com o dado "
        "organizado. A análise crítica continua sendo da entidade; o que o sistema tira é a "
        "digitação repetida. Endereço público, no domínio da entidade: "
        "https://licitacoes.osbrasiluberlandia.org/. Além do software, a intervenção levou à "
        "equipe um conhecimento que não estava no pedido inicial: o que é uma API e que já "
        "existe API pública para dar transparência às operações de compras, em especial a de "
        "Dados Abertos do Compras.gov.br. Na extensão, esse tipo de aprendizado entra de "
        "forma indireta, junto com o objetivo principal."
    ),
    "sol1": (
        "O sistema de apoio não nasceu na API federal. Primeiro ficou de pé a consulta local e a "
        "carga dos CSVs do Power BI da prefeitura: licitações, contratos, gestores e fiscais. O "
        "site weblicitacoes.uberlandia.mg.gov.br ficou para conferir na mão. Não usei ele como "
        "origem da carga. A API de Dados Abertos do Compras.gov/PNCP entrou depois, filtrada pelo "
        "município (IBGE 3170206) e pelas UASGs que eles já acompanhavam no portal Comprasnet, "
        "para ampliar o recorte federal. Essa API não substitui o portal de pesquisa nem a "
        "consulta do PNCP: é uma réplica em dados abertos. Para a entidade isso era novo. "
        "Eles conheciam o portal Comprasnet. Não conheciam a API de transparência que replica "
        "o mesmo universo para o sistema consumir. Como toda fonte externa, pode falhar. "
        "Se falhar, só aquele módulo deixa de atualizar. O restante do sistema continua."
    ),
    "sol2": (
        "A coleta do sistema de apoio roda sozinha, de preferência de madrugada. Quem observa "
        "entra já na base local. A análise continua sendo deles. O que sai é a digitação repetida. "
        "A planilha Cronograma ainda pode ficar do lado enquanto eles pegam confiança. Trocar "
        "ferramenta de uma hora para outra, nesse tipo de entidade, costuma virar ferramenta "
        "abandonada. O mesmo vale para o vocabulário novo: API, dado aberto, coleta automática. "
        "Entra no uso, não numa aula isolada."
    ),
    "tec1": (
        "Escolhi ferramenta barata de manter. FastAPI, SQLite no arquivo data/licitacoes.db, tela "
        "estática no mesmo serviço, Docker Compose numa VM da AWS (São Paulo, 1 GB de RAM, "
        "2 vCPUs, 40 GB). No máximo quatro "
        "contas: um admin e três de consulta. Coleta, Setup, disparo de CNPJ e token de IA só o "
        "admin mexe. Não coloquei RabbitMQ, MinIO nem banco gerenciado. Não precisava. Anotação "
        "que o observador faz na tela (tipo observador_id) não some quando a coleta roda de novo. "
        "Backup diário do SQLite, com pouca retenção, porque disco grátis acaba rápido. Quando o "
        "plano gratuito da AWS terminar, em janeiro de 2027, essa máquina deixa de ser de graça."
    ),
    "lim1": (
        "O sistema não aponta irregularidade sozinho. Organiza consulta. A API federal só devolve "
        "resultado classificado ou homologado, não a lista inteira de quem participou. Então o "
        "estudo de rodízio fica pela metade, e eu falei isso para eles. Há um segundo limite, "
        "este de acoplamento: a parte Compras.gov depende da API de Dados Abertos. Se a API "
        "trava, essa parte para de receber atualização, mesmo quando o registro já está no portal. "
        "Isso é passível de acontecer em qualquer consumo de dado aberto por API. Não é o "
        "problema original da entidade; é um risco da arquitetura. Um caso ocorrido em agosto "
        "está na seção 4.6, só como demonstração. Depois que a ideia de API ficou clara para "
        "a equipe, esse limite passou a fazer sentido: o portal humano pode mostrar a compra "
        "e o canal da máquina, não. A rotina de preço com modelo de linguagem é "
        "rascunho. O servidor refaz a conta com faixa de 15%. Não substitui pesquisa formal nem "
        "nota fiscal."
    ),
    "lim2": (
        "Em resumo: o sistema de apoio está no ar, no domínio da entidade, em "
        "https://licitacoes.osbrasiluberlandia.org/. Na minha máquina de desenvolvimento: "
        "docker compose up --build -d, endereço http://localhost:8096/."
    ),
    "ev0": (
        "O que segue é prova de que o sistema de apoio foi criado, implantado e usado, e não só "
        "descrito em relatório. Não é protótipo de disciplina. Tem print da tela do software "
        "(login, mapa de localidade e CNPJs vencedores), desenho do fluxo antes e depois, "
        "arquitetura, volume carregado, recado da entidade e o artigo que eu preparei para a "
        "RETII. No fim, a seção 4.6 registra um exemplo de risco da arquitetura: se a API "
        "externa usada pelo módulo Compras.gov para, aquele módulo deixa de atualizar. Esse "
        "exemplo também firmou, na prática, o que a equipe passou a entender no projeto: API "
        "não é o site; é o canal oficial de transparência que o sistema consome."
    ),
    "ev1": (
        "A Figura 1 é a tela de login, com a logo do Observatório. Sem usuário e senha a base não "
        "fica aberta na internet. Tirei o print na instância local, porta 8096, a mesma aplicação "
        "que está no endereço público."
    ),
    "ev1b": (
        "A Figura 2 é a tela Mapa de localidade. Mostra a sede dos vencedores das licitações, "
        "a partir dos itens homologados. Em cima ficam filtros de período, ano, órgão, modalidade, "
        "UF do vencedor, porte da empresa e a métrica do mapa (itens homologados, contratações ou "
        "valor). Os cards resumem o recorte: itens, contratações, valor homologado, UFs, municípios "
        "e a fatia que ficou em Uberlândia versus o que veio de fora. No centro, o mapa do Brasil, "
        "por estado ou calor de município. Do lado, o ranking por UF. No print: 6.797 itens "
        "homologados, 1.270 contratações, cerca de R$ 1,2 bilhão, 25 UFs e 390 municípios. Cerca "
        "de 25,5% dos itens em Uberlândia e 74,5% de fora. A lista nominal desses vencedores, "
        "por CNPJ, fica na Figura 3. Mesma instância local da Figura 1."
    ),
    "ev1c": (
        "A Figura 3 é a tela CNPJs vencedores. Não é o mapa da sede: é a lista de quem ganhou. "
        "Cada linha é um fornecedor (CNPJ ou CPF), com quantos itens levou, em quantas compras, "
        "o valor homologado e o município da sede. Em cima ficam filtros de período, ano, órgão, "
        "modalidade, UF do vencedor, cache do cadastro e porte da empresa, além da busca por nome "
        "ou CNPJ. O filtro de porte foi o que a sede do Observatório pediu em 3 de agosto. No "
        "print: 2.037 fornecedores consolidados, 16 com cadastro atualizado, 1.833 com cache "
        "vencido e 188 CPF. A origem é o resultado homologado do Compras.gov, enriquecido com "
        "dados públicos do CNPJ. Essa tela responde à pergunta do Marco sobre estatística de "
        "vitória: quem ganha, com que frequência e em que valor. Mesma instância local das "
        "Figuras 1 e 2."
    ),
    "ev2": (
        "A Figura 4 mostra o antes e o depois da intervenção. Antes, a análise só começava depois "
        "de copiar processo do portal para a planilha. Depois da implantação do sistema de apoio, "
        "a coleta alimenta a base local e o tempo da reunião vai para ler o processo, e não para "
        "digitar de novo o que o portal já tinha. O antes também era conceitual: o dado só "
        "existia na tela do portal. O depois inclui saber que o governo já publica o mesmo "
        "conteúdo por API de transparência, em especial a de Dados Abertos do Compras.gov.br."
    ),
    "ev3": (
        "A Figura 5 mostra de onde vem o dado (API federal, CSV do painel, cadastro de UASG), o "
        "hub de coleta de madrugada, o SQLite e as telas. A ponta Compras.gov foi encaixada depois "
        "que o sistema já existia. Nenhuma peça é novidade sozinha. O que importa é caber no "
        "tamanho da entidade e não precisar de gente de infraestrutura o tempo todo. A figura "
        "também serviu para mostrar à equipe de onde o dado vem, inclusive a API federal, que "
        "antes não fazia parte do vocabulário deles. Se essa API deixar de responder, só "
        "aquela coleta para; o restante do sistema segue com o CSV municipal."
    ),
    "ev4": (
        "A Tabela 1 sai dos CSVs oficiais do painel da prefeitura, corte em 18 de agosto de 2026. "
        "Não é o universo do PNCP. É o recorte com o qual eles já trabalhavam. Gestores e fiscais "
        "passam de 36 mil linhas. Ninguém ia conferir isso na planilha Cronograma."
    ),
    "ev5": (
        "A Figura 6 é a modalidade em 2025. Pregão eletrônico 559, dispensa 448, no total de 1.215 "
        "processos. Quem só olha concorrência grande deixa passar o miolo, que está no eletrônico "
        "e na dispensa (BRASIL, 2021). Quem mais solicita: DMAE com 333, Saúde com 233."
    ),
    "ev6": (
        "Fora os recados da seção 2, teve o convite para ser voluntário, o pedido do link no site "
        "e duas apresentações formais: uma na ACIUB, em 23 de julho, para a diretoria, e outra "
        "no Teams, em 3 de agosto, com a sede do Observatório. Nas duas a ferramenta foi bem "
        "recebida. Na segunda pediram filtro por porte, e isso entrou na tela de CNPJs vencedores. "
        "Escrevi também "
        "um artigo para a Revista de Engenharia, TI e Inovação (RETII, ISSN 2966-2508), da Uniube, "
        "com o mesmo tema deste relatório. Não medi com relógio quanto tempo eles economizaram. "
        "O que eles relatam é que o trabalho se rearranjou. A frente dos vereadores eu deixei de "
        "fora de propósito. Outra coisa que fica aberta: quando acabar o plano gratuito da AWS, "
        "em janeiro de 2027, o sistema como está hoje sai do ar, a menos que a entidade pague "
        "ou mude de hospedagem. Nas apresentações o ganho não foi só a tela. Diretoria e sede "
        "viram que existe API de transparência para operação pública, e que o sistema a usa. "
        "Esse tipo de conhecimento entra indireto na extensão, junto com o objetivo principal."
    ),
    "ev7": (
        "Esta seção não descreve o problema da entidade. O problema da entidade era a coleta "
        "manual na planilha Cronograma. O que segue é só um exemplo, ocorrido em agosto, de um "
        "risco passível da arquitetura: o módulo Compras.gov consome a API de Dados Abertos; se "
        "essa API deixa de devolver o registro, aquele módulo para de atualizar, mesmo com a "
        "compra visível no portal. Para documentar o caso, em 7 de agosto de 2026 abri o chamado "
        "56398424 no Portal de Serviços. A contratação 18431312000620-1-000341/2026, da UASG "
        "926922, Pregão Eletrônico, Lei 14.133, estava no PNCP desde 27 de julho. No portal "
        "Comprasnet (https://cnetmobile.estaleiro.serpro.gov.br/comprasnet-web/public/compras) "
        "a mesma compra aparece no acompanhamento 92692205002262026. Consulta na API do PNCP: "
        "HTTP 200. Consulta na API Dados Abertos: totalRegistros=0. Os sequenciais 000333 e "
        "000334 voltavam normalmente. Em 13 de agosto cobrei o andamento. Em 17 de agosto o "
        "SIASG informou que o caso ainda estava em análise. Serve para mostrar o limite, não "
        "para redefinir o diagnóstico do projeto. Serviu também, na prática, para a equipe "
        "entender o que é API: não é o portal que a pessoa vê; é o canal que o sistema "
        "consome. Se esse canal falha, a tela humana continua e a coleta automática não."
    ),
}

def prepare_assets() -> None:
    EVID.mkdir(parents=True, exist_ok=True)
    mapping = {
        "04_fluxo.png": FIG_ARTIGO / "fig1_fluxo.png",
        "05_arquitetura.png": FIG_ARTIGO / "fig2_arquitetura.png",
        "06_modalidades_2025.png": FIG_ARTIGO / "fig3_modalidades_2025.png",
        "00_logo_osb.png": LOGO_OSB,
    }
    for dest_name, src in mapping.items():
        dest = EVID / dest_name
        if src.exists():
            shutil.copy2(src, dest)


def styles():
    base = getSampleStyleSheet()
    return {
        "cover_inst": ParagraphStyle(
            "CoverInst", parent=base["Normal"], fontName="Times-Bold",
            fontSize=13, leading=17, alignment=TA_CENTER, textColor=NAVY, spaceAfter=2,
        ),
        "cover_sub": ParagraphStyle(
            "CoverSub", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=16, alignment=TA_CENTER, spaceAfter=2,
        ),
        "cover_comp": ParagraphStyle(
            "CoverComp", parent=base["Normal"], fontName="Times-Italic",
            fontSize=11, leading=15, alignment=TA_CENTER, textColor=GRAY, spaceAfter=4,
        ),
        "cover_title": ParagraphStyle(
            "CoverTitle", parent=base["Normal"], fontName="Times-Bold",
            fontSize=16, leading=22, alignment=TA_CENTER, textColor=NAVY, spaceAfter=8,
        ),
        "cover_author": ParagraphStyle(
            "CoverAuthor", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=16, alignment=TA_CENTER, spaceAfter=2,
        ),
        "h1": ParagraphStyle(
            "H1", parent=base["Heading1"], fontName="Times-Bold",
            fontSize=13, leading=17, spaceBefore=14, spaceAfter=8, textColor=NAVY,
        ),
        "h2": ParagraphStyle(
            "H2", parent=base["Heading2"], fontName="Times-Bold",
            fontSize=12, leading=16, spaceBefore=10, spaceAfter=6,
            textColor=colors.HexColor("#2c5282"),
        ),
        "body": ParagraphStyle(
            "Body", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=18, alignment=TA_JUSTIFY, firstLineIndent=1.25 * cm, spaceAfter=8,
        ),
        "quote": ParagraphStyle(
            "Quote", parent=base["Normal"], fontName="Times-Italic",
            fontSize=11, leading=16, alignment=TA_JUSTIFY,
            leftIndent=1.2 * cm, rightIndent=0.8 * cm, spaceBefore=4, spaceAfter=4,
        ),
        "quote_src": ParagraphStyle(
            "QuoteSrc", parent=base["Normal"], fontName="Times-Roman",
            fontSize=10, leading=13, alignment=TA_RIGHT, leftIndent=1.2 * cm,
            spaceAfter=10, textColor=GRAY,
        ),
        "caption": ParagraphStyle(
            "Caption", parent=base["Normal"], fontName="Times-Bold",
            fontSize=10, leading=13, alignment=TA_CENTER, spaceBefore=6, spaceAfter=2,
        ),
        "fonte": ParagraphStyle(
            "Fonte", parent=base["Normal"], fontName="Times-Roman",
            fontSize=9, leading=12, alignment=TA_CENTER, spaceAfter=10, textColor=GRAY,
        ),
        "ref": ParagraphStyle(
            "Ref", parent=base["Normal"], fontName="Times-Roman",
            fontSize=11, leading=15, alignment=TA_JUSTIFY,
            leftIndent=1.25 * cm, firstLineIndent=-1.25 * cm, spaceAfter=8,
        ),
        "toc": ParagraphStyle(
            "Toc", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=20, alignment=TA_LEFT, spaceAfter=2,
        ),
        "cell": ParagraphStyle(
            "Cell", parent=base["Normal"], fontName="Times-Roman",
            fontSize=9, leading=12, alignment=TA_LEFT,
        ),
        "cell_c": ParagraphStyle(
            "CellC", parent=base["Normal"], fontName="Times-Roman",
            fontSize=9, leading=12, alignment=TA_CENTER,
        ),
        "cell_h": ParagraphStyle(
            "CellH", parent=base["Normal"], fontName="Times-Bold",
            fontSize=9, leading=12, alignment=TA_CENTER, textColor=colors.white,
        ),
    }


def P(text, style):
    return Paragraph(text, style)


def fitted_image(path: Path, max_w_cm: float, max_h_cm: float) -> Image:
    img = Image(str(path))
    iw, ih = float(img.imageWidth), float(img.imageHeight)
    max_w, max_h = max_w_cm * cm, max_h_cm * cm
    ratio = min(max_w / iw, max_h / ih)
    img.drawWidth = iw * ratio
    img.drawHeight = ih * ratio
    img.hAlign = "CENTER"
    return img


def abnt_table(headers, rows, col_widths, s):
    head = [P(h, s["cell_h"]) for h in headers]
    body = []
    for row in rows:
        line = []
        for i, val in enumerate(row):
            st = s["cell"] if i == 0 else s["cell_c"]
            line.append(P(str(val), st))
        body.append(line)
    t = Table([head] + body, colWidths=col_widths, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e0")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f7fb")]),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return t


def later_pages(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(NAVY)
    canvas.setLineWidth(0.6)
    canvas.line(2 * cm, A4[1] - 1.5 * cm, A4[0] - 2 * cm, A4[1] - 1.5 * cm)
    canvas.setFont("Times-Italic", 8)
    canvas.setFillColor(GRAY)
    canvas.drawString(
        2 * cm, A4[1] - 1.35 * cm,
        "UNIUBE  Relato de experiência: sistema de apoio ao OSB",
    )
    canvas.line(2 * cm, 1.5 * cm, A4[0] - 2 * cm, 1.5 * cm)
    canvas.setFont("Times-Roman", 8)
    canvas.drawString(2 * cm, 1.15 * cm, "Diôgo Ferreira Moura  RA 1030125-2")
    canvas.drawRightString(A4[0] - 2 * cm, 1.15 * cm, f"{doc.page}")
    canvas.restoreState()


def first_page(canvas, doc):
    canvas.saveState()
    canvas.restoreState()


Q1_ORG = [
    ["Nome", "Observatório Social do Brasil de Uberlândia"],
    ["CNPJ", "23.497.346/0001-42"],
    ["Natureza", "OSC apartidária, sem fins lucrativos"],
    ["CNAE principal", "94.99-5-00  Atividades associativas"],
    ["Contato operacional", "Marco Aurélio Freitas Santos  (34) 9979-6169"],
    ["Ponte institucional", "Lécia Queiroz (CDL Uberlândia)"],
    ["Site", "https://www.osbrasiluberlandia.org/"],
]
Q2_UASG = [
    ["PMU", "Prefeitura Municipal de Uberlândia", "926922"],
    ["DMAE", "Departamento Municipal de Água e Esgoto", "926287"],
    ["FUTEL", "Fundação Uberlandense do Turismo, Esporte e Lazer", "926038"],
    ["ARESAN", "Agência de Regulação dos Serviços de Saneamento", "931351"],
    ["FERUB", "Fundação de Excelência Rural de Uberlândia", "930403"],
    ["EMAM", "Empresa Municipal de Apoio e Manutenção", "929315"],
    ["IPREMU", "Instituto de Previdência dos Servidores Municipais", "929301"],
    ["CAM", "Câmara Municipal de Uberlândia", "925010"],
]
Q3_INT = [
    ["16/03/2026", "Conversa com Antonio Carlos Oliveira, meu ex-professor de Administração na ESAMC. Pedi orientação; ele sugeriu procurar a Lécia na CDL."],
    ["20/05/2026", "Reunião na CDL com Lécia Queiroz, marcada por mim. Encaminhamento ao OSB e contato com Marco Aurélio."],
    ["21/05/2026", "Primeira reunião no Observatório, 14h30. Apresentação da entidade."],
    ["10/06/2026", "Reunião com datashow. Mostrei o que já estava andando e comecei a explicar o que é API e que já existe API federal de transparência das compras."],
    ["11/06/2026", "Pedido de estatística de vencedores. Convite a voluntariado. Pedido de link no site."],
    ["17/06/2026", "Convite para a prestação de contas do 1o quadrimestre de 2026 (Teams)."],
    ["19 a 25/06/2026", "Conversa sobre hospedagem gratuita. O sistema já existia (CSV municipal). A API Compras.gov entrou depois; para a equipe, o conceito de API ainda era novo."],
    ["15/07/2026", "Retorno da Lécia: o Marco estava contente com os resultados."],
    ["23/07/2026", "Apresentação na ACIUB. Além da ferramenta, mostrei a API de Dados Abertos do Compras.gov como fonte oficial de transparência, pouco conhecida na plateia."],
    ["03/08/2026", "Reunião no Teams com a sede do OSB. Pediram filtro por porte. Expliquei de novo o que é API e que essa API de transparência já existia."],
    ["07 a 17/08/2026", "Exemplo de risco da arquitetura (não do problema da entidade): a API de Dados Abertos parou de devolver registro e o módulo Compras.gov deixou de atualizar. Chamado 56398424; detalhe na seção 4.6."],
]
Q5_CHAMADO = [
    ["Chamado", "56398424 (chave 94441)"],
    ["Portal", "https://portaldeservicos.gestao.gov.br/"],
    ["Abertura", "07/08/2026, 13h14"],
    ["Contratação PNCP", "18431312000620-1-000341/2026"],
    ["Órgão / UASG", "Município de Uberlândia / 926922"],
    ["Modalidade", "Pregão Eletrônico (Lei 14.133, art. 28, I)"],
    ["Divulgação no PNCP", "27/07/2026"],
    ["Portal Comprasnet", "Aparece normalmente (pesquisa que a entidade já usa)"],
    ["Acompanhamento", "https://cnetmobile.estaleiro.serpro.gov.br/comprasnet-web/public/compras/acompanhamento-compra?compra=92692205002262026"],
    ["API PNCP", "HTTP 200, registro presente"],
    ["API Dados Abertos", "HTTP 200, totalRegistros=0"],
    ["Lacuna observada", "Sequenciais 000335 em diante ausentes; último ok: 000334 (23/07/2026)"],
    ["Efeito no sistema", "Só o módulo Compras.gov deixa de atualizar (risco da API, não da rotina da entidade)"],
    ["Situação em 17/08/2026", "Em análise (SIASG 3o Nível)"],
]
Q4_MOD = [
    ["Painel", "Quantidade e valor por situação, órgão e modalidade, nas duas fontes, com filtro de período."],
    ["Cobertura entre bases", "O que está no Compras.gov e não no painel municipal, e o inverso, pela chave órgão + ano + processo."],
    ["Consulta por processo", "Junta o que houver nas duas bases sobre o mesmo número de processo."],
    ["CNPJs vencedores", "Lista de quem ganhou (CNPJ ou CPF): itens, compras, valor homologado, porte, CNAE e município da sede. Filtro por porte pedido pela equipe em 03/08."],
    ["Mapa de localidade", "Sede dos vencedores das licitações (itens homologados). Filtros de período, órgão, modalidade, UF, porte e métrica. Cards com totais e mapa do Brasil, com ranking por UF."],
    ["Propostas abertas", "Itens com prazo PNCP vigente e rotina auxiliar de preço (rascunho, não pesquisa formal)."],
    ["Coleta e Setup", "Orquestração das fontes, agendamento noturno, UASGs, usuários e backup (só admin). A ponta Compras.gov usa a API de Dados Abertos, canal de transparência que a equipe não conhecia pelo nome. Se a API para, só essa coleta para."],
]
TAB1 = [
    ["Processos licitatórios", "1.215", "628", "Campos oficiais do painel"],
    ["Contratos (linhas)", "2.152", "556", "Inclui aditivos/parcelas na origem"],
    ["Gestores e fiscais", "n/a", "36.399", "Base acumulada, não anualizada"],
]
REFS = [
    "<b>BATINI, Carlo; SCANNAPIECO, Monica.</b> <i>Data and information quality: dimensions, principles and techniques.</i> Cham: Springer, 2016.",
    "<b>BRASIL.</b> <i>Decreto n. 8.777, de 11 de maio de 2016.</i> Institui a Política de Dados Abertos do Poder Executivo federal. Diário Oficial da União: seção 1, Brasília, DF, 12 maio 2016. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2016/decreto/d8777.htm. Acesso em: 18 ago. 2026.",
    "<b>BRASIL.</b> <i>Lei n. 12.527, de 18 de novembro de 2011.</i> Regula o acesso a informações. Diário Oficial da União: seção 1, Brasília, DF, 18 nov. 2011. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2011/lei/l12527.htm. Acesso em: 18 ago. 2026.",
    "<b>BRASIL.</b> <i>Lei n. 14.133, de 1 de abril de 2021.</i> Lei de Licitações e Contratos Administrativos. Diário Oficial da União: seção 1, Brasília, DF, 1 abr. 2021. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm. Acesso em: 18 ago. 2026.",
    "<b>BRASIL.</b> <i>Portal de Serviços.</i> Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://portaldeservicos.gestao.gov.br/. Acesso em: 18 ago. 2026.",
    "<b>COMPRAS.GOV.BR.</b> <i>Acompanhamento de compras (Comprasnet).</i> Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://cnetmobile.estaleiro.serpro.gov.br/comprasnet-web/public/compras. Acesso em: 18 ago. 2026.",
    "<b>COMPRAS.GOV.BR.</b> <i>Compras públicas em dados abertos.</i> Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://dadosabertos.compras.gov.br/swagger-ui/index.html. Acesso em: 18 ago. 2026.",
    "<b>LAUDON, Kenneth C.; LAUDON, Jane P.</b> <i>Sistemas de informação gerenciais.</i> 11. ed. São Paulo: Pearson, 2014.",
    "<b>OBSERVATÓRIO SOCIAL DO BRASIL DE UBERLÂNDIA.</b> <i>Página institucional.</i> Uberlândia, [2026]. Disponível em: https://www.osbrasiluberlandia.org/. Acesso em: 18 ago. 2026.",
    "<b>PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS.</b> <i>API de consulta: contratação 18431312000620-1-000341/2026.</i> Brasília, DF: Governo Federal, 2026. Disponível em: https://pncp.gov.br/api/consulta/v1/orgaos/18431312000620/compras/2026/341. Acesso em: 18 ago. 2026.",
    "<b>PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS.</b> <i>Manuais do PNCP.</i> Brasília, DF: Governo Federal, [2026]. Disponível em: https://www.gov.br/pncp/pt-br/pncp/manuais. Acesso em: 18 ago. 2026.",
    "<b>PREFEITURA MUNICIPAL DE UBERLÂNDIA.</b> <i>Painel de licitações e contratos (dados abertos).</i> Uberlândia, [2026]. Disponível em: https://app.powerbi.com/. Acesso em: 18 ago. 2026.",
    "<b>PRESSMAN, Roger S.; MAXIM, Bruce R.</b> <i>Engenharia de software: uma abordagem profissional.</i> 9. ed. Porto Alegre: AMGH, 2021.",
    "<b>SOMMERVILLE, Ian.</b> <i>Engenharia de software.</i> 10. ed. São Paulo: Pearson, 2018.",
    "<b>ZUIDERWIJK, Anneke; JANSSEN, Marijn.</b> Open data policies, their implementation and impact: a framework for comparison. <i>Government Information Quarterly</i>, [s. l.], v. 31, n. 1, p. 17-29, 2014. Disponível em: https://doi.org/10.1016/j.giq.2013.04.003. Acesso em: 18 ago. 2026.",
]
REFS_PLAIN = [
    "BATINI, Carlo; SCANNAPIECO, Monica. Data and information quality: dimensions, principles and techniques. Cham: Springer, 2016.",
    "BRASIL. Decreto n. 8.777, de 11 de maio de 2016. Institui a Política de Dados Abertos do Poder Executivo federal. Diário Oficial da União: seção 1, Brasília, DF, 12 maio 2016. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2016/decreto/d8777.htm. Acesso em: 18 ago. 2026.",
    "BRASIL. Lei n. 12.527, de 18 de novembro de 2011. Regula o acesso a informações. Diário Oficial da União: seção 1, Brasília, DF, 18 nov. 2011. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2011/lei/l12527.htm. Acesso em: 18 ago. 2026.",
    "BRASIL. Lei n. 14.133, de 1 de abril de 2021. Lei de Licitações e Contratos Administrativos. Diário Oficial da União: seção 1, Brasília, DF, 1 abr. 2021. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm. Acesso em: 18 ago. 2026.",
    "BRASIL. Portal de Serviços. Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://portaldeservicos.gestao.gov.br/. Acesso em: 18 ago. 2026.",
    "COMPRAS.GOV.BR. Acompanhamento de compras (Comprasnet). Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://cnetmobile.estaleiro.serpro.gov.br/comprasnet-web/public/compras. Acesso em: 18 ago. 2026.",
    "COMPRAS.GOV.BR. Compras públicas em dados abertos. Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://dadosabertos.compras.gov.br/swagger-ui/index.html. Acesso em: 18 ago. 2026.",
    "LAUDON, Kenneth C.; LAUDON, Jane P. Sistemas de informação gerenciais. 11. ed. São Paulo: Pearson, 2014.",
    "OBSERVATÓRIO SOCIAL DO BRASIL DE UBERLÂNDIA. Página institucional. Uberlândia, [2026]. Disponível em: https://www.osbrasiluberlandia.org/. Acesso em: 18 ago. 2026.",
    "PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS. API de consulta: contratação 18431312000620-1-000341/2026. Brasília, DF: Governo Federal, 2026. Disponível em: https://pncp.gov.br/api/consulta/v1/orgaos/18431312000620/compras/2026/341. Acesso em: 18 ago. 2026.",
    "PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS. Manuais do PNCP. Brasília, DF: Governo Federal, [2026]. Disponível em: https://www.gov.br/pncp/pt-br/pncp/manuais. Acesso em: 18 ago. 2026.",
    "PREFEITURA MUNICIPAL DE UBERLÂNDIA. Painel de licitações e contratos (dados abertos). Uberlândia, [2026]. Disponível em: https://app.powerbi.com/. Acesso em: 18 ago. 2026.",
    "PRESSMAN, Roger S.; MAXIM, Bruce R. Engenharia de software: uma abordagem profissional. 9. ed. Porto Alegre: AMGH, 2021.",
    "SOMMERVILLE, Ian. Engenharia de software. 10. ed. São Paulo: Pearson, 2018.",
    "ZUIDERWIJK, Anneke; JANSSEN, Marijn. Open data policies, their implementation and impact: a framework for comparison. Government Information Quarterly, [s. l.], v. 31, n. 1, p. 17-29, 2014. Disponível em: https://doi.org/10.1016/j.giq.2013.04.003. Acesso em: 18 ago. 2026.",
]


def build_pdf() -> None:
    s = styles()
    story = []
    story.append(Spacer(1, 1.6 * cm))
    story.append(P("UNIVERSIDADE DE UBERABA", s["cover_inst"]))
    story.append(P("UNIUBE", s["cover_sub"]))
    story.append(Spacer(1, 0.25 * cm))
    story.append(P("Curso de Sistemas de Informação", s["cover_sub"]))
    story.append(P("Polo Uberlândia", s["cover_sub"]))
    story.append(Spacer(1, 2.2 * cm))
    story.append(P(T["cover_comp"], s["cover_comp"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(P(T["cover_title"], s["cover_title"]))
    story.append(P(T["cover_sub"], s["cover_comp"]))
    story.append(Spacer(1, 2.4 * cm))
    story.append(P("Diôgo Ferreira Moura", s["cover_author"]))
    story.append(P("RA 1030125-2", s["cover_author"]))
    story.append(Spacer(1, 3.2 * cm))
    story.append(P(T["cover_city"], s["cover_author"]))
    story.append(P("Agosto de 2026", s["cover_author"]))
    story.append(PageBreak())

    story.append(P("SUMÁRIO", s["h1"]))
    for k in ("toc1", "toc2", "toc3", "toc4", "toc5"):
        story.append(P(T[k], s["toc"]))
    story.append(PageBreak())

    story.append(P(T["h1"], s["h1"]))
    story.append(P(T["intro"], s["body"]))
    story.append(P("1.1 Identificação da organização parceira", s["h2"]))
    if (EVID / "00_logo_osb.png").exists():
        story.append(fitted_image(EVID / "00_logo_osb.png", 7.2, 2.6))
        story.append(Spacer(1, 0.15 * cm))
    story.append(P(T["org1"], s["body"]))
    story.append(P(T["org2"], s["body"]))
    story.append(abnt_table(["Campo", "Informação"], Q1_ORG, [5.2 * cm, 11.6 * cm], s))
    story.append(P("Quadro 1. Identificação da organização parceira", s["caption"]))
    story.append(P("Fonte: registros do projeto e material da entidade (2026).", s["fonte"]))

    story.append(P("1.2 Problema identificado", s["h2"]))
    story.append(P(T["prob1"], s["body"]))
    story.append(P(T["prob2"], s["body"]))
    story.append(P("1.3 Cenário e contexto encontrados", s["h2"]))
    story.append(P(T["cen1"], s["body"]))
    story.append(P("As unidades compradoras que eles já acompanhavam no portal Comprasnet, e que orientei o recorte do sistema, são estas. O portal eles já usavam; a API de transparência dessas mesmas UASGs, não:", s["body"]))
    story.append(abnt_table(["Sigla", "Órgão", "UASG"], Q2_UASG, [2.6 * cm, 11.4 * cm, 2.8 * cm], s))
    story.append(P("Quadro 2. Unidades compradoras acompanhadas em Uberlândia", s["caption"]))
    story.append(P("Fonte: rotina da entidade e API Compras.gov (2026).", s["fonte"]))
    story.append(P(T["cen2"], s["body"]))

    story.append(P("2 INTERAÇÃO COM A ORGANIZAÇÃO PARCEIRA", s["h1"]))
    story.append(P("2.1 Como o contato se estabeleceu", s["h2"]))
    story.append(P(T["int1"], s["body"]))
    story.append(P(T["int2"], s["body"]))
    story.append(P("2.2 Facilidades encontradas", s["h2"]))
    story.append(P(T["fac1"], s["body"]))
    story.append(P(T["fac2"], s["body"]))
    story.append(P(T["q1"], s["quote"]))
    story.append(P(T["q1src"], s["quote_src"]))
    story.append(P(T["fac3"], s["body"]))
    story.append(P(T["q2"], s["quote"]))
    story.append(P(T["q2src"], s["quote_src"]))
    story.append(P(T["fac4"], s["body"]))
    story.append(P(T["fac5"], s["body"]))
    story.append(P("2.3 Dificuldades e o que precisei mudar", s["h2"]))
    story.append(P(T["dif1"], s["body"]))
    story.append(P(T["dif1b"], s["body"]))
    story.append(P(T["dif2"], s["body"]))
    story.append(P(T["dif3"], s["body"]))
    story.append(abnt_table(["Data", "O que aconteceu"], Q3_INT, [3.4 * cm, 13.4 * cm], s))
    story.append(P("Quadro 3. Síntese da interação com a organização parceira", s["caption"]))
    story.append(P("Fonte: WhatsApp, reuniões do projeto e Portal de Serviços (2026).", s["fonte"]))

    story.append(P("3 PROPOSTA DE SOLUÇÃO (INTERVENÇÃO) ADOTADA", s["h1"]))
    story.append(P("3.1 O sistema de apoio criado", s["h2"]))
    story.append(P(T["sol0"], s["body"]))
    story.append(P(T["sol1"], s["body"]))
    story.append(P(T["sol2"], s["body"]))
    story.append(P("3.2 Escolha técnica e tamanho da operação", s["h2"]))
    story.append(P(T["tec1"], s["body"]))
    story.append(P("3.3 Telas que ficaram na rotina", s["h2"]))
    story.append(abnt_table(["Módulo", "Para que serve"], Q4_MOD, [4.2 * cm, 12.6 * cm], s))
    story.append(P("Quadro 4. Módulos do sistema de apoio e uso previsto", s["caption"]))
    story.append(P("Fonte: sistema de apoio implantado (2026).", s["fonte"]))
    story.append(P(T["lim1"], s["body"]))
    story.append(P(T["lim2"], s["body"]))

    story.append(P("4 REGISTROS DE EVIDÊNCIAS DA AÇÃO REALIZADA", s["h1"]))
    story.append(P(T["ev0"], s["body"]))
    story.append(P("4.1 Print da tela do software", s["h2"]))
    story.append(P(T["ev1"], s["body"]))
    bloc = [P("Figura 1. Tela de acesso do sistema de apoio (login)", s["caption"])]
    if LOGIN_PNG.exists():
        bloc.append(fitted_image(LOGIN_PNG, 16.0, 10.0))
    bloc.append(P("Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).", s["fonte"]))
    story.append(KeepTogether(bloc))
    story.append(P(T["ev1b"], s["body"]))
    bloc = [P("Figura 2. Mapa de localidade dos vencedores das licitações (itens homologados)", s["caption"])]
    if MAPA_PNG.exists():
        bloc.append(fitted_image(MAPA_PNG, 16.2, 11.2))
    bloc.append(P("Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).", s["fonte"]))
    story.append(KeepTogether(bloc))
    story.append(P(T["ev1c"], s["body"]))
    bloc = [P("Figura 3. Tela de CNPJs vencedores das licitações (fornecedores homologados)", s["caption"])]
    if CNPJS_PNG.exists():
        bloc.append(fitted_image(CNPJS_PNG, 16.2, 11.2))
    bloc.append(P("Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P("4.2 Fluxo antes e depois", s["h2"]))
    story.append(P(T["ev2"], s["body"]))
    bloc = [P("Figura 4. Rotina de acompanhamento antes e depois do sistema de apoio", s["caption"])]
    if (EVID / "04_fluxo.png").exists():
        bloc.append(fitted_image(EVID / "04_fluxo.png", 16.2, 7.2))
    bloc.append(P("Fonte: registros do projeto (2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P("4.3 Arquitetura do sistema", s["h2"]))
    story.append(P(T["ev3"], s["body"]))
    bloc = [P("Figura 5. Arquitetura lógica do sistema de apoio", s["caption"])]
    if (EVID / "05_arquitetura.png").exists():
        bloc.append(fitted_image(EVID / "05_arquitetura.png", 16.2, 8.2))
    bloc.append(P("Fonte: registros do projeto (2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P("4.4 Volume carregado das fontes oficiais", s["h2"]))
    story.append(P(T["ev4"], s["body"]))
    story.append(abnt_table(["Conjunto", "2025", "2026 (parcial)", "Observação"], TAB1, [4.4 * cm, 2.8 * cm, 3.6 * cm, 6.0 * cm], s))
    story.append(P("Tabela 1. Volume carregado a partir do painel municipal (corte em 18/08/2026)", s["caption"]))
    story.append(P("Fonte: CSVs oficiais do painel da PMU, consolidados pelo sistema (2026).", s["fonte"]))
    story.append(P(T["ev5"], s["body"]))
    bloc = [P("Figura 6. Processos licitatórios de 2025 por modalidade (painel municipal)", s["caption"])]
    if (EVID / "06_modalidades_2025.png").exists():
        bloc.append(fitted_image(EVID / "06_modalidades_2025.png", 15.4, 8.0))
    bloc.append(P("Fonte: CSVs oficiais do painel da PMU (2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P("4.5 Recados da entidade e o artigo", s["h2"]))
    story.append(P(T["ev6"], s["body"]))

    story.append(P("4.6 Exemplo de risco: se a API externa para, o módulo Compras.gov para", s["h2"]))
    story.append(P(T["ev7"], s["body"]))
    story.append(abnt_table(["Campo", "Registro"], Q5_CHAMADO, [5.2 * cm, 11.6 * cm], s))
    story.append(P("Quadro 5. Demonstração de falha da API externa (chamado 56398424)", s["caption"]))
    story.append(P("Fonte: Portal de Serviços, Comprasnet, APIs PNCP e Compras.gov (ago. 2026).", s["fonte"]))

    story.append(P("5 REFERÊNCIAS", s["h1"]))
    for r in REFS:
        story.append(P(r, s["ref"]))

    doc = SimpleDocTemplate(
        str(OUT_PDF), pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm, topMargin=2.2 * cm, bottomMargin=2.2 * cm,
        title="Relatorio Semana 1: sistematizacao de evidencias",
        author="Diogo Ferreira Moura",
        subject="TCC UNIUBE Relato de experiencia do Projeto Integrado",
    )
    doc.build(story, onFirstPage=first_page, onLaterPages=later_pages)


def _set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _docx_p(doc, text, *, size=12, bold=False, italic=False, align="justify",
            first_line=1.25, space_after=8, space_before=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.5
    if first_line:
        pf.first_line_indent = Cm(first_line)
    p.alignment = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def _docx_caption(doc, text):
    return _docx_p(doc, text, size=10, bold=True, align="center", first_line=0, space_before=8, space_after=2)


def _docx_fonte(doc, text):
    return _docx_p(doc, text, size=9, align="center", first_line=0, space_after=10)


def _docx_img(doc, path: Path, width_cm: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def _docx_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _set_run_font(run, size=10, bold=False)
    doc.add_paragraph()
    return table


def build_docx() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    _docx_p(doc, "UNIVERSIDADE DE UBERABA", size=14, bold=True, align="center", first_line=0, space_after=2)
    _docx_p(doc, "UNIUBE", size=12, align="center", first_line=0, space_after=2)
    _docx_p(doc, "Curso de Sistemas de Informação", size=12, align="center", first_line=0, space_after=2)
    _docx_p(doc, "Polo Uberlândia", size=12, align="center", first_line=0, space_after=18)
    _docx_p(doc, T["cover_comp_docx"], size=12, italic=True, align="center", first_line=0, space_after=12)
    _docx_p(doc, T["cover_title"], size=16, bold=True, align="center", first_line=0, space_after=8)
    _docx_p(doc, T["cover_sub"], size=12, italic=True, align="center", first_line=0, space_after=24)
    _docx_p(doc, "Diôgo Ferreira Moura", size=12, align="center", first_line=0, space_after=2)
    _docx_p(doc, "RA 1030125-2", size=12, align="center", first_line=0, space_after=36)
    _docx_p(doc, T["cover_city"], size=12, align="center", first_line=0, space_after=2)
    _docx_p(doc, "Agosto de 2026", size=12, align="center", first_line=0, space_after=12)
    doc.add_page_break()

    _docx_p(doc, "SUMÁRIO", size=13, bold=True, align="left", first_line=0, space_after=12)
    for k in ("toc1", "toc2", "toc3", "toc4", "toc5"):
        _docx_p(doc, T[k], size=12, align="left", first_line=0, space_after=4)
    doc.add_page_break()

    _docx_p(doc, T["h1"], size=13, bold=True, align="left", first_line=0)
    _docx_p(doc, T["intro"])
    _docx_p(doc, "1.1 Identificação da organização parceira", size=12, bold=True, align="left", first_line=0, space_before=10)
    if (EVID / "00_logo_osb.png").exists():
        _docx_img(doc, EVID / "00_logo_osb.png", 6.5)
    _docx_p(doc, T["org1"])
    _docx_p(doc, T["org2"])
    _docx_table(doc, ["Campo", "Informação"], Q1_ORG)
    _docx_caption(doc, "Quadro 1. Identificação da organização parceira")
    _docx_fonte(doc, "Fonte: registros do projeto e material da entidade (2026).")

    _docx_p(doc, "1.2 Problema identificado", size=12, bold=True, align="left", first_line=0, space_before=10)
    _docx_p(doc, T["prob1"])
    _docx_p(doc, T["prob2"])
    _docx_p(doc, "1.3 Cenário e contexto encontrados", size=12, bold=True, align="left", first_line=0, space_before=10)
    _docx_p(doc, T["cen1"])
    _docx_p(doc, "As unidades compradoras que eles já acompanhavam no portal Comprasnet, e que orientei o recorte do sistema, são estas. O portal eles já usavam; a API de transparência dessas mesmas UASGs, não:")
    _docx_table(doc, ["Sigla", "Órgão", "UASG"], Q2_UASG)
    _docx_caption(doc, "Quadro 2. Unidades compradoras acompanhadas em Uberlândia")
    _docx_fonte(doc, "Fonte: rotina da entidade e API Compras.gov (2026).")
    _docx_p(doc, T["cen2"])

    _docx_p(doc, "2 INTERAÇÃO COM A ORGANIZAÇÃO PARCEIRA", size=13, bold=True, align="left", first_line=0, space_before=14)
    _docx_p(doc, "2.1 Como o contato se estabeleceu", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["int1"])
    _docx_p(doc, T["int2"])
    _docx_p(doc, "2.2 Facilidades encontradas", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["fac1"])
    _docx_p(doc, T["fac2"])
    q = _docx_p(doc, T["q1"], italic=True, first_line=0)
    q.paragraph_format.left_indent = Cm(1.25)
    _docx_p(doc, T["q1src"], size=10, align="right", first_line=0)
    _docx_p(doc, T["fac3"])
    q = _docx_p(doc, T["q2"], italic=True, first_line=0)
    q.paragraph_format.left_indent = Cm(1.25)
    _docx_p(doc, T["q2src"], size=10, align="right", first_line=0)
    _docx_p(doc, T["fac4"])
    _docx_p(doc, T["fac5"])
    _docx_p(doc, "2.3 Dificuldades e o que precisei mudar", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["dif1"])
    _docx_p(doc, T["dif1b"])
    _docx_p(doc, T["dif2"])
    _docx_p(doc, T["dif3"])
    _docx_table(doc, ["Data", "O que aconteceu"], Q3_INT)
    _docx_caption(doc, "Quadro 3. Síntese da interação com a organização parceira")
    _docx_fonte(doc, "Fonte: WhatsApp, reuniões do projeto e Portal de Serviços (2026).")

    _docx_p(doc, "3 PROPOSTA DE SOLUÇÃO (INTERVENÇÃO) ADOTADA", size=13, bold=True, align="left", first_line=0, space_before=14)
    _docx_p(doc, "3.1 O sistema de apoio criado", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["sol0"])
    _docx_p(doc, T["sol1"])
    _docx_p(doc, T["sol2"])
    _docx_p(doc, "3.2 Escolha técnica e tamanho da operação", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["tec1"])
    _docx_p(doc, "3.3 Telas que ficaram na rotina", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_table(doc, ["Módulo", "Para que serve"], Q4_MOD)
    _docx_caption(doc, "Quadro 4. Módulos do sistema de apoio e uso previsto")
    _docx_fonte(doc, "Fonte: sistema de apoio implantado (2026).")
    _docx_p(doc, T["lim1"])
    _docx_p(doc, T["lim2"])

    _docx_p(doc, "4 REGISTROS DE EVIDÊNCIAS DA AÇÃO REALIZADA", size=13, bold=True, align="left", first_line=0, space_before=14)
    _docx_p(doc, T["ev0"])
    _docx_p(doc, "4.1 Print da tela do software", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["ev1"])
    if LOGIN_PNG.exists():
        _docx_img(doc, LOGIN_PNG, 15.5)
    _docx_caption(doc, "Figura 1. Tela de acesso do sistema de apoio (login)")
    _docx_fonte(doc, "Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).")
    _docx_p(doc, T["ev1b"])
    if MAPA_PNG.exists():
        _docx_img(doc, MAPA_PNG, 16.0)
    _docx_caption(doc, "Figura 2. Mapa de localidade dos vencedores das licitações (itens homologados)")
    _docx_fonte(doc, "Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).")
    _docx_p(doc, T["ev1c"])
    if CNPJS_PNG.exists():
        _docx_img(doc, CNPJS_PNG, 16.0)
    _docx_caption(doc, "Figura 3. Tela de CNPJs vencedores das licitações (fornecedores homologados)")
    _docx_fonte(doc, "Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).")
    _docx_p(doc, "4.2 Fluxo antes e depois", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["ev2"])
    if (EVID / "04_fluxo.png").exists():
        _docx_img(doc, EVID / "04_fluxo.png", 16.0)
    _docx_caption(doc, "Figura 4. Rotina de acompanhamento antes e depois do sistema de apoio")
    _docx_fonte(doc, "Fonte: registros do projeto (2026).")
    _docx_p(doc, "4.3 Arquitetura do sistema", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["ev3"])
    if (EVID / "05_arquitetura.png").exists():
        _docx_img(doc, EVID / "05_arquitetura.png", 16.0)
    _docx_caption(doc, "Figura 5. Arquitetura lógica do sistema de apoio")
    _docx_fonte(doc, "Fonte: registros do projeto (2026).")
    _docx_p(doc, "4.4 Volume carregado das fontes oficiais", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["ev4"])
    _docx_table(doc, ["Conjunto", "2025", "2026 (parcial)", "Observação"], TAB1)
    _docx_caption(doc, "Tabela 1. Volume carregado a partir do painel municipal (corte em 18/08/2026)")
    _docx_fonte(doc, "Fonte: CSVs oficiais do painel da PMU, consolidados pelo sistema (2026).")
    _docx_p(doc, T["ev5"])
    if (EVID / "06_modalidades_2025.png").exists():
        _docx_img(doc, EVID / "06_modalidades_2025.png", 15.2)
    _docx_caption(doc, "Figura 6. Processos licitatórios de 2025 por modalidade (painel municipal)")
    _docx_fonte(doc, "Fonte: CSVs oficiais do painel da PMU (2026).")
    _docx_p(doc, "4.5 Recados da entidade e o artigo", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["ev6"])
    _docx_p(doc, "4.6 Exemplo de risco: se a API externa para, o módulo Compras.gov para", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["ev7"])
    _docx_table(doc, ["Campo", "Registro"], Q5_CHAMADO)
    _docx_caption(doc, "Quadro 5. Demonstração de falha da API externa (chamado 56398424)")
    _docx_fonte(doc, "Fonte: Portal de Serviços, Comprasnet, APIs PNCP e Compras.gov (ago. 2026).")

    _docx_p(doc, "5 REFERÊNCIAS", size=13, bold=True, align="left", first_line=0, space_before=14)
    for r in REFS_PLAIN:
        p = _docx_p(doc, r, size=11, first_line=0, space_after=8)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)

    doc.save(str(OUT_DOCX))


def assert_no_forbidden(path: Path) -> None:
    if path.suffix.lower() == ".pdf":
        from pypdf import PdfReader
        text = "".join((p.extract_text() or "") for p in PdfReader(str(path)).pages)
    elif path.suffix.lower() == ".docx":
        d = Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = "\n".join(parts)
    else:
        text = path.read_text(encoding="utf-8")
    hits = sorted({ch for ch in text if ch in FORBIDDEN})
    if hits:
        raise SystemExit(f"Ainda ha traco/aspas tipograficas em {path.name}: {hits!r}")


def main() -> None:
    src = Path(__file__).read_text(encoding="utf-8")
    for ch in FORBIDDEN:
        if ch in src:
            raise SystemExit(f"O gerador ainda contem o caractere {ch!r}. Tire antes de gerar.")
    prepare_assets()
    build_pdf()
    build_docx()
    assert_no_forbidden(OUT_PDF)
    assert_no_forbidden(OUT_DOCX)
    print("PDF:", OUT_PDF, OUT_PDF.stat().st_size)
    print("DOCX:", OUT_DOCX, OUT_DOCX.stat().st_size)


if __name__ == "__main__":
    main()
