#!/usr/bin/env python3
"""Relatorio Semana 5: documento completo com elementos pre, textuais e pos-textuais (ABNT / ACQA / TCC)."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path("/home/dfmoura/Documents/test_several1/uniube/4")
EVID = ROOT / "evidencias"
FIG_ARTIGO = Path("/home/dfmoura/Documents/test_several1/uniube/3/figuras")
LOGO_OSB = Path(
    "/home/dfmoura/Documents/test_several1/trigger/23/app/static/logo_oficial_udi.png"
)
LOGIN_PNG = EVID / "01_tela_login.png"
MAPA_PNG = EVID / "02_tela_mapa_homologados.png"
CNPJS_PNG = EVID / "07_tela_cnpjs_vencedores.png"
OUT_PDF = ROOT / "Relatorio_Semana5_Elementos_Relatorio_ABNT.pdf"
OUT_DOCX = ROOT / "Relatorio_Semana5_Elementos_Relatorio_ABNT.docx"

NAVY = colors.HexColor("#1F4E79")
GRAY = colors.HexColor("#4A4A4A")
FORBIDDEN = "\u2014\u2013\u2012\u2212\u2015\u2026\u00b7\u201c\u201d"

# Texto unico para PDF e DOCX. Sem traco tipografico (em dash).
T = {
    "cover_comp": "Trabalho de Conclusão de Curso:<br/>Relato de Experiência do Projeto Integrado",
    "cover_comp_docx": "Trabalho de Conclusão de Curso: Relato de Experiência do Projeto Integrado",
    "cover_title": (
        "Relatório sobre o desenvolvimento e a implantação de um sistema de informação de "
        "apoio ao Observatório Social do Brasil em Uberlândia/MG: coleta, consulta e "
        "análise de licitações públicas"
    ),
    "cover_sub": "Relatório acadêmico da atividade extensionista",
    "cover_city": "Uberlândia / MG",
    "toc1": "1 Introdução",
    "toc2": "2 Fundamentação teórica",
    "toc3": "3 Material e métodos",
    "toc4": "4 Resultados",
    "toc5": "5 Conclusão",
    "toc6": "Referências",
    "resumo": (
        "Este relatório apresenta a experiência extensionista desenvolvida junto ao Observatório "
        "Social do Brasil de Uberlândia, organização da sociedade civil que acompanha licitações "
        "dos órgãos municipais. O problema identificado na rotina da entidade não era a ausência "
        "de dado público, e sim a coleta manual: a equipe transcrevia processos do portal para a "
        "planilha Cronograma antes de analisar. O objetivo foi desenvolver, implantar e "
        "disponibilizar um sistema de informação de apoio à coleta, à consulta e à análise dessas "
        "licitações, com custo compatível com uma OSC sem equipe de TI. Do ponto de vista "
        "metodológico, trata-se de pesquisa aplicada com estudo de caso único. O material incluiu "
        "VM na AWS, FastAPI, SQLite, Docker Compose e as fontes oficiais (CSVs do painel Power BI "
        "da Prefeitura e, depois, a API de Dados Abertos do Compras.gov.br). O método combinou "
        "convívio semanal na entidade, desenvolvimento incremental e duas apresentações formais "
        "(ACIUB, 30 min; Teams com a sede, 1h30). O resultado principal é o sistema em produção "
        "em https://licitacoes.osbrasiluberlandia.org/, com telas de painel, processo, CNPJs "
        "vencedores, mapa de localidade e cobertura entre bases. No recorte municipal de 2025 "
        "constam 1.215 processos; gestores e fiscais passam de 36 mil linhas. Houve retorno "
        "positivo da entidade, pedido de filtro por porte e convite a voluntariado. Conclui-se "
        "que a intervenção atendeu ao pedido da organização parceira, com limites explícitos de "
        "hospedagem gratuita (free tier até janeiro de 2027) e de dependência da API externa."
    ),
    "palavras": (
        "Palavras-chave: sistema de informação; controle social; dados abertos; licitações "
        "públicas; extensão universitária."
    ),
    "intro0": (
        "Este relatório descreve, em formato acadêmico, a experiência que tive no Projeto Integrado "
        "de Extensão e que agora entra no Trabalho de Conclusão de Curso. A organização parceira foi "
        "o Observatório Social do Brasil de Uberlândia, o OSB. A ação que executei foi criar, implantar "
        "e colocar em uso um sistema de informação de apoio ao acompanhamento das licitações dos órgãos "
        "de Uberlândia. O endereço público, no domínio da entidade, é "
        "https://licitacoes.osbrasiluberlandia.org/. O documento organiza-se com elementos pré-textuais "
        "(capa, sumário e resumo), textuais (introdução, fundamentação teórica, material e métodos, "
        "resultados e conclusão) e pós-textuais (referências), conforme a NBR 14724, com citações "
        "pela NBR 10520 e referências pela NBR 6023 (ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS, "
        "2011, 2018, 2023). O que organizei nas semanas anteriores deste componente reaparece aqui "
        "já na estrutura completa de um trabalho acadêmico."
    ),
    "tema1": (
        "O tema deste estudo é o desenvolvimento e a implantação de um sistema de informação de apoio "
        "a uma organização da sociedade civil que acompanha compra pública municipal. Não é um curso, "
        "não é oficina e não é cartilha. É um software em produção, feito para a rotina de quem observa "
        "licitação em Uberlândia. A entidade parceira foi o Observatório Social do Brasil de Uberlândia. "
        "Organização da sociedade civil, apartidária, sem fins lucrativos. CNPJ 23.497.346/0001-42. "
        "Endereço que usei no acompanhamento: Av. Vasconcelos Costa, 1500, Sala 3, Anexo I, Bairro "
        "Martins, Uberlândia-MG, CEP 38400-452. Telefone (34) 3239-1529. Site: "
        "https://www.osbrasiluberlandia.org/. A unidade local existe desde 2015, ligada à associação G7 "
        "(a CDL entra nisso) e faz parte da rede nacional do Observatório Social do Brasil."
    ),
    "tema2": (
        "O que eles se propõem é despertar cidadania fiscal, no sentido de a sociedade acompanhar o uso "
        "do dinheiro público. Em Uberlândia isso aparece em duas frentes. Uma é o acompanhamento das "
        "licitações da Prefeitura e das autarquias e fundações (DMAE, FUTEL, EMAM, IPREMU, FERUB, "
        "ARESAN, PRODAUB e a Câmara). A outra é olhar a atuação dos vereadores. Eu atuei na primeira "
        "frente. Foi a que mais consumia tempo de planilha e foi o pedido mais claro da entidade: um "
        "sistema de apoio para reunir, consultar e analisar essas licitações, em vez de copiar processo "
        "por processo do portal para a planilha Cronograma. O pedido era a ferramenta. O que a API era, "
        "e o fato de já existir uma API oficial de transparência para essas compras, não estava no "
        "vocabulário da equipe. Isso veio depois, de forma indireta, junto com o uso do sistema."
    ),
    "prob1": (
        "Dado público, no papel, não falta. Tem a Lei de Acesso à Informação, o decreto de dados abertos "
        "e a Lei 14.133, com o PNCP (BRASIL, 2011, 2016, 2021). O problema que eu vi no escritório era "
        "outro: o dado existia, mas o uso ainda era manual. Toda semana alguém entra no portal, copia "
        "processo, cola na planilha Cronograma e só depois analisa. Silva (2024), ao tratar do potencial "
        "de reúso dos dados abertos do governo brasileiro, insiste no mesmo ponto: publicar não é o "
        "mesmo que usar. Foi esse desencontro que justificou a intervenção: "
        "criar um sistema de apoio à coleta e à consulta, no tamanho da entidade. Não era falha da API "
        "federal. A API de Dados Abertos do Compras.gov entrou depois, quando o sistema já existia. "
        "O problema da entidade era a digitação repetida."
    ),
    "prob2": (
        "A rotina que me mostraram é mais ou menos assim. Abre o site da prefeitura, transcreve para o "
        "Cronograma, monta o acompanhamento. Ao mesmo tempo consulta o Comprasnet pelas UASGs de "
        "Uberlândia e o Power BI da PMU. Esse painel atrasava, eles mesmos comentaram. Cada fonte "
        "descreve o processo de um jeito, com campo e código diferentes. O relatório de quadrimestre "
        "depende disso. Se a semana atrasava, o quadrimestre atrasava junto. O Comprasnet, para eles, "
        "era o site em que a pessoa pesquisa. Não havia, na rotina, a ideia de API: um canal em que o "
        "dado oficial já sai pronto para o sistema consumir, sem copiar a tela. Essa API de transparência "
        "das compras públicas já existia; o que faltava era o conhecimento e o uso."
    ),
    "just1": (
        "A justificativa é prosaica. Toda semana o tempo da equipe ia para copiar processo. Sem gente "
        "de TI e sem verba fixa para servidor, um sistema grande não ia vingar. O componente pedia ação "
        "extensionista em Sistemas de Informação. O pedido da entidade era uma ferramenta de apoio. "
        "Essas duas coisas bateram. Pinho e Sacramento (2009), ao tratarem de accountability, lembram "
        "que prestar contas, ser transparente e responder por isso não cabem numa palavra só no "
        "português. No OSB isso vira coisa de mesa: ler edital, conferir quem ganhou, fechar o "
        "quadrimestre. Um sistema que organize a consulta entra exatamente aí, sem pretender substituir "
        "o olhar de quem observa."
    ),
    "just2": (
        "Há ainda o tamanho da operação. É voluntariado. Reunião quase toda semana, em geral na "
        "associação comercial. Prestação de contas a cada quatro meses. Quatro usuários no sistema já "
        "dão conta. Laudon e Laudon (2014) tratam o sistema de informação como conjunto que coleta, "
        "processa, armazena e distribui dado para apoiar decisão. Aqui isso só se sustenta se a "
        "ferramenta couber no bolso e na cabeça de quem vai usar. Por isso a intervenção não foi um "
        "diagnóstico encerrado em relatório: foi software no ar, no domínio deles, com coleta que pode "
        "rodar de madrugada e tela em que a análise começa já com o dado organizado."
    ),
    "obj1": (
        "O objetivo geral foi desenvolver, implantar e disponibilizar um sistema de informação de apoio "
        "ao Observatório Social do Brasil de Uberlândia, para a coleta, a consulta e a análise das "
        "licitações dos órgãos do município, a partir de fontes oficiais e com custo de operação "
        "compatível com uma OSC sem equipe de TI."
    ),
    "obj2": (
        "Os objetivos específicos foram quatro. Primeiro, identificar a rotina da entidade, o problema "
        "da coleta manual e o cenário em que o trabalho acontece. Segundo, construir o sistema a partir "
        "das fontes que eles já usavam, sem apagar o rastro da origem. Terceiro, colocar a ferramenta "
        "em uso, com contas limitadas e endereço no domínio da entidade. Quarto, registrar a interação, "
        "as evidências da ação e o retorno da equipe. Não entrou como objetivo consertar API federal, "
        "nem cobrir a frente dos vereadores. Isso ficou de fora de propósito."
    ),
    "teo1": (
        "Controle social, no vocabulário da administração pública brasileira, é a participação da "
        "sociedade na fiscalização da gestão. Não substitui controle interno nem tribunal de contas. "
        "Ocupa um lugar ao lado, com outro tempo. A rede dos Observatórios Sociais do Brasil se "
        "apresenta nesse hiato: cidadania fiscal, núcleo municipal, voluntariado, reunião semanal e "
        "prestação de contas quadrimestral. Um sistema que ignore essa restrição de gente e de horário "
        "tende a ficar sem uso, mesmo com modelo de dados correto."
    ),
    "teo2": (
        "A literatura de dados abertos insiste no mesmo ponto por outro ângulo. Silva (2024), ao "
        "revisar a métrica DGABr, parte da premissa de que publicar dado aberto não basta: o que "
        "importa é o potencial de reúso. Sem capacidade analítica, o portal é pouco utilizado. No "
        "município de médio porte existem painel em Power BI, portal de "
        "licitações e PNCP, e mesmo assim o observador continua recortando informação à mão, porque o "
        "dado oficial chega em recortes distintos e o cruzamento precisa ser feito por quem consulta. "
        "A Lei n. 12.527/2011 impôs a publicidade como regra (BRASIL, 2011). O Decreto n. 8.777/2016 "
        "tratou da política federal de dados abertos (BRASIL, 2016). A Lei n. 14.133/2021 reorganizou "
        "modalidades e consolidou o PNCP (BRASIL, 2021). Lei nenhuma, sozinha, tira a planilha Cronograma "
        "da mesa."
    ),
    "teo3": (
        "Sistemas de informação, nesse contexto, não são um fim. Laudon e Laudon (2014) classificam "
        "sistemas de apoio à decisão como aqueles que combinam dados e modelos para ajudar em problema "
        "pouco estruturado. O problema do Observatório é desse tipo. O objetivo não é emitir empenho. "
        "É observar um conjunto de pregões e verificar se o padrão de vitórias se repete, se o preço "
        "destoa, se o fiscal está nomeado. O sistema não responde sozinho. Ele organiza a consulta. "
        "Moreira et al. (2020), em estudo de acesso aberto sobre qualidade na recuperação de dados "
        "governamentais, tratam dimensões como acurácia, completude, consistência e pontualidade. No "
        "projeto, isso virou regra simples: não inventar campo que a fonte não tem; não apagar anotação "
        "do observador quando a coleta roda de novo; aceitar que o painel municipal atrasa em relação "
        "ao PNCP e mostrar as duas pontas, em vez de escolher uma e fingir que a outra não existe."
    ),
    "teo4": (
        "Do lado da engenharia, Pressman e Maxim (2021) e Sommerville (2018) sustentam o desenvolvimento "
        "que vai se ajustando pelo uso, quando o requisito não nasce de um edital interno. Foi o caso. "
        "A primeira ideia de raspar o portal da prefeitura com navegador quebrou fácil: layout, sessão, "
        "máquina com tela virtual. A fonte estável foram os CSVs oficiais do painel. A API de Dados "
        "Abertos do Compras.gov.br entrou depois, para o recorte federal das UASGs que eles já "
        "acompanhavam no portal. Troca de caminho no meio, sim. Porque o requisito ficou mais claro "
        "depois que a ferramenta começou a ser vista na reunião, e não antes."
    ),
    "met0": (
        "A ação extensionista não foi um curso. Não houve turma, apostila nem carga horária de sala. "
        "O que executei foi o desenvolvimento, a implantação e a disponibilização de um sistema de "
        "informação de apoio, com convívio semanal na entidade e duas apresentações formais. Esta "
        "seção descreve o material empregado e o método seguido, no sentido de permitir que outra "
        "pessoa compreenda como a ação foi conduzida e, se quiser, reproduza o percurso em outro "
        "contexto."
    ),
    "met1": (
        "O trabalho mistura pesquisa aplicada com estudo de caso único (GIL, 2022; YIN, 2015). O caso "
        "é o Observatório Social do Brasil de Uberlândia/MG (CNPJ 23.497.346/0001-42). Não pretendi "
        "amostra estatística de OSCs brasileiras. O intuito foi entender uma operação concreta, "
        "construir o sistema e devolvê-lo a quem pediu. Yin (2015) cabe porque as perguntas foram do "
        "tipo como e por que: como a equipe junta as fontes hoje, por que o painel municipal não basta "
        "e como um sistema pequeno pode entrar na rotina sem ficar abandonado."
    ),
    "met2": (
        "O contato com a entidade começou em 21 de maio de 2026, às 14h30, depois de uma articulação "
        "que deu volta pela CDL. De maio a agosto o convívio foi quase toda semana, em geral na ACIUB, "
        "com datashow. A coleta de requisito não começou com gravador. Começou por sentar na rotina "
        "deles, ver a planilha Cronograma, o acompanhamento, o Comprasnet e o que mais tomava tempo: "
        "achar o processo, cruzar fonte, olhar o fornecedor e fechar o quadrimestre. Material "
        "instrucional, no sentido de apostila, eu não preparei. O material foi o próprio sistema, as "
        "telas, e depois o artigo que enviei à Revista de Engenharia, TI e Inovação (RETII, ISSN "
        "2966-2508), da Uniube, com o mesmo tema deste relatório."
    ),
    "met3": (
        "Houve duas ofertas formais da ferramenta, além do uso corrente nas reuniões. No dia 23 de "
        "julho de 2026, quinta-feira, das 15h30 às 16h, apresentei na sala de reunião da ACIUB, para "
        "a diretoria do Observatório, gente da CDL e da superintendência da ACIUB. Cerca de meia hora. "
        "Em 3 de agosto, segunda-feira, das 16h às 17h30, apresentei de novo, desta vez por Microsoft "
        "Teams, com o pessoal da sede do Observatório. Uma hora e meia. Nas duas a forma foi a mesma: "
        "mostrar o sistema no ar, as telas de consulta e, de passagem, o que é uma API e que o governo "
        "já publica dado aberto de compra por esse canal. Não foi aula isolada. Entrou no uso."
    ),
    "met4": (
        "O material foi escolhido para caber no bolso e na operação da entidade. Hardware: uma VM "
        "na AWS, região de São Paulo, Ubuntu, 1 GB de RAM, 2 vCPUs e 40 GB de disco, no plano "
        "gratuito. Software: FastAPI, SQLite no arquivo data/licitacoes.db, tela estática no mesmo "
        "serviço e Docker Compose. Contas: no máximo quatro (um admin e três de consulta). Coleta, "
        "Setup, disparo de CNPJ e token de IA só o admin mexe. Não usei RabbitMQ, MinIO nem banco "
        "gerenciado. Não precisava. Backup diário do SQLite, com pouca retenção, porque disco grátis "
        "acaba rápido. O free tier desta conta acaba em 13 de janeiro de 2027, ou antes se o crédito "
        "acabar. Falei isso para eles. Fontes oficiais adotadas: CSVs do painel Power BI da "
        "Prefeitura (licitações, contratos e gestores/fiscais) e, depois, a API de Dados Abertos do "
        "Compras.gov.br, filtrada para Uberlândia (IBGE 3170206) e para as UASGs que eles já "
        "acompanhavam. O portal weblicitacoes.uberlandia.mg.gov.br ficou para conferência manual."
    ),
    "met5": (
        "O método foi incremental. Cada fatia (coleta, consulta, painel, autenticação, agendamento, "
        "perfil de CNPJ, mapa, análise auxiliar de preço) só avançava depois de a anterior estar "
        "utilizável. A coleta de requisito começou por sentar na rotina deles: planilha Cronograma, "
        "acompanhamento, Comprasnet e o que mais tomava tempo. Desenvolvimento e implantação "
        "correram de maio a agosto de 2026, com reunião quase toda semana. Houve duas ofertas "
        "formais: 23 de julho, presencial na ACIUB (cerca de 30 min); 3 de agosto, Teams com a sede "
        "(1h30). Nas duas mostrei o sistema no ar e, de passagem, o que é uma API. A regra na "
        "persistência foi manter o significado dos campos da fonte e não apagar anotação do "
        "observador quando a coleta roda de novo. A avaliação foi de uso e de consistência, não de "
        "experimento com cronômetro. Não medi em minutos o tempo antes e depois. O ganho que eles "
        "relatam é de rearranjo do trabalho."
    ),
    "res0": (
        "Os resultados a seguir registram o que foi obtido na execução da atividade extensionista: "
        "a interação com a organização parceira, as evidências do sistema em uso e, ao final, um "
        "exemplo pontual de risco da arquitetura quando a API externa falha. Esse exemplo não "
        "redefine o problema da entidade. Não é protótipo de disciplina. Há print de tela, desenho "
        "do fluxo, arquitetura, volume carregado e retorno da equipe."
    ),
    "int1": (
        "Chegar no Observatório deu volta. Em março eu tinha pensado num treinamento de inteligência "
        "artificial para associados da CDL. Pensei no Antonio Carlos Oliveira porque ele tinha sido "
        "meu professor no curso de Administração, onde me formei, na ESAMC Uberlândia. Conversei com "
        "ele no sentido de pedir orientação: como eu poderia dar andamento no trabalho, a quem falar, "
        "por onde começar. Ele sugeriu procurar a Lécia Queiroz, da CDL. O resto foi por minha conta. "
        "No dia 20 de maio tive reunião presencial com a Lécia na CDL, Av. Belo Horizonte 1290. Ficou "
        "claro que colocar o projeto dentro da CDL naquele momento emperrava. Ela me passou o Marco "
        "Aurélio Freitas Santos, do OSB. Marcamos no mesmo dia. Fui em 21 de maio, 14h30."
    ),
    "int2": (
        "O que emperrou foi essa articulação inicial, não o convívio depois. O Antonio Carlos só "
        "apontou o caminho. Marcar a Lécia, ir na CDL, falar com o Marco e tocar o sistema foi comigo. "
        "Com o OSB o contato foi direto. Quase toda semana."
    ),
    "fac1": (
        "O Marco recebeu o projeto como sistema de apoio ao trabalho de acompanhamento, não como visita "
        "de faculdade. A gente se via na rotina deles, na ACIUB, com datashow. Mostrei o que o sistema "
        "já fazia e o que ainda estava no meio. Eles me explicaram a planilha Cronograma, o "
        "acompanhamento, e o que mais tomava tempo: achar o processo, cruzar com o Comprasnet, olhar o "
        "fornecedor e fechar o quadrimestre. Foi essa rotina que o software passou a apoiar. Nas mesmas "
        "reuniões eu expliquei, sem jargão de curso, o que é uma API: o jeito de o sistema buscar o "
        "dado direto da fonte, sem a pessoa copiar a tela. Mostrei que o governo já oferece isso para "
        "transparência das operações públicas, principalmente a API de Dados Abertos do Compras.gov.br. "
        "Esse conhecimento não era o objetivo da atividade; veio junto, de forma indireta."
    ),
    "fac2": (
        "No dia 11 de junho, depois de uma dessas reuniões, o Marco mandou a pergunta que virou tela "
        "no sistema:"
    ),
    "q1": (
        '"Seria possível extrairmos estatísticas de participação de empresas, % de vitória, '
        "quais tipos de objetos ou licitação participa e % de empresas que participam quando "
        "esta participa? o que gostaria de avaliar, existem histórias que empresas participam "
        "para favorecer uma e vao fazendo rodízio nas vitórias, então gostaria de analisar "
        'pela estatística."'
    ),
    "q1src": "Marco Aurélio Freitas Santos, WhatsApp, 11 jun. 2026.",
    "fac3": (
        "Na mesma conversa ele me chamou para ser voluntário. Combinamos de deixar isso mais formal "
        "depois que a faculdade acabar, para não virar obrigação de curso interno da rede. Também "
        "pediram um link no site deles. Ficou https://licitacoes.osbrasiluberlandia.org/. Em 15 de "
        "julho a Lécia me escreveu dizendo que o Marco estava contente com o que estava saindo:"
    ),
    "q2": '"Que legal! O Marco tem me informado, ele está muito contente com os resultados!"',
    "q2src": "Lécia Queiroz, CDL Uberlândia, WhatsApp, 15 jul. 2026.",
    "fac4": (
        "Isso não ficou só no WhatsApp. No dia 23 de julho, quinta-feira, das 15h30 às 16h, organizei "
        "uma apresentação na sala de reunião da ACIUB. O convite foi para a diretoria do Observatório, "
        "gente da CDL e da superintendência da ACIUB. O tema era o sistema de apoio que eu havia "
        "criado: a busca das licitações e a automação da coleta. Gostaram bastante da ferramenta. "
        "Além da tela, mostrei que a coleta automática não é um truque no site: usa uma API já "
        "existente de transparência das compras públicas, a de Dados Abertos do Compras.gov.br, que a "
        "maior parte da plateia não conhecia pelo nome."
    ),
    "fac5": (
        "Em 3 de agosto, segunda-feira, das 16h às 17h30, apresentei de novo, desta vez por Microsoft "
        "Teams, com o pessoal da sede do Observatório. Também gostaram. Chegaram a pedir um filtro por "
        "porte de empresa, para olhar se quem ganha é ME, EPP ou empresa maior. Eu coloquei isso "
        "depois, na tela de CNPJs vencedores. Na mesma reunião expliquei de novo o que é API e que essa "
        "API federal de transparência já estava disponível, mesmo antes do nosso sistema. Para quem "
        "observa licitação no portal, isso costuma passar despercebido."
    ),
    "dif1": (
        "Tive dificuldade no começo com a CDL. O Antonio Carlos orientou, mas não intermediou a "
        "reunião. Sem o encaminhamento da Lécia, depois que eu mesmo a procurei, acho que o prazo da "
        "disciplina teria travado. No OSB eu não quis chegar com sistema pronto. Quis entender a "
        "planilha primeiro. Isso atrasou eu mostrar tudo na nuvem, mas evitou empurrar um jeito de "
        "trabalhar que não era o deles."
    ),
    "dif1b": (
        "Hospedagem de graça enrolou de verdade. Testei plano da Oracle de madrugada. Tive uma reunião "
        "com o pessoal da Nuvolli para ver se dava para ter nuvem sem custo. No fim subi numa VM da "
        "AWS, região de São Paulo, Ubuntu, 1 GB de RAM, 2 vCPUs e 40 GB de disco, no plano gratuito. "
        "Tem crédito e prazo. O free tier desta conta acaba em 13 de janeiro de 2027, ou antes se o "
        "crédito acabar. Quando isso acontecer, se ninguém pagar o plano ou mudar de máquina, o "
        "sistema sai do ar. Falei isso para eles. Não é solução eterna. É o que cabia no bolso da "
        "entidade agora."
    ),
    "dif2": (
        "Troquei de caminho técnico no meio. Primeiro pensei em raspar o portal da prefeitura com "
        "navegador. Quebrou fácil: layout, sessão, máquina com tela virtual. O sistema nasceu assim, "
        "com os CSVs oficiais do painel municipal. A API de Dados Abertos do Compras.gov.br entrou "
        "depois, quando a ferramenta já existia, para puxar o recorte federal das UASGs que eles já "
        "acompanhavam no portal Comprasnet. Pressman e Maxim (2021) e Sommerville (2018) tratam isso "
        "como desenvolvimento que vai se ajustando pelo uso. Foi o caso. Quando a API entrou, precisei "
        "explicar o que ela é e por que o portal que eles já usavam e a API não são a mesma coisa. "
        "Para a equipe, isso era novidade: existia um canal oficial de transparência feito para "
        "máquina, e não só a tela de pesquisa."
    ),
    "dif3": (
        "Depois que o sistema já estava no ar, apareceu uma dificuldade que não era da rotina da "
        "entidade, e sim do tipo de solução que eu escolhi. O módulo Compras.gov depende de uma API "
        "federal de dados abertos. Se essa API deixa de devolver registro, aquela ponta da coleta para "
        "de atualizar, mesmo com a compra visível no portal. Isso é um risco passível de qualquer "
        "sistema que consome fonte externa. Em agosto ocorreu um caso desses; o detalhe fica mais "
        "abaixo, só como exemplo, não como o problema que o Observatório me pediu para resolver. "
        "Avisei a entidade do limite. Só deu para avisar com clareza porque, no meio do projeto, a "
        "ideia de API já tinha sido apresentada: não é o site que a pessoa vê; é o canal que o "
        "sistema consome."
    ),
    "sol0": (
        "A proposta de solução que adotei foi desenvolver, implantar e disponibilizar o sistema de "
        "apoio. O software junta as fontes oficiais, grava numa base local e oferece telas de consulta "
        "(painel, processo, CNPJs vencedores, mapa de localidade e cobertura entre bases). Quem observa "
        "deixa de começar o trabalho copiando portal para planilha e passa a começar já com o dado "
        "organizado. A análise crítica continua sendo da entidade; o que o sistema tira é a digitação "
        "repetida. A planilha Cronograma ainda pode ficar do lado enquanto eles pegam confiança. Trocar "
        "ferramenta de uma hora para outra, nesse tipo de entidade, costuma virar ferramenta abandonada."
    ),
    "ev1": (
        "A Figura 1 é a tela de login, com a logo do Observatório. Sem usuário e senha a base não fica "
        "aberta na internet. Tirei o print na instância local, porta 8096, a mesma aplicação que está "
        "no endereço público."
    ),
    "ev1b": (
        "A Figura 2 é a tela Mapa de localidade. Mostra a sede dos vencedores das licitações, a partir "
        "dos itens homologados. Em cima ficam filtros de período, ano, órgão, modalidade, UF do "
        "vencedor, porte da empresa e a métrica do mapa (itens homologados, contratações ou valor). Os "
        "cards resumem o recorte: itens, contratações, valor homologado, UFs, municípios e a fatia que "
        "ficou em Uberlândia versus o que veio de fora. No centro, o mapa do Brasil, por estado ou "
        "calor de município. Do lado, o ranking por UF. No print: 6.797 itens homologados, 1.270 "
        "contratações, cerca de R$ 1,2 bilhão, 25 UFs e 390 municípios. Cerca de 25,5% dos itens em "
        "Uberlândia e 74,5% de fora. A lista nominal desses vencedores, por CNPJ, fica na Figura 3. "
        "Mesma instância local da Figura 1."
    ),
    "ev1c": (
        "A Figura 3 é a tela CNPJs vencedores. Não é o mapa da sede: é a lista de quem ganhou. Cada "
        "linha é um fornecedor (CNPJ ou CPF), com quantos itens levou, em quantas compras, o valor "
        "homologado e o município da sede. Em cima ficam filtros de período, ano, órgão, modalidade, "
        "UF do vencedor, cache do cadastro e porte da empresa, além da busca por nome ou CNPJ. O filtro "
        "de porte foi o que a sede do Observatório pediu em 3 de agosto. No print: 2.037 fornecedores "
        "consolidados, 16 com cadastro atualizado, 1.833 com cache vencido e 188 CPF. A origem é o "
        "resultado homologado do Compras.gov, enriquecido com dados públicos do CNPJ. Essa tela "
        "responde à pergunta do Marco sobre estatística de vitória: quem ganha, com que frequência e "
        "em que valor. Mesma instância local das Figuras 1 e 2."
    ),
    "ev2": (
        "A Figura 4 mostra o antes e o depois da intervenção. Antes, a análise só começava depois de "
        "copiar processo do portal para a planilha. Depois da implantação do sistema de apoio, a coleta "
        "alimenta a base local e o tempo da reunião vai para ler o processo, e não para digitar de novo "
        "o que o portal já tinha. O antes também era conceitual: o dado só existia na tela do portal. "
        "O depois inclui saber que o governo já publica o mesmo conteúdo por API de transparência, em "
        "especial a de Dados Abertos do Compras.gov.br."
    ),
    "ev3": (
        "A Figura 5 mostra de onde vem o dado (API federal, CSV do painel, cadastro de UASG), o hub de "
        "coleta de madrugada, o SQLite e as telas. A ponta Compras.gov foi encaixada depois que o "
        "sistema já existia. Nenhuma peça é novidade sozinha. O que importa é caber no tamanho da "
        "entidade e não precisar de gente de infraestrutura o tempo todo. A figura também serviu para "
        "mostrar à equipe de onde o dado vem, inclusive a API federal, que antes não fazia parte do "
        "vocabulário deles. Se essa API deixar de responder, só aquela coleta para; o restante do "
        "sistema segue com o CSV municipal."
    ),
    "ev4": (
        "A Tabela 1 sai dos CSVs oficiais do painel da prefeitura, corte em 18 de agosto de 2026. Não "
        "é o universo do PNCP. É o recorte com o qual eles já trabalhavam. Gestores e fiscais passam de "
        "36 mil linhas. Ninguém ia conferir isso na planilha Cronograma."
    ),
    "ev5": (
        "A Figura 6 é a modalidade em 2025. Pregão eletrônico 559, dispensa 448, no total de 1.215 "
        "processos. Quem só olha concorrência grande deixa passar o miolo, que está no eletrônico e na "
        "dispensa (BRASIL, 2021). Quem mais solicita: DMAE com 333, Saúde com 233."
    ),
    "disc1": (
        "Discutindo o que isso mostra, o recorte municipal de 2025 concentra volume em pregão "
        "eletrônico e em dispensa, e em dois solicitantes (água e esgoto, saúde). O mapa da Figura 2 "
        "puxa o olhar para fora da cidade: a maior parte dos itens homologados no recorte da tela não "
        "tem sede em Uberlândia. A tela de CNPJs vencedores atende o pedido de 11 de junho, com o "
        "limite que eu já tinha falado: a API federal só devolve resultado classificado ou homologado, "
        "não a lista inteira de quem participou. Então o estudo de rodízio fica pela metade. O sistema "
        "não aponta irregularidade sozinho. Organiza consulta. A rotina de preço com modelo de "
        "linguagem é rascunho. O servidor refaz a conta com faixa de 15%. Não substitui pesquisa "
        "formal nem nota fiscal."
    ),
    "ev6": (
        "Fora os recados já transcritos, teve o convite para ser voluntário, o pedido do link no site "
        "e as duas apresentações formais. Nas duas a ferramenta foi bem recebida. Na segunda pediram "
        "filtro por porte, e isso entrou na tela de CNPJs vencedores. Não medi com relógio quanto tempo "
        "eles economizaram. O que eles relatam é que o trabalho se rearranjou. A frente dos vereadores "
        "eu deixei de fora de propósito. Outra coisa que fica aberta: quando acabar o plano gratuito "
        "da AWS, em janeiro de 2027, o sistema como está hoje sai do ar, a menos que a entidade pague "
        "ou mude de hospedagem. Nas apresentações o ganho não foi só a tela. Diretoria e sede viram "
        "que existe API de transparência para operação pública, e que o sistema a usa. Esse tipo de "
        "conhecimento entra indireto na extensão, junto com o objetivo principal."
    ),
    "ev7": (
        "Esta parte não descreve o problema da entidade. O problema da entidade era a coleta manual na "
        "planilha Cronograma. O que segue é só um exemplo, ocorrido em agosto, de um risco passível da "
        "arquitetura: o módulo Compras.gov consome a API de Dados Abertos; se essa API deixa de "
        "devolver o registro, aquele módulo para de atualizar, mesmo com a compra visível no portal. "
        "Para documentar o caso, em 7 de agosto de 2026 abri o chamado 56398424 no Portal de Serviços. "
        "A contratação 18431312000620-1-000341/2026, da UASG 926922, Pregão Eletrônico, Lei 14.133, "
        "estava no PNCP desde 27 de julho. No portal Comprasnet "
        "(https://cnetmobile.estaleiro.serpro.gov.br/comprasnet-web/public/compras) a mesma compra "
        "aparece no acompanhamento 92692205002262026. Consulta na API do PNCP: HTTP 200. Consulta na "
        "API Dados Abertos: totalRegistros=0. Os sequenciais 000333 e 000334 voltavam normalmente. Em "
        "13 de agosto cobrei o andamento. Em 17 de agosto o SIASG informou que o caso ainda estava em "
        "análise. Serve para mostrar o limite, não para redefinir o diagnóstico do projeto. Serviu "
        "também, na prática, para a equipe entender o que é API: não é o portal que a pessoa vê; é o "
        "canal que o sistema consome. Se esse canal falha, a tela humana continua e a coleta "
        "automática não."
    ),
    "conc1": (
        "O que ficou de pé é um sistema de informação de apoio, no ar, no domínio da entidade, em "
        "https://licitacoes.osbrasiluberlandia.org/. A intervenção atendeu ao que o Observatório pediu: "
        "reunir fonte oficial, consultar sem começar pela cópia para a planilha Cronograma e olhar "
        "vencedor com algum recorte (mapa, CNPJ, porte). A análise continua sendo deles. O software "
        "não denuncia irregularidade sozinho. Na base municipal de 2025 entram 1.215 processos; o miolo "
        "está no pregão eletrônico e na dispensa. Gestores e fiscais passam de 36 mil linhas, volume "
        "que a planilha não dava conta de conferir. O retorno da entidade, nas reuniões, no WhatsApp e "
        "nas duas apresentações, foi de uso e de pedido de ajuste, não de recusa da ferramenta."
    ),
    "conc2": (
        "Houve um ganho que não estava no objetivo inicial. A equipe passou a saber o que é uma API e "
        "que o governo já publica, por esse canal, dado aberto para transparência das compras, em "
        "especial a API de Dados Abertos do Compras.gov.br. Isso entrou no convívio, não numa aula. "
        "Também ficou explícito o limite: se essa API para, só o módulo Compras.gov deixa de atualizar. "
        "O restante segue com o CSV municipal. Outro limite, este de hospedagem: o plano gratuito da "
        "AWS acaba em 13 de janeiro de 2027, ou antes se o crédito acabar. Sem pagamento ou mudança de "
        "máquina, o sistema sai do ar."
    ),
    "conc3": (
        "Para desdobramento, combinamos deixar o voluntariado mais formal depois que a faculdade "
        "acabar, para não virar obrigação de curso interno da rede. A planilha Cronograma ainda pode "
        "ficar do lado enquanto eles pegam confiança; o critério prático de sucesso é ela ser usada "
        "cada vez menos. Ficou em aberto medir, com a rotina já apoiada, o tempo de preparação do "
        "quadrimestre. A frente dos vereadores continua fora, de propósito: outro recorte, outra fonte. "
        "Se a API federal um dia devolver a lista completa de proponentes, o estudo de rodízio que o "
        "Marco pediu em 11 de junho deixa de ficar pela metade. Enquanto isso, o que cabe à entidade é "
        "decidir a hospedagem depois de janeiro de 2027. O endereço local de desenvolvimento, na minha "
        "máquina, permanece o mesmo: docker compose up --build -d, http://localhost:8096/."
    ),
}

Q1_ORG = [
    ["Nome", "Observatório Social do Brasil de Uberlândia"],
    ["CNPJ", "23.497.346/0001-42"],
    ["Natureza", "OSC apartidária, sem fins lucrativos"],
    ["CNAE principal", "94.99-5-00  Atividades associativas"],
    ["Contato operacional", "Marco Aurélio Freitas Santos  (34) 9979-6169"],
    ["Ponte institucional", "Lécia Queiroz (CDL Uberlândia)"],
    ["Site", "https://www.osbrasiluberlandia.org/"],
    ["Sistema de apoio", "https://licitacoes.osbrasiluberlandia.org/"],
]
Q2_UASG = [
    ["PMU", "Prefeitura Municipal de Uberlândia", "926922"],
    ["DMAE", "Departamento Municipal de Água e Esgoto", "926287"],
    ["FUTEL", "Fundação Uberlandense do Turismo, Esporte e Lazer", "926038"],
    ["ARESAN", "Agência de Regulação dos Serviços de Saneamento", "931351"],
    ["FERUB", "Fundação de Excelência Rural de Uberlândia", "930403"],
    ["EMAM", "Empresa Municipal de Apoio e Manutenção", "929315"],
    ["IPREMU", "Instituto de Previdência dos Servidores Municipais", "929301"],
    ["CAM", "Câmara Municipal de Uberlândia", "925010"],
]
Q3_INT = [
    ["16/03/2026", "Conversa com Antonio Carlos Oliveira, meu ex-professor de Administração na ESAMC. Pedi orientação; ele sugeriu procurar a Lécia na CDL."],
    ["20/05/2026", "Reunião na CDL com Lécia Queiroz, marcada por mim. Encaminhamento ao OSB e contato com Marco Aurélio."],
    ["21/05/2026", "Primeira reunião no Observatório, 14h30. Apresentação da entidade."],
    ["10/06/2026", "Reunião com datashow. Mostrei o que já estava andando e comecei a explicar o que é API e que já existe API federal de transparência das compras."],
    ["11/06/2026", "Pedido de estatística de vencedores. Convite a voluntariado. Pedido de link no site."],
    ["17/06/2026", "Convite para a prestação de contas do 1o quadrimestre de 2026 (Teams)."],
    ["19 a 25/06/2026", "Conversa sobre hospedagem gratuita. O sistema já existia (CSV municipal). A API Compras.gov entrou depois; para a equipe, o conceito de API ainda era novo."],
    ["15/07/2026", "Retorno da Lécia: o Marco estava contente com os resultados."],
    ["23/07/2026", "Apresentação na ACIUB, 15h30 às 16h. Além da ferramenta, mostrei a API de Dados Abertos do Compras.gov como fonte oficial de transparência, pouco conhecida na plateia."],
    ["03/08/2026", "Reunião no Teams com a sede do OSB, 16h às 17h30. Pediram filtro por porte. Expliquei de novo o que é API e que essa API de transparência já existia."],
    ["07 a 17/08/2026", "Exemplo de risco da arquitetura (não do problema da entidade): a API de Dados Abertos parou de devolver registro e o módulo Compras.gov deixou de atualizar. Chamado 56398424."],
]
Q4_MET = [
    ["Natureza da ação", "Desenvolvimento e implantação de sistema de informação de apoio (não foi curso)."],
    ["Período com a entidade", "21/05/2026 a 03/08/2026 (convívio semanal); articulação desde 16/03/2026."],
    ["Forma de oferta", "Software em produção no domínio da entidade; reuniões na ACIUB; duas apresentações formais."],
    ["Apresentação ACIUB", "23/07/2026, 15h30 às 16h (cerca de 30 min), presencial, diretoria / CDL / ACIUB."],
    ["Apresentação sede", "03/08/2026, 16h às 17h30 (1h30), Microsoft Teams."],
    ["Material", "O próprio sistema (telas e base local). Sem apostila. Artigo enviado à RETII (ISSN 2966-2508)."],
    ["Usuários", "Até quatro contas: um administrador e três de consulta."],
    ["Hospedagem", "VM AWS, São Paulo, Ubuntu, 1 GB RAM, 2 vCPUs, 40 GB. Free tier até 13/01/2027."],
]
Q5_MOD = [
    ["Painel", "Quantidade e valor por situação, órgão e modalidade, nas duas fontes, com filtro de período."],
    ["Cobertura entre bases", "O que está no Compras.gov e não no painel municipal, e o inverso, pela chave órgão + ano + processo."],
    ["Consulta por processo", "Junta o que houver nas duas bases sobre o mesmo número de processo."],
    ["CNPJs vencedores", "Lista de quem ganhou (CNPJ ou CPF): itens, compras, valor homologado, porte, CNAE e município da sede. Filtro por porte pedido pela equipe em 03/08."],
    ["Mapa de localidade", "Sede dos vencedores das licitações (itens homologados). Filtros de período, órgão, modalidade, UF, porte e métrica. Cards com totais e mapa do Brasil, com ranking por UF."],
    ["Propostas abertas", "Itens com prazo PNCP vigente e rotina auxiliar de preço (rascunho, não pesquisa formal)."],
    ["Coleta e Setup", "Orquestração das fontes, agendamento noturno, UASGs, usuários e backup (só admin). A ponta Compras.gov usa a API de Dados Abertos. Se a API para, só essa coleta para."],
]
Q6_CHAMADO = [
    ["Chamado", "56398424 (chave 94441)"],
    ["Portal", "https://portaldeservicos.gestao.gov.br/"],
    ["Abertura", "07/08/2026, 13h14"],
    ["Contratação PNCP", "18431312000620-1-000341/2026"],
    ["Órgão / UASG", "Município de Uberlândia / 926922"],
    ["Modalidade", "Pregão Eletrônico (Lei 14.133, art. 28, I)"],
    ["Divulgação no PNCP", "27/07/2026"],
    ["Portal Comprasnet", "Aparece normalmente (pesquisa que a entidade já usa)"],
    ["Acompanhamento", "https://cnetmobile.estaleiro.serpro.gov.br/comprasnet-web/public/compras/acompanhamento-compra?compra=92692205002262026"],
    ["API PNCP", "HTTP 200, registro presente"],
    ["API Dados Abertos", "HTTP 200, totalRegistros=0"],
    ["Lacuna observada", "Sequenciais 000335 em diante ausentes; último ok: 000334 (23/07/2026)"],
    ["Efeito no sistema", "Só o módulo Compras.gov deixa de atualizar (risco da API, não da rotina da entidade)"],
    ["Situação em 17/08/2026", "Em análise (SIASG 3o Nível)"],
]
TAB1 = [
    ["Processos licitatórios", "1.215", "628", "Campos oficiais do painel"],
    ["Contratos (linhas)", "2.152", "556", "Inclui aditivos/parcelas na origem"],
    ["Gestores e fiscais", "n/a", "36.399", "Base acumulada, não anualizada"],
]
# Referencias NBR 6023:2018.
# DOI so quando Crossref/editor confirma a mesma obra citada.
# Livro comercial sem DOI: ISBN da edicao citada (nao inventar DOI).
# Lei, norma ABNT e portal: URL oficial, sem DOI.
REFS = [
    "<b>ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS.</b> <i>NBR 14724: informação e documentação: trabalhos acadêmicos: apresentação.</i> Rio de Janeiro: ABNT, 2011.",
    "<b>ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS.</b> <i>NBR 6023: informação e documentação: referências: elaboração.</i> Rio de Janeiro: ABNT, 2018.",
    "<b>ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS.</b> <i>NBR 10520: informação e documentação: citações em documentos.</i> Rio de Janeiro: ABNT, 2023.",
    "<b>BRASIL.</b> <i>Decreto n. 8.777, de 11 de maio de 2016.</i> Institui a Política de Dados Abertos do Poder Executivo federal. Diário Oficial da União: seção 1, Brasília, DF, 12 maio 2016. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2016/decreto/d8777.htm. Acesso em: 21 ago. 2026.",
    "<b>BRASIL.</b> <i>Lei n. 12.527, de 18 de novembro de 2011.</i> Regula o acesso a informações. Diário Oficial da União: seção 1, Brasília, DF, 18 nov. 2011. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2011/lei/l12527.htm. Acesso em: 21 ago. 2026.",
    "<b>BRASIL.</b> <i>Lei n. 14.133, de 1 de abril de 2021.</i> Lei de Licitações e Contratos Administrativos. Diário Oficial da União: seção 1, Brasília, DF, 1 abr. 2021. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm. Acesso em: 21 ago. 2026.",
    "<b>BRASIL.</b> <i>Portal de Serviços.</i> Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://portaldeservicos.gestao.gov.br/. Acesso em: 21 ago. 2026.",
    "<b>COMPRAS.GOV.BR.</b> <i>Acompanhamento de compras (Comprasnet).</i> Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://cnetmobile.estaleiro.serpro.gov.br/comprasnet-web/public/compras. Acesso em: 21 ago. 2026.",
    "<b>COMPRAS.GOV.BR.</b> <i>Compras públicas em dados abertos.</i> Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://dadosabertos.compras.gov.br/swagger-ui/index.html. Acesso em: 21 ago. 2026.",
    "<b>GIL, Antonio Carlos.</b> <i>Como elaborar projetos de pesquisa.</i> 7. ed. São Paulo: Atlas, 2022. ISBN 978-65-5977-163-9.",
    "<b>LAUDON, Kenneth C.; LAUDON, Jane P.</b> <i>Sistemas de informação gerenciais.</i> 11. ed. São Paulo: Pearson, 2014. ISBN 978-85-4300-585-0.",
    "<b>MOREIRA, Fábio Mosso et al.</b> A qualidade na recuperação de dados governamentais: um estudo sobre dados de políticas públicas na internet. <i>Perspectivas em Ciência da Informação</i>, Belo Horizonte, v. 25, n. 2, p. 103-132, abr./jun. 2020. DOI: https://doi.org/10.1590/1981-5344/3994. Acesso em: 21 ago. 2026.",
    "<b>OBSERVATÓRIO SOCIAL DO BRASIL DE UBERLÂNDIA.</b> <i>Página institucional.</i> Uberlândia, [2026]. Disponível em: https://www.osbrasiluberlandia.org/. Acesso em: 21 ago. 2026.",
    "<b>PINHO, José Antonio Gomes de; SACRAMENTO, Ana Rita Silva.</b> Accountability: já podemos traduzi-la para o português? <i>Revista de Administração Pública</i>, Rio de Janeiro, v. 43, n. 6, p. 1343-1368, nov./dez. 2009. DOI: https://doi.org/10.1590/S0034-76122009000600006. Acesso em: 21 ago. 2026.",
    "<b>PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS.</b> <i>API de consulta: contratação 18431312000620-1-000341/2026.</i> Brasília, DF: Governo Federal, 2026. Disponível em: https://pncp.gov.br/api/consulta/v1/orgaos/18431312000620/compras/2026/341. Acesso em: 21 ago. 2026.",
    "<b>PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS.</b> <i>Manuais do PNCP.</i> Brasília, DF: Governo Federal, [2026]. Disponível em: https://www.gov.br/pncp/pt-br/pncp/manuais. Acesso em: 21 ago. 2026.",
    "<b>PREFEITURA MUNICIPAL DE UBERLÂNDIA.</b> <i>Painel de licitações e contratos (dados abertos).</i> Uberlândia, [2026]. Disponível em: https://app.powerbi.com/. Acesso em: 21 ago. 2026.",
    "<b>PRESSMAN, Roger S.; MAXIM, Bruce R.</b> <i>Engenharia de software: uma abordagem profissional.</i> 9. ed. Porto Alegre: AMGH, 2021. ISBN 978-65-5804-010-1.",
    "<b>SILVA, Patrícia Nascimento.</b> Reúso de dados abertos do governo brasileiro: atualização da métrica DGABr. <i>Em Questão</i>, Porto Alegre, v. 30, e-138332, 2024. DOI: https://doi.org/10.1590/1808-5245.30.138332. Acesso em: 31 ago. 2026.",
    "<b>SOMMERVILLE, Ian.</b> <i>Engenharia de software.</i> 10. ed. São Paulo: Pearson, 2018. ISBN 978-85-4302-497-4.",
    "<b>YIN, Robert K.</b> <i>Estudo de caso: planejamento e métodos.</i> 5. ed. Porto Alegre: Bookman, 2015. ISBN 978-85-8260-231-7.",
]
REFS_PLAIN = [
    "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. NBR 14724: informação e documentação: trabalhos acadêmicos: apresentação. Rio de Janeiro: ABNT, 2011.",
    "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. NBR 6023: informação e documentação: referências: elaboração. Rio de Janeiro: ABNT, 2018.",
    "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. NBR 10520: informação e documentação: citações em documentos. Rio de Janeiro: ABNT, 2023.",
    "BRASIL. Decreto n. 8.777, de 11 de maio de 2016. Institui a Política de Dados Abertos do Poder Executivo federal. Diário Oficial da União: seção 1, Brasília, DF, 12 maio 2016. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2016/decreto/d8777.htm. Acesso em: 21 ago. 2026.",
    "BRASIL. Lei n. 12.527, de 18 de novembro de 2011. Regula o acesso a informações. Diário Oficial da União: seção 1, Brasília, DF, 18 nov. 2011. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2011/lei/l12527.htm. Acesso em: 21 ago. 2026.",
    "BRASIL. Lei n. 14.133, de 1 de abril de 2021. Lei de Licitações e Contratos Administrativos. Diário Oficial da União: seção 1, Brasília, DF, 1 abr. 2021. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm. Acesso em: 21 ago. 2026.",
    "BRASIL. Portal de Serviços. Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://portaldeservicos.gestao.gov.br/. Acesso em: 21 ago. 2026.",
    "COMPRAS.GOV.BR. Acompanhamento de compras (Comprasnet). Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://cnetmobile.estaleiro.serpro.gov.br/comprasnet-web/public/compras. Acesso em: 21 ago. 2026.",
    "COMPRAS.GOV.BR. Compras públicas em dados abertos. Brasília, DF: Ministério da Gestão e da Inovação em Serviços Públicos, [2026]. Disponível em: https://dadosabertos.compras.gov.br/swagger-ui/index.html. Acesso em: 21 ago. 2026.",
    "GIL, Antonio Carlos. Como elaborar projetos de pesquisa. 7. ed. São Paulo: Atlas, 2022. ISBN 978-65-5977-163-9.",
    "LAUDON, Kenneth C.; LAUDON, Jane P. Sistemas de informação gerenciais. 11. ed. São Paulo: Pearson, 2014. ISBN 978-85-4300-585-0.",
    "MOREIRA, Fábio Mosso et al. A qualidade na recuperação de dados governamentais: um estudo sobre dados de políticas públicas na internet. Perspectivas em Ciência da Informação, Belo Horizonte, v. 25, n. 2, p. 103-132, abr./jun. 2020. DOI: https://doi.org/10.1590/1981-5344/3994. Acesso em: 21 ago. 2026.",
    "OBSERVATÓRIO SOCIAL DO BRASIL DE UBERLÂNDIA. Página institucional. Uberlândia, [2026]. Disponível em: https://www.osbrasiluberlandia.org/. Acesso em: 21 ago. 2026.",
    "PINHO, José Antonio Gomes de; SACRAMENTO, Ana Rita Silva. Accountability: já podemos traduzi-la para o português? Revista de Administração Pública, Rio de Janeiro, v. 43, n. 6, p. 1343-1368, nov./dez. 2009. DOI: https://doi.org/10.1590/S0034-76122009000600006. Acesso em: 21 ago. 2026.",
    "PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS. API de consulta: contratação 18431312000620-1-000341/2026. Brasília, DF: Governo Federal, 2026. Disponível em: https://pncp.gov.br/api/consulta/v1/orgaos/18431312000620/compras/2026/341. Acesso em: 21 ago. 2026.",
    "PORTAL NACIONAL DE CONTRATAÇÕES PÚBLICAS. Manuais do PNCP. Brasília, DF: Governo Federal, [2026]. Disponível em: https://www.gov.br/pncp/pt-br/pncp/manuais. Acesso em: 21 ago. 2026.",
    "PREFEITURA MUNICIPAL DE UBERLÂNDIA. Painel de licitações e contratos (dados abertos). Uberlândia, [2026]. Disponível em: https://app.powerbi.com/. Acesso em: 21 ago. 2026.",
    "PRESSMAN, Roger S.; MAXIM, Bruce R. Engenharia de software: uma abordagem profissional. 9. ed. Porto Alegre: AMGH, 2021. ISBN 978-65-5804-010-1.",
    "SILVA, Patrícia Nascimento. Reúso de dados abertos do governo brasileiro: atualização da métrica DGABr. Em Questão, Porto Alegre, v. 30, e-138332, 2024. DOI: https://doi.org/10.1590/1808-5245.30.138332. Acesso em: 31 ago. 2026.",
    "SOMMERVILLE, Ian. Engenharia de software. 10. ed. São Paulo: Pearson, 2018. ISBN 978-85-4302-497-4.",
    "YIN, Robert K. Estudo de caso: planejamento e métodos. 5. ed. Porto Alegre: Bookman, 2015. ISBN 978-85-8260-231-7.",
]


def prepare_assets() -> None:
    EVID.mkdir(parents=True, exist_ok=True)
    mapping = {
        "04_fluxo.png": FIG_ARTIGO / "fig1_fluxo.png",
        "05_arquitetura.png": FIG_ARTIGO / "fig2_arquitetura.png",
        "06_modalidades_2025.png": FIG_ARTIGO / "fig3_modalidades_2025.png",
        "00_logo_osb.png": LOGO_OSB,
    }
    for dest_name, src in mapping.items():
        dest = EVID / dest_name
        if src.exists():
            shutil.copy2(src, dest)


def styles():
    base = getSampleStyleSheet()
    return {
        "cover_inst": ParagraphStyle(
            "CoverInst", parent=base["Normal"], fontName="Times-Bold",
            fontSize=13, leading=17, alignment=TA_CENTER, textColor=NAVY, spaceAfter=2,
        ),
        "cover_sub": ParagraphStyle(
            "CoverSub", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=16, alignment=TA_CENTER, spaceAfter=2,
        ),
        "cover_comp": ParagraphStyle(
            "CoverComp", parent=base["Normal"], fontName="Times-Italic",
            fontSize=11, leading=15, alignment=TA_CENTER, textColor=GRAY, spaceAfter=4,
        ),
        "cover_title": ParagraphStyle(
            "CoverTitle", parent=base["Normal"], fontName="Times-Bold",
            fontSize=16, leading=22, alignment=TA_CENTER, textColor=NAVY, spaceAfter=8,
        ),
        "cover_author": ParagraphStyle(
            "CoverAuthor", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=16, alignment=TA_CENTER, spaceAfter=2,
        ),
        "h1": ParagraphStyle(
            "H1", parent=base["Heading1"], fontName="Times-Bold",
            fontSize=13, leading=17, spaceBefore=14, spaceAfter=8, textColor=NAVY,
        ),
        "h2": ParagraphStyle(
            "H2", parent=base["Heading2"], fontName="Times-Bold",
            fontSize=12, leading=16, spaceBefore=10, spaceAfter=6,
            textColor=colors.HexColor("#2c5282"),
        ),
        "h3": ParagraphStyle(
            "H3", parent=base["Heading3"], fontName="Times-Bold",
            fontSize=11, leading=15, spaceBefore=8, spaceAfter=5,
            textColor=colors.HexColor("#2c5282"),
        ),
        "body": ParagraphStyle(
            "Body", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=18, alignment=TA_JUSTIFY, firstLineIndent=1.25 * cm, spaceAfter=8,
        ),
        "quote": ParagraphStyle(
            "Quote", parent=base["Normal"], fontName="Times-Italic",
            fontSize=11, leading=16, alignment=TA_JUSTIFY,
            leftIndent=1.2 * cm, rightIndent=0.8 * cm, spaceBefore=4, spaceAfter=4,
        ),
        "quote_src": ParagraphStyle(
            "QuoteSrc", parent=base["Normal"], fontName="Times-Roman",
            fontSize=10, leading=13, alignment=TA_RIGHT, leftIndent=1.2 * cm,
            spaceAfter=10, textColor=GRAY,
        ),
        "caption": ParagraphStyle(
            "Caption", parent=base["Normal"], fontName="Times-Bold",
            fontSize=10, leading=13, alignment=TA_CENTER, spaceBefore=6, spaceAfter=2,
        ),
        "fonte": ParagraphStyle(
            "Fonte", parent=base["Normal"], fontName="Times-Roman",
            fontSize=9, leading=12, alignment=TA_CENTER, spaceAfter=10, textColor=GRAY,
        ),
        "ref": ParagraphStyle(
            "Ref", parent=base["Normal"], fontName="Times-Roman",
            fontSize=11, leading=15, alignment=TA_JUSTIFY,
            leftIndent=1.25 * cm, firstLineIndent=-1.25 * cm, spaceAfter=8,
        ),
        "toc": ParagraphStyle(
            "Toc", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=20, alignment=TA_LEFT, spaceAfter=2,
        ),
        "resumo": ParagraphStyle(
            "Resumo", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=18, alignment=TA_JUSTIFY, firstLineIndent=0, spaceAfter=10,
        ),
        "palavras": ParagraphStyle(
            "Palavras", parent=base["Normal"], fontName="Times-Roman",
            fontSize=12, leading=18, alignment=TA_JUSTIFY, firstLineIndent=0, spaceBefore=8, spaceAfter=4,
        ),
        "cell": ParagraphStyle(
            "Cell", parent=base["Normal"], fontName="Times-Roman",
            fontSize=9, leading=12, alignment=TA_LEFT,
        ),
        "cell_c": ParagraphStyle(
            "CellC", parent=base["Normal"], fontName="Times-Roman",
            fontSize=9, leading=12, alignment=TA_CENTER,
        ),
        "cell_h": ParagraphStyle(
            "CellH", parent=base["Normal"], fontName="Times-Bold",
            fontSize=9, leading=12, alignment=TA_CENTER, textColor=colors.white,
        ),
    }


def P(text, style):
    return Paragraph(text, style)


def fitted_image(path: Path, max_w_cm: float, max_h_cm: float) -> Image:
    img = Image(str(path))
    iw, ih = float(img.imageWidth), float(img.imageHeight)
    max_w, max_h = max_w_cm * cm, max_h_cm * cm
    ratio = min(max_w / iw, max_h / ih)
    img.drawWidth = iw * ratio
    img.drawHeight = ih * ratio
    img.hAlign = "CENTER"
    return img


def abnt_table(headers, rows, col_widths, s):
    head = [P(h, s["cell_h"]) for h in headers]
    body = []
    for row in rows:
        line = []
        for i, val in enumerate(row):
            st = s["cell"] if i == 0 else s["cell_c"]
            line.append(P(str(val), st))
        body.append(line)
    t = Table([head] + body, colWidths=col_widths, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e0")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f7fb")]),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return t


def later_pages(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(NAVY)
    canvas.setLineWidth(0.6)
    canvas.line(3 * cm, A4[1] - 1.8 * cm, A4[0] - 2 * cm, A4[1] - 1.8 * cm)
    canvas.setFont("Times-Italic", 8)
    canvas.setFillColor(GRAY)
    canvas.drawString(
        3 * cm, A4[1] - 1.65 * cm,
        "UNIUBE  Relato academico: sistema de apoio ao OSB",
    )
    canvas.line(3 * cm, 1.8 * cm, A4[0] - 2 * cm, 1.8 * cm)
    canvas.setFont("Times-Roman", 8)
    canvas.drawString(3 * cm, 1.45 * cm, "Diôgo Ferreira Moura  RA 1030125-2")
    canvas.drawRightString(A4[0] - 2 * cm, 1.45 * cm, f"{doc.page}")
    canvas.restoreState()


def first_page(canvas, doc):
    canvas.saveState()
    canvas.restoreState()


def build_pdf() -> None:
    s = styles()
    story = []

    # ---- PRE-TEXTUAIS: capa ----
    story.append(Spacer(1, 1.6 * cm))
    story.append(P("UNIVERSIDADE DE UBERABA", s["cover_inst"]))
    story.append(P("UNIUBE", s["cover_sub"]))
    story.append(Spacer(1, 0.25 * cm))
    story.append(P("Curso de Sistemas de Informação", s["cover_sub"]))
    story.append(P("Polo Uberlândia", s["cover_sub"]))
    story.append(Spacer(1, 2.2 * cm))
    story.append(P(T["cover_comp"], s["cover_comp"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(P(T["cover_title"], s["cover_title"]))
    story.append(P(T["cover_sub"], s["cover_comp"]))
    story.append(Spacer(1, 2.4 * cm))
    story.append(P("Diôgo Ferreira Moura", s["cover_author"]))
    story.append(P("RA 1030125-2", s["cover_author"]))
    story.append(Spacer(1, 3.2 * cm))
    story.append(P(T["cover_city"], s["cover_author"]))
    story.append(P("Setembro de 2026", s["cover_author"]))
    story.append(PageBreak())

    # ---- PRE-TEXTUAIS: sumario ----
    story.append(P("SUMÁRIO", s["h1"]))
    story.append(P(T["toc1"], s["toc"]))
    story.append(P("1.1 Apresentação do tema", s["toc"]))
    story.append(P("1.2 Problema", s["toc"]))
    story.append(P("1.3 Justificativa", s["toc"]))
    story.append(P("1.4 Objetivos", s["toc"]))
    story.append(P(T["toc2"], s["toc"]))
    story.append(P(T["toc3"], s["toc"]))
    story.append(P("3.1 Natureza e percurso da ação", s["toc"]))
    story.append(P("3.2 Material", s["toc"]))
    story.append(P("3.3 Métodos", s["toc"]))
    story.append(P(T["toc4"], s["toc"]))
    story.append(P("4.1 Interação com a organização parceira", s["toc"]))
    story.append(P("4.2 Evidências da ação realizada", s["toc"]))
    story.append(P("4.3 Exemplo de risco da arquitetura", s["toc"]))
    story.append(P(T["toc5"], s["toc"]))
    story.append(P(T["toc6"], s["toc"]))
    story.append(PageBreak())

    # ---- PRE-TEXTUAIS: resumo ----
    story.append(P("RESUMO", s["h1"]))
    story.append(P(T["resumo"], s["resumo"]))
    story.append(P(T["palavras"], s["palavras"]))
    story.append(PageBreak())

    # ---- TEXTUAIS: introducao ----
    story.append(P("1 INTRODUÇÃO", s["h1"]))
    story.append(P(T["intro0"], s["body"]))
    story.append(P("1.1 Apresentação do tema", s["h2"]))
    if (EVID / "00_logo_osb.png").exists():
        story.append(fitted_image(EVID / "00_logo_osb.png", 7.2, 2.6))
        story.append(Spacer(1, 0.15 * cm))
    story.append(P(T["tema1"], s["body"]))
    story.append(P(T["tema2"], s["body"]))
    story.append(abnt_table(["Campo", "Informação"], Q1_ORG, [5.0 * cm, 11.0 * cm], s))
    story.append(P("Quadro 1. Identificação da organização parceira e do sistema de apoio", s["caption"]))
    story.append(P("Fonte: registros do projeto e material da entidade (2026).", s["fonte"]))

    story.append(P("1.2 Problema", s["h2"]))
    story.append(P(T["prob1"], s["body"]))
    story.append(P(T["prob2"], s["body"]))
    story.append(P(
        "As unidades compradoras que eles já acompanhavam no portal Comprasnet, e que orientei "
        "o recorte do sistema, são estas. O portal eles já usavam; a API de transparência dessas "
        "mesmas UASGs, não:",
        s["body"],
    ))
    story.append(abnt_table(["Sigla", "Órgão", "UASG"], Q2_UASG, [2.4 * cm, 10.8 * cm, 2.8 * cm], s))
    story.append(P("Quadro 2. Unidades compradoras acompanhadas em Uberlândia", s["caption"]))
    story.append(P("Fonte: rotina da entidade e API Compras.gov (2026).", s["fonte"]))

    story.append(P("1.3 Justificativa", s["h2"]))
    story.append(P(T["just1"], s["body"]))
    story.append(P(T["just2"], s["body"]))

    story.append(P("1.4 Objetivos", s["h2"]))
    story.append(P(T["obj1"], s["body"]))
    story.append(P(T["obj2"], s["body"]))

    # ---- TEXTUAIS: fundamentacao teorica ----
    story.append(P("2 FUNDAMENTAÇÃO TEÓRICA", s["h1"]))
    story.append(P(T["teo1"], s["body"]))
    story.append(P(T["teo2"], s["body"]))
    story.append(P(T["teo3"], s["body"]))
    story.append(P(T["teo4"], s["body"]))

    # ---- TEXTUAIS: material e metodos ----
    story.append(P("3 MATERIAL E MÉTODOS", s["h1"]))
    story.append(P(T["met0"], s["body"]))
    story.append(P("3.1 Natureza e percurso da ação", s["h2"]))
    story.append(P(T["met1"], s["body"]))
    story.append(P(T["met2"], s["body"]))
    story.append(P(T["met3"], s["body"]))
    story.append(abnt_table(["Item", "Descrição"], Q4_MET, [4.2 * cm, 11.8 * cm], s))
    story.append(P("Quadro 3. Síntese da execução da ação extensionista", s["caption"]))
    story.append(P("Fonte: registros do projeto (2026).", s["fonte"]))

    story.append(P("3.2 Material", s["h2"]))
    story.append(P(T["met4"], s["body"]))

    story.append(P("3.3 Métodos", s["h2"]))
    story.append(P(T["met5"], s["body"]))

    # ---- TEXTUAIS: resultados ----
    story.append(P("4 RESULTADOS", s["h1"]))
    story.append(P(T["res0"], s["body"]))
    story.append(P("4.1 Interação com a organização parceira", s["h2"]))
    story.append(P(T["int1"], s["body"]))
    story.append(P(T["int2"], s["body"]))
    story.append(P(T["fac1"], s["body"]))
    story.append(P(T["fac2"], s["body"]))
    story.append(P(T["q1"], s["quote"]))
    story.append(P(T["q1src"], s["quote_src"]))
    story.append(P(T["fac3"], s["body"]))
    story.append(P(T["q2"], s["quote"]))
    story.append(P(T["q2src"], s["quote_src"]))
    story.append(P(T["fac4"], s["body"]))
    story.append(P(T["fac5"], s["body"]))
    story.append(P(T["dif1"], s["body"]))
    story.append(P(T["dif1b"], s["body"]))
    story.append(P(T["dif2"], s["body"]))
    story.append(P(T["dif3"], s["body"]))
    story.append(abnt_table(["Data", "O que aconteceu"], Q3_INT, [3.2 * cm, 12.8 * cm], s))
    story.append(P("Quadro 4. Síntese da interação com a organização parceira", s["caption"]))
    story.append(P("Fonte: WhatsApp, reuniões do projeto e Portal de Serviços (2026).", s["fonte"]))

    story.append(P("4.2 Evidências da ação realizada", s["h2"]))
    story.append(P(T["sol0"], s["body"]))
    story.append(abnt_table(["Módulo", "Para que serve"], Q5_MOD, [4.0 * cm, 12.0 * cm], s))
    story.append(P("Quadro 5. Módulos do sistema de apoio e uso previsto", s["caption"]))
    story.append(P("Fonte: sistema de apoio implantado (2026).", s["fonte"]))

    story.append(P(T["ev1"], s["body"]))
    bloc = []
    if LOGIN_PNG.exists():
        bloc.append(fitted_image(LOGIN_PNG, 16.0, 10.0))
    bloc.append(P("Figura 1. Tela de acesso do sistema de apoio (login)", s["caption"]))
    bloc.append(P("Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P(T["ev1b"], s["body"]))
    bloc = []
    if MAPA_PNG.exists():
        bloc.append(fitted_image(MAPA_PNG, 16.2, 11.2))
    bloc.append(P("Figura 2. Mapa de localidade dos vencedores das licitações (itens homologados)", s["caption"]))
    bloc.append(P("Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P(T["ev1c"], s["body"]))
    bloc = []
    if CNPJS_PNG.exists():
        bloc.append(fitted_image(CNPJS_PNG, 16.2, 11.2))
    bloc.append(P("Figura 3. Tela de CNPJs vencedores das licitações (fornecedores homologados)", s["caption"]))
    bloc.append(P("Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P(T["ev2"], s["body"]))
    bloc = []
    if (EVID / "04_fluxo.png").exists():
        bloc.append(fitted_image(EVID / "04_fluxo.png", 16.2, 7.2))
    bloc.append(P("Figura 4. Rotina de acompanhamento antes e depois do sistema de apoio", s["caption"]))
    bloc.append(P("Fonte: elaborado pelo autor (2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P(T["ev3"], s["body"]))
    bloc = []
    if (EVID / "05_arquitetura.png").exists():
        bloc.append(fitted_image(EVID / "05_arquitetura.png", 16.2, 8.2))
    bloc.append(P("Figura 5. Arquitetura lógica do sistema de apoio", s["caption"]))
    bloc.append(P("Fonte: elaborado pelo autor (2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P(T["ev4"], s["body"]))
    story.append(abnt_table(
        ["Conjunto", "2025", "2026 (parcial)", "Observação"],
        TAB1,
        [4.2 * cm, 2.6 * cm, 3.4 * cm, 5.8 * cm],
        s,
    ))
    story.append(P("Tabela 1. Volume carregado a partir do painel municipal (corte em 18/08/2026)", s["caption"]))
    story.append(P("Fonte: CSVs oficiais do painel da PMU, consolidados pelo sistema (2026).", s["fonte"]))

    story.append(P(T["ev5"], s["body"]))
    bloc = []
    if (EVID / "06_modalidades_2025.png").exists():
        bloc.append(fitted_image(EVID / "06_modalidades_2025.png", 15.4, 8.0))
    bloc.append(P("Figura 6. Processos licitatórios de 2025 por modalidade (painel municipal)", s["caption"]))
    bloc.append(P("Fonte: CSVs oficiais do painel da PMU (2026).", s["fonte"]))
    story.append(KeepTogether(bloc))

    story.append(P(T["disc1"], s["body"]))
    story.append(P(T["ev6"], s["body"]))

    story.append(P("4.3 Exemplo de risco da arquitetura (não é o problema da entidade)", s["h2"]))
    story.append(P(T["ev7"], s["body"]))
    story.append(abnt_table(["Campo", "Registro"], Q6_CHAMADO, [5.0 * cm, 11.0 * cm], s))
    story.append(P("Quadro 6. Demonstração de falha da API externa (chamado 56398424)", s["caption"]))
    story.append(P("Fonte: Portal de Serviços, Comprasnet, APIs PNCP e Compras.gov (ago. 2026).", s["fonte"]))

    # ---- TEXTUAIS: conclusao ----
    story.append(P("5 CONCLUSÃO", s["h1"]))
    story.append(P(T["conc1"], s["body"]))
    story.append(P(T["conc2"], s["body"]))
    story.append(P(T["conc3"], s["body"]))

    # ---- POS-TEXTUAIS: referencias ----
    story.append(P("REFERÊNCIAS", s["h1"]))
    for r in REFS:
        story.append(P(r, s["ref"]))

    doc = SimpleDocTemplate(
        str(OUT_PDF), pagesize=A4,
        leftMargin=3.0 * cm, rightMargin=2.0 * cm, topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title="Relatorio Semana 5: elementos do relatorio ABNT",
        author="Diogo Ferreira Moura",
        subject="TCC UNIUBE Relato de experiencia do Projeto Integrado",
    )
    doc.build(story, onFirstPage=first_page, onLaterPages=later_pages)


def _set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _docx_p(doc, text, *, size=12, bold=False, italic=False, align="justify",
            first_line=1.25, space_after=8, space_before=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.5
    if first_line:
        pf.first_line_indent = Cm(first_line)
    p.alignment = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def _docx_caption(doc, text):
    return _docx_p(doc, text, size=10, bold=True, align="center", first_line=0, space_before=8, space_after=2)


def _docx_fonte(doc, text):
    return _docx_p(doc, text, size=9, align="center", first_line=0, space_after=10)


def _docx_img(doc, path: Path, width_cm: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def _docx_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _set_run_font(run, size=10, bold=False)
    doc.add_paragraph()
    return table


def build_docx() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    _docx_p(doc, "UNIVERSIDADE DE UBERABA", size=14, bold=True, align="center", first_line=0, space_after=2)
    _docx_p(doc, "UNIUBE", size=12, align="center", first_line=0, space_after=2)
    _docx_p(doc, "Curso de Sistemas de Informação", size=12, align="center", first_line=0, space_after=2)
    _docx_p(doc, "Polo Uberlândia", size=12, align="center", first_line=0, space_after=18)
    _docx_p(doc, T["cover_comp_docx"], size=12, italic=True, align="center", first_line=0, space_after=12)
    _docx_p(doc, T["cover_title"], size=16, bold=True, align="center", first_line=0, space_after=8)
    _docx_p(doc, T["cover_sub"], size=12, italic=True, align="center", first_line=0, space_after=24)
    _docx_p(doc, "Diôgo Ferreira Moura", size=12, align="center", first_line=0, space_after=2)
    _docx_p(doc, "RA 1030125-2", size=12, align="center", first_line=0, space_after=36)
    _docx_p(doc, T["cover_city"], size=12, align="center", first_line=0, space_after=2)
    _docx_p(doc, "Setembro de 2026", size=12, align="center", first_line=0, space_after=12)
    doc.add_page_break()

    _docx_p(doc, "SUMÁRIO", size=13, bold=True, align="left", first_line=0, space_after=12)
    for item in (
        T["toc1"],
        "1.1 Apresentação do tema",
        "1.2 Problema",
        "1.3 Justificativa",
        "1.4 Objetivos",
        T["toc2"],
        T["toc3"],
        "3.1 Natureza e percurso da ação",
        "3.2 Material",
        "3.3 Métodos",
        T["toc4"],
        "4.1 Interação com a organização parceira",
        "4.2 Evidências da ação realizada",
        "4.3 Exemplo de risco da arquitetura",
        T["toc5"],
        T["toc6"],
    ):
        _docx_p(doc, item, size=12, align="left", first_line=0, space_after=4)
    doc.add_page_break()

    _docx_p(doc, "RESUMO", size=13, bold=True, align="left", first_line=0, space_after=12)
    _docx_p(doc, T["resumo"], first_line=0)
    _docx_p(doc, T["palavras"], first_line=0, space_before=8)
    doc.add_page_break()

    _docx_p(doc, "1 INTRODUÇÃO", size=13, bold=True, align="left", first_line=0)
    _docx_p(doc, T["intro0"])
    _docx_p(doc, "1.1 Apresentação do tema", size=12, bold=True, align="left", first_line=0, space_before=10)
    if (EVID / "00_logo_osb.png").exists():
        _docx_img(doc, EVID / "00_logo_osb.png", 6.5)
    _docx_p(doc, T["tema1"])
    _docx_p(doc, T["tema2"])
    _docx_table(doc, ["Campo", "Informação"], Q1_ORG)
    _docx_caption(doc, "Quadro 1. Identificação da organização parceira e do sistema de apoio")
    _docx_fonte(doc, "Fonte: registros do projeto e material da entidade (2026).")

    _docx_p(doc, "1.2 Problema", size=12, bold=True, align="left", first_line=0, space_before=10)
    _docx_p(doc, T["prob1"])
    _docx_p(doc, T["prob2"])
    _docx_p(
        doc,
        "As unidades compradoras que eles já acompanhavam no portal Comprasnet, e que orientei "
        "o recorte do sistema, são estas. O portal eles já usavam; a API de transparência dessas "
        "mesmas UASGs, não:",
    )
    _docx_table(doc, ["Sigla", "Órgão", "UASG"], Q2_UASG)
    _docx_caption(doc, "Quadro 2. Unidades compradoras acompanhadas em Uberlândia")
    _docx_fonte(doc, "Fonte: rotina da entidade e API Compras.gov (2026).")

    _docx_p(doc, "1.3 Justificativa", size=12, bold=True, align="left", first_line=0, space_before=10)
    _docx_p(doc, T["just1"])
    _docx_p(doc, T["just2"])
    _docx_p(doc, "1.4 Objetivos", size=12, bold=True, align="left", first_line=0, space_before=10)
    _docx_p(doc, T["obj1"])
    _docx_p(doc, T["obj2"])

    _docx_p(doc, "2 FUNDAMENTAÇÃO TEÓRICA", size=13, bold=True, align="left", first_line=0, space_before=14)
    _docx_p(doc, T["teo1"])
    _docx_p(doc, T["teo2"])
    _docx_p(doc, T["teo3"])
    _docx_p(doc, T["teo4"])

    _docx_p(doc, "3 MATERIAL E MÉTODOS", size=13, bold=True, align="left", first_line=0, space_before=14)
    _docx_p(doc, T["met0"])
    _docx_p(doc, "3.1 Natureza e percurso da ação", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["met1"])
    _docx_p(doc, T["met2"])
    _docx_p(doc, T["met3"])
    _docx_table(doc, ["Item", "Descrição"], Q4_MET)
    _docx_caption(doc, "Quadro 3. Síntese da execução da ação extensionista")
    _docx_fonte(doc, "Fonte: registros do projeto (2026).")
    _docx_p(doc, "3.2 Material", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["met4"])
    _docx_p(doc, "3.3 Métodos", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["met5"])

    _docx_p(doc, "4 RESULTADOS", size=13, bold=True, align="left", first_line=0, space_before=14)
    _docx_p(doc, T["res0"])
    _docx_p(doc, "4.1 Interação com a organização parceira", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["int1"])
    _docx_p(doc, T["int2"])
    _docx_p(doc, T["fac1"])
    _docx_p(doc, T["fac2"])
    q = _docx_p(doc, T["q1"], italic=True, first_line=0)
    q.paragraph_format.left_indent = Cm(1.25)
    _docx_p(doc, T["q1src"], size=10, align="right", first_line=0)
    _docx_p(doc, T["fac3"])
    q = _docx_p(doc, T["q2"], italic=True, first_line=0)
    q.paragraph_format.left_indent = Cm(1.25)
    _docx_p(doc, T["q2src"], size=10, align="right", first_line=0)
    _docx_p(doc, T["fac4"])
    _docx_p(doc, T["fac5"])
    _docx_p(doc, T["dif1"])
    _docx_p(doc, T["dif1b"])
    _docx_p(doc, T["dif2"])
    _docx_p(doc, T["dif3"])
    _docx_table(doc, ["Data", "O que aconteceu"], Q3_INT)
    _docx_caption(doc, "Quadro 4. Síntese da interação com a organização parceira")
    _docx_fonte(doc, "Fonte: WhatsApp, reuniões do projeto e Portal de Serviços (2026).")

    _docx_p(doc, "4.2 Evidências da ação realizada", size=12, bold=True, align="left", first_line=0, space_before=8)
    _docx_p(doc, T["sol0"])
    _docx_table(doc, ["Módulo", "Para que serve"], Q5_MOD)
    _docx_caption(doc, "Quadro 5. Módulos do sistema de apoio e uso previsto")
    _docx_fonte(doc, "Fonte: sistema de apoio implantado (2026).")
    _docx_p(doc, T["ev1"])
    if LOGIN_PNG.exists():
        _docx_img(doc, LOGIN_PNG, 15.5)
    _docx_caption(doc, "Figura 1. Tela de acesso do sistema de apoio (login)")
    _docx_fonte(doc, "Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).")
    _docx_p(doc, T["ev1b"])
    if MAPA_PNG.exists():
        _docx_img(doc, MAPA_PNG, 16.0)
    _docx_caption(doc, "Figura 2. Mapa de localidade dos vencedores das licitações (itens homologados)")
    _docx_fonte(doc, "Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).")
    _docx_p(doc, T["ev1c"])
    if CNPJS_PNG.exists():
        _docx_img(doc, CNPJS_PNG, 16.0)
    _docx_caption(doc, "Figura 3. Tela de CNPJs vencedores das licitações (fornecedores homologados)")
    _docx_fonte(doc, "Fonte: captura de tela na instância local, porta 8096 (18 ago. 2026).")
    _docx_p(doc, T["ev2"])
    if (EVID / "04_fluxo.png").exists():
        _docx_img(doc, EVID / "04_fluxo.png", 16.0)
    _docx_caption(doc, "Figura 4. Rotina de acompanhamento antes e depois do sistema de apoio")
    _docx_fonte(doc, "Fonte: elaborado pelo autor (2026).")
    _docx_p(doc, T["ev3"])
    if (EVID / "05_arquitetura.png").exists():
        _docx_img(doc, EVID / "05_arquitetura.png", 16.0)
    _docx_caption(doc, "Figura 5. Arquitetura lógica do sistema de apoio")
    _docx_fonte(doc, "Fonte: elaborado pelo autor (2026).")
    _docx_p(doc, T["ev4"])
    _docx_table(doc, ["Conjunto", "2025", "2026 (parcial)", "Observação"], TAB1)
    _docx_caption(doc, "Tabela 1. Volume carregado a partir do painel municipal (corte em 18/08/2026)")
    _docx_fonte(doc, "Fonte: CSVs oficiais do painel da PMU, consolidados pelo sistema (2026).")
    _docx_p(doc, T["ev5"])
    if (EVID / "06_modalidades_2025.png").exists():
        _docx_img(doc, EVID / "06_modalidades_2025.png", 15.2)
    _docx_caption(doc, "Figura 6. Processos licitatórios de 2025 por modalidade (painel municipal)")
    _docx_fonte(doc, "Fonte: CSVs oficiais do painel da PMU (2026).")
    _docx_p(doc, T["disc1"])
    _docx_p(doc, T["ev6"])
    _docx_p(
        doc,
        "4.3 Exemplo de risco da arquitetura (não é o problema da entidade)",
        size=12, bold=True, align="left", first_line=0, space_before=8,
    )
    _docx_p(doc, T["ev7"])
    _docx_table(doc, ["Campo", "Registro"], Q6_CHAMADO)
    _docx_caption(doc, "Quadro 6. Demonstração de falha da API externa (chamado 56398424)")
    _docx_fonte(doc, "Fonte: Portal de Serviços, Comprasnet, APIs PNCP e Compras.gov (ago. 2026).")

    _docx_p(doc, "5 CONCLUSÃO", size=13, bold=True, align="left", first_line=0, space_before=14)
    _docx_p(doc, T["conc1"])
    _docx_p(doc, T["conc2"])
    _docx_p(doc, T["conc3"])

    _docx_p(doc, "REFERÊNCIAS", size=13, bold=True, align="left", first_line=0, space_before=14)
    for r in REFS_PLAIN:
        p = _docx_p(doc, r, size=11, first_line=0, space_after=8)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)

    doc.save(str(OUT_DOCX))


def assert_no_forbidden(path: Path) -> None:
    if path.suffix.lower() == ".pdf":
        from pypdf import PdfReader
        text = "".join((p.extract_text() or "") for p in PdfReader(str(path)).pages)
    elif path.suffix.lower() == ".docx":
        d = Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = "\n".join(parts)
    else:
        text = path.read_text(encoding="utf-8")
    hits = sorted({ch for ch in text if ch in FORBIDDEN})
    if hits:
        raise SystemExit(f"Ainda ha traco/aspas tipograficas em {path.name}: {hits!r}")


def main() -> None:
    src = Path(__file__).read_text(encoding="utf-8")
    for ch in FORBIDDEN:
        if ch in src:
            raise SystemExit(f"O gerador ainda contem o caractere {ch!r}. Tire antes de gerar.")
    prepare_assets()
    build_pdf()
    build_docx()
    assert_no_forbidden(OUT_PDF)
    assert_no_forbidden(OUT_DOCX)
    print("PDF:", OUT_PDF, OUT_PDF.stat().st_size)
    print("DOCX:", OUT_DOCX, OUT_DOCX.stat().st_size)


if __name__ == "__main__":
    main()
